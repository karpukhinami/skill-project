# -*- coding: utf-8 -*-
"""Автономное веб-приложение: извлечение данных из ФРП и кодификатора"""

import pandas as pd
import json
import re
import io
import os
import copy
import urllib.parse as _urlparse
import streamlit as st
import requests
from typing import Dict, List, Tuple, Optional
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

# ─── DB helpers ──────────────────────────────────────────────────────────────

def _clean_db_url(url: str) -> str:
    """Убирает channel_binding=require, который не поддерживает psycopg2."""
    parsed = _urlparse.urlparse(url)
    qs = _urlparse.parse_qs(parsed.query)
    qs.pop('channel_binding', None)
    new_q = _urlparse.urlencode({k: v[0] for k, v in qs.items()})
    return _urlparse.urlunparse(parsed._replace(query=new_q))


def get_db_conn():
    """Возвращает psycopg2-соединение или None."""
    try:
        import psycopg2
        url = os.environ.get('DATABASE_URL', '')
        if not url:
            return None
        return psycopg2.connect(_clean_db_url(url), connect_timeout=10)
    except Exception:
        return None


@st.cache_data(ttl=300, show_spinner=False)
def load_frp_topics_cached() -> pd.DataFrame:
    """Загружает frp_topics из БД и кэширует на 5 минут."""
    conn = get_db_conn()
    if conn is None:
        return pd.DataFrame(columns=['id', 'grade_class', 'subject_id', 'subject', 'section', 'topic', 'program'])
    try:
        df = pd.read_sql(
            "SELECT f.id, f.grade_class, f.subject_id, s.name AS subject, f.section, f.topic, f.program "
            "FROM frp_topics f "
            "JOIN subjects s ON f.subject_id = s.id "
            "ORDER BY f.grade_class, s.name, f.section, f.topic",
            conn
        )
        conn.close()
        return df
    except Exception:
        conn.close()
        return pd.DataFrame(columns=['id', 'grade_class', 'subject_id', 'subject', 'section', 'topic', 'program'])


@st.cache_data(ttl=300, show_spinner=False)
def load_subjects_cached() -> pd.DataFrame:
    """Загружает subjects из БД и кэширует на 5 минут."""
    conn = get_db_conn()
    if conn is None:
        return pd.DataFrame(columns=['id', 'name', 'parent_id'])
    try:
        df = pd.read_sql(
            "SELECT id, name, parent_id FROM subjects WHERE is_archived = FALSE ORDER BY name",
            conn
        )
        conn.close()
        return df
    except Exception:
        conn.close()
        return pd.DataFrame(columns=['id', 'name', 'parent_id'])


@st.cache_data(ttl=60, show_spinner=False)
def load_view_data_cached(table: str) -> pd.DataFrame:
    """Загружает skill_defs или content_element_defs с JOIN frp_topics."""
    import psycopg2
    url = os.environ.get('DATABASE_URL', '')
    if not url:
        return pd.DataFrame()
    try:
        conn = psycopg2.connect(_clean_db_url(url), connect_timeout=10)
        with conn.cursor() as cur:
            cur.execute(f"""
                SELECT s.id,
                       s.label_normalized,
                       s.frp_label,
                       s.frp_topic_id,
                       COALESCE(sub.name,      '(без предмета)') AS subject,
                       COALESCE(f.grade_class, '—')              AS grade_class,
                       COALESCE(f.section,     '(без раздела)')  AS section,
                       COALESCE(f.topic,       '(без темы)')     AS topic
                FROM {table} s
                LEFT JOIN frp_topics f ON s.frp_topic_id = f.id
                LEFT JOIN subjects sub ON f.subject_id = sub.id
            """)
            rows = cur.fetchall()
            cols  = [d[0] for d in cur.description]
        conn.close()
        df = pd.DataFrame(rows, columns=cols)
        df['_grade_sort'] = pd.to_numeric(df['grade_class'], errors='coerce').fillna(99)
        df['_min_id'] = df.groupby(
            ['subject', 'grade_class', 'section', 'topic', 'frp_label']
        )['id'].transform('min')
        df = df.sort_values(
            ['subject', '_grade_sort', 'section', 'topic', '_min_id', 'id']
        ).drop(columns=['_grade_sort', '_min_id']).reset_index(drop=True)
        return df
    except Exception:
        return pd.DataFrame()



def normalize_db_text(s: str) -> str:
    """Нижний регистр, схлопнуть пробелы, убрать точки и пробелы в конце."""
    s = re.sub(r'\s+', ' ', str(s).strip().lower())
    s = re.sub(r'[\.\s]+$', '', s)
    return s


def split_into_sentences(text: str) -> List[str]:
    """
    Разбивает текст на предложения:
      - по переводу строки
      - по '. ЗАГЛАВНАЯ', если перед точкой НЕ одиночная заглавная буква (инициал)
    """
    sentences = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        parts = []
        last = 0
        for m in re.finditer(r'\.\s+(?=[А-ЯA-Z])', line):
            pos = m.start()
            before = line[:pos]
            # Проверяем: предшествует ли точке одиночная заглавная буква (инициал)?
            if re.search(r'(?:^|[\s.])[А-ЯA-Z]$', before):
                continue
            parts.append(line[last:pos])
            last = m.end() - 1  # начало с заглавной буквы
        parts.append(line[last:])
        for p in parts:
            p = p.strip().rstrip('.')
            if p:
                sentences.append(p)
    return sentences


def load_atomize_prompt(mode_type: str = 'skills') -> str:
    filename = 'atomize_skill.txt' if mode_type == 'skills' else 'atomize_content.txt'
    prompt_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'prompts', filename)
    try:
        with open(prompt_path, encoding='utf-8') as f:
            return f.read()
    except Exception:
        return "Разбей текст на атомарные единицы. Верни JSON: {\"atomic_skills\": [...]}"


def load_tag_prompt(step: int) -> str:
    filename = {
        1: "tag_prompt_1_extract.txt",
        2: "tag_prompt_2_normalize.txt",
        3: "tag_prompt_3_assign.txt",
    }.get(int(step), "tag_prompt_1_extract.txt")
    prompt_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "prompts", filename)
    try:
        with open(prompt_path, encoding="utf-8") as f:
            return f.read()
    except Exception:
        return ""

st.set_page_config(page_title="Извлечение ФРП", layout="wide")

# --- Вспомогательные функции для JSON ---

def get_json_info(data_dict):
    """Возвращает информацию о JSON: тип, предметы, классы."""
    if not data_dict:
        return {'type': '?', 'subjects': [], 'classes': [], 'count': 0}
    sample_key = next(iter(data_dict.keys()), '')
    item_type = 'навыки' if 'skill' in sample_key.lower() else 'содержание'
    subjects = set()
    classes = set()
    for item in data_dict.values():
        if isinstance(item, dict):
            s = item.get('subject', '')
            if s:
                subjects.add(str(s).strip())
            c = item.get('class', '')
            if c:
                classes.add(str(c).strip())
    return {
        'type': item_type,
        'subjects': sorted(subjects) if subjects else ['—'],
        'classes': sorted(classes, key=lambda x: int(x) if str(x).isdigit() else 0),
        'count': len(data_dict)
    }


# --- Логика извлечения ФРП (таблица) ---

def detect_table_structure(df):
    num_cols = len(df.columns)
    if num_cols < 5:
        return {'type': '4_columns', 'col_section': None, 'col_topic': 0,
                'col_hours': 1, 'col_content': 2, 'col_skills': 3}
    last_col = df.iloc[:, 4]
    non_empty = last_col.dropna()
    if len(non_empty) > 10:
        return {'type': '5_columns', 'col_section': 0, 'col_topic': 1,
                'col_hours': 2, 'col_content': 3, 'col_skills': 4}
    return {'type': '4_columns', 'col_section': None, 'col_topic': 0,
            'col_hours': 1, 'col_content': 2, 'col_skills': 3}


def extract_frp_from_df(df, subject, program):
    """Извлекает навыки и содержание из DataFrame (универсально для Excel и PDF)."""
    structure = detect_table_structure(df)
    
    header_row = None
    for idx in range(min(30, len(df))):
        row_text = ' '.join([str(x).lower() for x in df.iloc[idx] if pd.notna(x)])
        if 'содержан' in row_text or 'деятельност' in row_text:
            header_row = idx
            break
    if header_row is None:
        header_row = 0
    
    current_class = None
    for idx in range(header_row):
        col0_text = str(df.iloc[idx, 0])
        if col0_text.lower().startswith('федера'):
            continue
        match = re.search(r'(\d+)\s*класс', col0_text.lower())
        if match:
            current_class = match.group(1)
            break
    
    skills_list = []
    content_list = []
    current_section = ""
    last_content_item = None
    last_skills_item = None
    
    for idx in range(header_row + 1, len(df)):
        row = df.iloc[idx]
        col0 = row[0] if pd.notna(row[0]) else None
        col1 = row[1] if pd.notna(row[1]) else None
        col_content = row[structure['col_content']] if structure['col_content'] < len(row) and pd.notna(row[structure['col_content']]) else None
        col_skills = row[structure['col_skills']] if structure['col_skills'] < len(row) and pd.notna(row[structure['col_skills']]) else None
        
        if all(pd.isna(row)):
            continue
        if col0 and str(col0).lower().startswith('итого'):
            continue
        if col0 and str(col0).lower().startswith('федера'):
            continue
        if col0 and str(col0).isdigit() and not col1:
            continue
        
        if col0:
            col0_str = str(col0).strip()
            match = re.match(r'^(\d+)\s*класс$', col0_str, re.IGNORECASE)
            if match:
                new_class = match.group(1)
                if 5 <= int(new_class) <= 11:
                    current_class = new_class
                continue
        
        if structure['col_section'] is not None and col0 and 'раздел' in str(col0).lower() and '.' in str(col0):
            section_text = re.sub(r'(?i)раздел\s*\d+\.\s*', '', str(col0))
            current_section = section_text.strip()
            continue
        
        is_topic = False
        topic_code = ""
        topic_name = ""
        
        if structure['type'] == '4_columns':
            if col0 and col1:
                try:
                    int(str(col1).strip())
                    is_topic = True
                    topic_name = str(col0).strip()
                except ValueError:
                    pass
        else:
            if col0 and re.match(r'^\d+\.\d+', str(col0)):
                is_topic = True
                topic_code = str(col0).strip()
                topic_name = str(col1).strip() if col1 else ""
        
        if is_topic:
            if col_content:
                content_item = {
                    'code': topic_code, 'text': str(col_content).strip(),
                    'class': current_class, 'section': current_section,
                    'topic': topic_name, 'subject': subject,
                    'program': program, 'sources': ['фрп_планирование']
                }
                content_list.append(content_item)
                last_content_item = content_item
            if col_skills:
                skills_item = {
                    'code': topic_code, 'text': str(col_skills).strip(),
                    'class': current_class, 'section': current_section,
                    'topic': topic_name, 'subject': subject,
                    'program': program, 'sources': ['фрп_планирование']
                }
                skills_list.append(skills_item)
                last_skills_item = skills_item
            continue
        
        if not col0 and not col1:
            if col_content and last_content_item:
                last_content_item['text'] += " " + str(col_content).strip()
            if col_skills and last_skills_item:
                last_skills_item['text'] += " " + str(col_skills).strip()
    
    return skills_list, content_list


def extract_from_sheet(filename, sheet_name, program, file_content=None):
    """Извлекает навыки и содержание из листа Excel."""
    if file_content is not None:
        df = pd.read_excel(io.BytesIO(file_content), sheet_name=sheet_name, header=None)
    else:
        df = pd.read_excel(filename, sheet_name=sheet_name, header=None)
    return extract_frp_from_df(df, sheet_name, program)


def split_text(text):
    # Разбиваем только если после точки идет пробел и заглавная буква
    sentences = re.split(r'\.\s+(?=[А-ЯЁA-Z])', text)
    return [s.strip() for s in sentences if s.strip()]


# --- Логика извлечения из кодификатора ---

def clean_section_name(s):
    """Преобразует «По теме «...»» в «...» — убирает обёртку и кавычки."""
    if not s or not isinstance(s, str):
        return s
    m = re.search(r'По теме\s*[«"]([^»"]+)[»"]', s, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return s


def parse_codifier_sheet(df_raw, subject='', program=''):
    """Парсит вкладку кодификатора (навыки или содержание)."""
    results = []
    current_class = None
    current_section = None
    current_section_code = None
    last_valid_row = None

    for idx, row in df_raw.iterrows():
        col0 = row[0] if 0 in row else None
        col1 = row[1] if 1 in row else None

        if pd.isna(col1):
            if pd.notna(col0) and 'класс' in str(col0).lower():
                match = re.search(r'(\d+)', str(col0))
                if match:
                    current_class = match.group(1)
            continue

        text_value = str(col1).strip()
        if not text_value or text_value == 'nan':
            continue
        if 'таблица' in text_value.lower() or 'результат' in text_value.lower() or 'содержан' in text_value.lower():
            continue

        if pd.isna(col0):
            if last_valid_row is not None:
                last_valid_row['text'] += " " + text_value
            continue

        col0_str = str(col0).strip()
        if col0_str.isdigit() and len(col0_str) <= 2:
            current_section_code = col0_str
            current_section = clean_section_name(text_value)
            continue

        if '.' in col0_str:
            new_row = {
                'code': col0_str,
                'text': text_value,
                'class': current_class,
                'section': current_section,
                'section_code': current_section_code,
                'subject': subject,
                'program': program,
                'sources': ['кодификатор'],
                'topic': ''
            }
            last_valid_row = new_row
            results.append(new_row)

    return results


def split_sentences(text):
    # Разбиваем только если после точки идет пробел и заглавная буква, или точка в конце строки
    sentences = re.split(r'\.\s+(?=[А-ЯЁA-Z])|\.$', text)
    return [s.strip() for s in sentences if s.strip()]


# --- Функции для сравнения и слияния JSON ---

def normalize_text_for_comparison(text):
    """Нормализация текста для сравнения: только удаление пунктуации в конце, регистр не меняем"""
    if not text:
        return ""
    # Удаляем знаки препинания в конце
    text = text.rstrip('.,;:!?')
    return text.strip()


def comprehensive_similarity(text1, text2):
    """
    Комплексная оценка похожести текстов.
    Возвращает словарь с метриками похожести.
    """
    from difflib import SequenceMatcher
    
    norm1 = normalize_text_for_comparison(text1)
    norm2 = normalize_text_for_comparison(text2)
    
    # Если после нормализации идентичны
    if norm1 == norm2:
        return {
            'similarity': 1.0,
            'method': 'exact_match',
            'normalized_match': True,
            'difference_type': 'none'
        }
    
    # SequenceMatcher для общего сходства
    seq_ratio = SequenceMatcher(None, norm1, norm2).ratio()
    
    # Токен-сходство (сравнение по словам) - коэффициент Сёренсена-Дайса
    tokens1 = set(norm1.split())
    tokens2 = set(norm2.split())
    token_intersection = len(tokens1 & tokens2)
    token_sum = len(tokens1) + len(tokens2)
    token_sim = (2 * token_intersection) / token_sum if token_sum > 0 else 0.0
    
    # Разница в количестве токенов
    token_diff = abs(len(tokens1) - len(tokens2))
    
    # Биграммы для учета порядка - коэффициент Сёренсена-Дайса
    def get_bigrams(text):
        return set(text[i:i+2] for i in range(len(text)-1))
    bigrams1 = get_bigrams(norm1)
    bigrams2 = get_bigrams(norm2)
    bigram_intersection = len(bigrams1 & bigrams2)
    bigram_sum = len(bigrams1) + len(bigrams2)
    bigram_sim = (2 * bigram_intersection) / bigram_sum if bigram_sum > 0 else 0.0
    
    # Взвешенная комбинация
    combined = (seq_ratio * 0.4 + token_sim * 0.4 + bigram_sim * 0.2)
    
    # Ограничиваем сверху значением 1.0 на всякий случай
    combined = min(combined, 1.0)
    
    # Определение типа различий
    if combined >= 0.95:
        diff_type = 'chars_only'  # Только различия в символах/буквах
    elif combined >= 0.85:
        diff_type = 'few_words'  # Одно-два слова отличаются
    else:
        diff_type = 'significant'  # Значительные различия
    
    return {
        'similarity': combined,
        'sequence_ratio': seq_ratio,
        'token_similarity': token_sim,
        'bigram_similarity': bigram_sim,
        'common_tokens': token_intersection,
        'total_tokens': token_sum,
        'token_diff': token_diff,
        'difference_type': diff_type
    }


def add_prefix_to_keys(data_dict, prefix):
    """
    Добавляет префикс ко всем ключам в словаре.
    
    Args:
        data_dict: словарь с записями
        prefix: префикс для добавления (например, 'frp_table_')
    
    Returns:
        Новый словарь с префиксами в ключах
    """
    return {f"{prefix}{key}": value for key, value in data_dict.items()}


def find_similar_records(target_record, base_data_dict, subject, class_num, section_filter=None):
    """
    Находит похожие записи в базовом наборе.
    
    Args:
        target_record: запись для сравнения
        base_data_dict: словарь базовых записей (с ключами!)
        subject: предмет для фильтрации
        class_num: класс для фильтрации
        section_filter: название раздела для фильтрации (опционально)
    
    Returns:
        Список из 3 наиболее похожих записей с метриками и ключами
    """
    target_text = target_record.get('text', '')
    if not target_text:
        return []
    
    candidates = []
    subject_norm = subject.strip().lower()
    class_num_str = str(class_num).strip() if class_num else ''
    skip_class_filter = not class_num_str or class_num_str == '0'
    
    for base_key, base_record in base_data_dict.items():
        if base_record.get('subject', '').strip().lower() != subject_norm:
            continue
        if not skip_class_filter and str(base_record.get('class', '')).strip() != class_num_str:
            continue
        
        # Фильтрация по разделу, если указан
        if section_filter:
            base_section = base_record.get('section', '').strip()
            if base_section and base_section != section_filter.strip():
                continue
        
        base_text = base_record.get('text', '')
        if not base_text:
            continue
        
        # Вычисление похожести
        similarity_data = comprehensive_similarity(target_text, base_text)
        similarity_data['record'] = base_record
        similarity_data['key'] = base_key  # Сохраняем ключ!
        candidates.append(similarity_data)
    
    # Сортировка по похожести и возврат топ-3
    candidates.sort(key=lambda x: x['similarity'], reverse=True)
    return candidates[:3]


def validate_json_source(data_dict, expected_source):
    """
    Проверяет, что в JSON файле есть записи с ожидаемым источником.
    
    Args:
        data_dict: словарь с записями
        expected_source: ожидаемый источник ('фрп_текст', 'фрп_планирование', 'кодификатор')
    
    Returns:
        (is_valid, message)
    """
    if not data_dict:
        return False, "Файл пустой"
    
    found_source = False
    for item in data_dict.values():
        if isinstance(item, dict):
            sources = item.get('sources', [])
            if sources and len(sources) > 0:
                first_source = str(sources[0]).strip()
                if first_source == expected_source:
                    found_source = True
                    break
    
    if not found_source:
        return False, f"В файле не найдено записей с источником '{expected_source}'. Проверьте правильность файла."
    
    return True, "Файл валиден"


def get_unique_sections(base_records, subject, class_num):
    """Получает уникальные разделы из базовых записей для заданного предмета и класса.
    При class_num=='0' берёт разделы из всех записей с данным предметом."""
    sections = set()
    subj_norm = subject.strip().lower()
    cls_norm = str(class_num).strip()
    for record in base_records:
        rsubj = (record.get('subject') or '').strip().lower()
        rcls = str(record.get('class', '')).strip() or '0'
        if rsubj == subj_norm and (cls_norm == '0' or rcls == cls_norm):
            section = (record.get('section') or '').strip()
            if section:
                sections.add(section)
    return sorted(list(sections))


def _check_and_transition_next_iteration():
    """Если все три словаря пусты и есть следующая итерация — переходим."""
    if not (st.session_state.compare_for_choice or st.session_state.compare_for_section_topic or st.session_state.compare_for_section_only):
        if st.session_state.compare_next_data:
            st.session_state.compare_iteration = 2
            st.session_state.compare_compare_data = copy.deepcopy(st.session_state.compare_next_data)
            st.session_state.compare_next_data = None
            compare_data = st.session_state.compare_compare_data
            fails = extract_fails_and_clean(compare_data)
            st.session_state.compare_fails.update(fails)
            st.session_state.compare_etalon_data = copy.deepcopy(st.session_state.compare_base_data)
            base_data = copy.deepcopy(st.session_state.compare_base_data)
            merged_data, for_choice, for_section_topic, for_section_only = process_comparison_iteration(
                base_data, compare_data,
                st.session_state.compare_report,
                st.session_state.compare_stats,
                etalon_data=st.session_state.compare_etalon_data,
                simple_mode=st.session_state.get('compare_simple_mode', False)
            )
            st.session_state.compare_base_data = merged_data
            st.session_state.compare_for_choice = for_choice
            st.session_state.compare_for_section_topic = for_section_topic
            st.session_state.compare_for_section_only = for_section_only
            if not (for_choice or for_section_topic or for_section_only):
                st.session_state.compare_merged_result = merged_data
            return True
        else:
            st.session_state.compare_merged_result = st.session_state.compare_base_data
            return True
    return False


def extract_fails_and_clean(compare_data):
    """
    Извлекает записи без предмета/класса/текста в отдельный словарь fails,
    удаляя их из compare_data.
    
    Returns:
        fails: dict {key: record} — записи с недостающими полями
    """
    fails = {}
    keys_to_remove = []
    for key, record in list(compare_data.items()):
        subject = (record.get('subject') or '').strip()
        class_val = str(record.get('class', '')).strip()
        text = (record.get('text') or '').strip()
        if not subject or not class_val or not text:
            fails[key] = record.copy()
            keys_to_remove.append(key)
    for k in keys_to_remove:
        del compare_data[k]
    return fails


def process_comparison_iteration(base_data, compare_data, report, stats, etalon_data=None, simple_mode=False):
    """
    Обрабатывает одну итерацию сравнения.
    Обработанные записи удаляются из compare_data.
    
    Args:
        base_data: рабочая копия для добавления/объединения (модифицируется в результате)
        compare_data: словарь записей для сравнения (модифицируется — обработанные удаляются)
        report: список для журнала отчёта
        stats: словарь статистики
        etalon_data: эталонный словарь для сопоставления (не изменяется). Если None — используется base_data.
    
    Returns:
        merged_data: объединённый словарь (результат слияния в base_data)
        for_choice: {compare_key: decision} — сличение двух и выбор (высокий порог сходства)
        for_section_topic: {compare_key: decision} — подбор раздела и темы (совпадения есть, но не близкие)
        for_section_only: {compare_key: decision} — выбор раздела (совпадений нет, раздел неизвестен)
    """
    etalon = etalon_data if etalon_data is not None else base_data
    merged_data = base_data.copy()
    for_choice = {}
    for_section_topic = {}  # В простом режиме остаётся пустым
    for_section_only = {}
    
    # Глобальный набор для отслеживания всех обработанных записей по нормализованному тексту
    global_processed_texts = set()
    
    # Группировка по классам и предметам (записи без класса → группа "0")
    classes_by_subject = {}
    for key, record in compare_data.items():
        subject = record.get('subject', '').strip()
        class_num = str(record.get('class', '')).strip()
        if not subject:
            continue
        if not class_num:
            class_num = '0'
        if subject not in classes_by_subject:
            classes_by_subject[subject] = set()
        classes_by_subject[subject].add(class_num)
    
    # Обработка по каждому предмету и классу
    for subject in classes_by_subject:
        for class_num in sorted(classes_by_subject[subject], key=lambda x: int(x) if x.isdigit() else 0):
            # Получаем записи текущего класса из compare_data (class_num "0" = пустой у записи)
            def _match(rec, subj, cls):
                rsubj = rec.get('subject', '').strip().lower()
                rcls = str(rec.get('class', '')).strip() or '0'
                return rsubj == subj.strip().lower() and rcls == cls
            compare_records = [
                (key, record) for key, record in compare_data.items()
                if _match(record, subject, class_num)
            ]
            
            # Отслеживаем уже обработанные записи по ключам, чтобы избежать дубликатов
            processed_keys = set()
            
            for compare_key, compare_record in compare_records:
                compare_text = compare_record.get('text', '')
                if not compare_text:
                    continue
                
                # Пропускаем, если эта запись уже обрабатывалась (по ключу)
                if compare_key in processed_keys or compare_key in global_processed_texts:
                    continue
                processed_keys.add(compare_key)
                global_processed_texts.add(compare_key)
                
                # Определяем фильтр по разделу
                compare_section = compare_record.get('section', '').strip()
                section_filter = None
                if compare_section and class_num != '0':
                    # Получаем список эталонных записей для проверки разделов
                    etalon_records_list = [
                        rec for rec in etalon.values()
                        if (rec.get('subject', '').strip().lower() == subject.strip().lower() and
                            str(rec.get('class', '')).strip() == class_num)
                    ]
                    base_sections = get_unique_sections(etalon_records_list, subject, class_num)
                    if compare_section.strip().lower() in [s.strip().lower() for s in base_sections]:
                        section_filter = compare_section
                
                # Сопоставляем с эталоном (не с рабочей копией)
                similar = find_similar_records(
                    compare_record, etalon, subject, class_num, section_filter
                )
                
                if not similar:
                    # Нет похожих записей: берём разделы из эталона (по предмету/классу, при пустоте — из всего эталона)
                    etalon_records_for_sections = [
                        r for r in etalon.values()
                        if (r.get('subject', '').strip().lower() == subject.strip().lower() and
                            (class_num == '0' or str(r.get('class', '')).strip() == class_num))
                    ]
                    base_sections_list = get_unique_sections(etalon_records_for_sections, subject, class_num) if etalon_records_for_sections else []
                    if not base_sections_list:
                        sc = {}
                        for r in etalon.values():
                            s = (r.get('section') or '').strip()
                            if s:
                                sc[s] = sc.get(s, 0) + 1
                        base_sections_list = sorted(sc.keys(), key=lambda x: (-sc.get(x, 0), x))
                    compare_section_val = compare_record.get('section', '').strip()
                    # Регистронезависимое совпадение раздела
                    section_matches = compare_section_val and any(
                        s.strip().lower() == compare_section_val.lower() for s in base_sections_list
                    )
                    if section_matches:
                        # Раздел совпадает — сохраняем запись как есть (тему оставляем или пусто)
                        new_record = compare_record.copy()
                        new_record['section'] = compare_section_val
                        new_record['topic'] = compare_record.get('topic', '') or ''
                        compare_sources = new_record.get('sources', [])
                        source_prefix = 'frp_text_' if (compare_sources and 'фрп_текст' in compare_sources) else 'codifier_'
                        max_num = 0
                        for k in merged_data.keys():
                            if k.startswith(source_prefix) and ('skill_' in k or 'content_' in k):
                                try:
                                    parts = k.split('_')
                                    if len(parts) >= 3 and parts[-1].isdigit():
                                        max_num = max(max_num, int(parts[-1]))
                                except Exception:
                                    pass
                        new_key = f"{source_prefix}skill_{max_num + 1:04d}"
                        merged_data[new_key] = new_record
                        report.append({
                            'action': 'no_match_section_ok',
                            'compare_key': compare_key,
                            'text': compare_text,
                            'section': compare_section_val,
                            'topic': new_record.get('topic', ''),
                            'note': 'Сохранено: раздел совпал с эталоном'
                        })
                        stats['section_assigned_auto'] = stats.get('section_assigned_auto', 0) + 1
                        if compare_key in compare_data:
                            del compare_data[compare_key]
                    else:
                        # Раздел неизвестен или не совпадает — на выбор пользователю (только раздел, тему не трогаем)
                        decision = {
                            'type': 'new_record',
                            'record': compare_record,
                            'compare_key': compare_key,
                            'similar_records': [],
                            'subject': subject,
                            'class': class_num,
                            'base_sections': base_sections_list
                        }
                        for_section_only[compare_key] = decision
                        if compare_key in compare_data:
                            del compare_data[compare_key]
                    continue
                
                best_match = similar[0]
                similarity = best_match['similarity']
                diff_type = best_match.get('difference_type', 'significant')
                base_record = best_match['record']
                
                # Проверяем полное совпадение текстов (100% или почти 100%)
                base_text_norm = normalize_text_for_comparison(base_record.get('text', ''))
                compare_text_norm = normalize_text_for_comparison(compare_text)
                is_exact_match = (base_text_norm == compare_text_norm) or (similarity >= 0.99)
                
                if is_exact_match or (diff_type == 'chars_only' and similarity >= 0.95):
                    # Автоматическое объединение
                    base_record = best_match['record']
                    base_key = best_match.get('key')  # Ключ уже есть в результатах поиска
                    
                    if base_key and base_key in merged_data:
                        # Объединяем источники
                        merged_record = merged_data[base_key].copy()
                        compare_sources = compare_record.get('sources', [])
                        base_sources = merged_record.get('sources', [])
                        merged_sources = list(set(base_sources + compare_sources))
                        merged_record['sources'] = merged_sources
                        
                        # При точном совпадении всегда присваиваем раздел/тему из эталонной записи
                        if base_record.get('section'):
                            merged_record['section'] = base_record.get('section')
                        if base_record.get('topic'):
                            merged_record['topic'] = base_record.get('topic')
                        
                        merged_data[base_key] = merged_record
                        
                        # Запись в журнал (compare_key нужен для точного подсчёта при одинаковых текстах)
                        report.append({
                            'action': 'auto_merge',
                            'compare_key': compare_key,
                            'base_text': base_record.get('text', ''),
                            'compare_text': compare_text,
                            'similarity': similarity,
                            'merged_sources': merged_sources,
                            'section': merged_record.get('section', ''),
                            'topic': merged_record.get('topic', '')
                        })
                        stats['auto_merged'] += 1
                        if compare_key in compare_data:
                            del compare_data[compare_key]
                    else:
                        # Точное совпадение найдено, но base_key не найден - добавляем как новую с автоматическим присвоением раздела/темы
                        new_record = compare_record.copy()
                        if base_record.get('section'):
                            new_record['section'] = base_record.get('section')
                        if base_record.get('topic'):
                            new_record['topic'] = base_record.get('topic')
                        # Объединяем источники
                        compare_sources = compare_record.get('sources', [])
                        base_sources = base_record.get('sources', [])
                        new_record['sources'] = list(set(base_sources + compare_sources))
                        
                        # Генерируем новый ключ с префиксом на основе источника
                        source_prefix = 'frp_text_' if 'фрп_текст' in compare_sources else 'codifier_'
                        # Находим максимальный номер среди существующих ключей с таким префиксом
                        max_num = 0
                        for k in merged_data.keys():
                            if k.startswith(source_prefix) and ('skill_' in k or 'content_' in k):
                                try:
                                    parts = k.split('_')
                                    if len(parts) >= 3 and parts[-1].isdigit():
                                        max_num = max(max_num, int(parts[-1]))
                                except:
                                    pass
                        
                        new_key = f"{source_prefix}skill_{max_num + 1:04d}"
                        merged_data[new_key] = new_record
                        
                        report.append({
                            'action': 'auto_merge_new',
                            'compare_key': compare_key,
                            'base_text': base_record.get('text', ''),
                            'compare_text': compare_text,
                            'similarity': similarity,
                            'section': new_record.get('section', ''),
                            'topic': new_record.get('topic', '')
                        })
                        stats['auto_merged'] += 1
                        if compare_key in compare_data:
                            del compare_data[compare_key]
                    
                    # После автоматического объединения переходим к следующей записи
                    continue
                
                elif diff_type == 'few_words' and similarity >= 0.85:
                    # Требуется выбор пользователя (сличение двух и выбора)
                    decision = {
                        'type': 'choice',
                        'compare_record': compare_record,
                        'compare_key': compare_key,
                        'base_record': best_match['record'],
                        'base_key': best_match.get('key'),
                        'similarity': similarity,
                        'similar_records': similar[:3],
                        'subject': subject,
                        'class': class_num
                    }
                    for_choice[compare_key] = decision
                    if compare_key in compare_data:
                        del compare_data[compare_key]
                
                else:
                    # Низкая похожесть
                    if simple_mode:
                        # Простое сравнение: обрабатываем как "нет совпадений" — раздел есть -> сохранить, иначе выбор раздела
                        etalon_records_for_sections = [
                            r for r in etalon.values()
                            if (r.get('subject', '').strip().lower() == subject.strip().lower() and
                                (class_num == '0' or str(r.get('class', '')).strip() == class_num))
                        ]
                        base_sections_list = get_unique_sections(etalon_records_for_sections, subject, class_num) if etalon_records_for_sections else []
                        if not base_sections_list:
                            sc = {}
                            for r in etalon.values():
                                s = (r.get('section') or '').strip()
                                if s:
                                    sc[s] = sc.get(s, 0) + 1
                            base_sections_list = sorted(sc.keys(), key=lambda x: (-sc.get(x, 0), x))
                        compare_section_val = compare_record.get('section', '').strip()
                        section_matches = compare_section_val and any(
                            s.strip().lower() == compare_section_val.lower() for s in base_sections_list
                        )
                        if section_matches:
                            new_record = compare_record.copy()
                            new_record['section'] = compare_section_val
                            new_record['topic'] = compare_record.get('topic', '') or ''
                            compare_sources = new_record.get('sources', [])
                            source_prefix = 'frp_text_' if (compare_sources and 'фрп_текст' in compare_sources) else 'codifier_'
                            max_num = 0
                            for k in merged_data.keys():
                                if k.startswith(source_prefix) and ('skill_' in k or 'content_' in k):
                                    try:
                                        parts = k.split('_')
                                        if len(parts) >= 3 and parts[-1].isdigit():
                                            max_num = max(max_num, int(parts[-1]))
                                    except Exception:
                                        pass
                            new_key = f"{source_prefix}skill_{max_num + 1:04d}"
                            merged_data[new_key] = new_record
                            report.append({
                                'action': 'no_match_section_ok',
                                'compare_key': compare_key,
                                'text': compare_text,
                                'section': compare_section_val,
                                'topic': new_record.get('topic', ''),
                                'note': 'Сохранено (простое сравнение): раздел совпал'
                            })
                            stats['section_assigned_auto'] = stats.get('section_assigned_auto', 0) + 1
                            if compare_key in compare_data:
                                del compare_data[compare_key]
                        else:
                            decision = {
                                'type': 'new_record',
                                'record': compare_record,
                                'compare_key': compare_key,
                                'similar_records': [],
                                'subject': subject,
                                'class': class_num,
                                'base_sections': base_sections_list
                            }
                            for_section_only[compare_key] = decision
                            if compare_key in compare_data:
                                del compare_data[compare_key]
                    else:
                        # Обычный режим: подбор раздела/темы на основе похожих записей
                        decision = {
                            'type': 'section_topic_choice',
                            'record': compare_record,
                            'compare_key': compare_key,
                            'similar_records': similar[:3],
                            'subject': subject,
                            'class': class_num
                        }
                        for_section_topic[compare_key] = decision
                        if compare_key in compare_data:
                            del compare_data[compare_key]
    
    return merged_data, for_choice, for_section_topic, for_section_only


def create_separate_elements(items, id_prefix):
    """Разбивает элементы на отдельные предложения."""
    result = {}
    counter = 1
    for item in items:
        sentences = split_sentences(item['text'])
        for sentence in sentences:
            new_item = item.copy()
            new_item['text'] = sentence
            result[f"{id_prefix}_{counter:04d}"] = new_item
            counter += 1
    return result


# --- Извлечение из PDF ---

def _table_to_df(table):
    """Список списков → DataFrame, выравнивание по максимальной ширине."""
    if not table:
        return pd.DataFrame()
    max_cols = max(len(r) for r in table)
    padded = [list(r) + [None] * (max_cols - len(r)) for r in table]
    return pd.DataFrame(padded)


def _detect_table_type(table):
    """Определяет тип: frp (4-5 кол), codifier (2 кол с кодами X.Y)."""
    if not table:
        return None
    flat = " ".join(str(c).lower() for row in table[:5] for c in row if c)
    n_col = max(len(r) for r in table) if table else 0
    if "содержан" in flat or "деятельност" in flat:
        return "frp"
    if n_col == 2 and any(
        c and "." in str(c) and str(c).split(".")[0].isdigit()
        for row in table[:15] for c in row if c
    ):
        return "codifier"
    if n_col >= 4 and ("содержан" in flat or "деятельност" in flat or "раздел" in flat):
        return "frp"
    return None


def _get_class_from_frp_table(table):
    """Находит класс в таблице ФРП (5, 6, ...)."""
    for row in table[:10]:
        for c in row:
            if c and re.search(r"(\d+)\s*класс", str(c), re.I):
                m = re.search(r"(\d+)", str(c))
                if m and 5 <= int(m.group(1)) <= 11:
                    return m.group(1)
    return None


def _get_class_from_codifier_table(table):
    """Кодификатор может содержать класс в заголовках."""
    for row in table[:5]:
        for c in row:
            if c and re.search(r"(\d+)\s*класс", str(c), re.I):
                m = re.search(r"(\d+)", str(c))
                if m:
                    return m.group(1)
    return None


def extract_and_merge_pdf_tables(pdf_bytes):
    """
    Извлекает таблицы из PDF, определяет тип, объединяет по классу.
    Возвращает: (doc_type, merged_dfs, stats)
    - doc_type: 'frp' | 'codifier'
    - merged_dfs: список (subject_or_class, df) для передачи в extract
    - stats: dict с таблицами, классами
    """
    import pdfplumber
    tables_by_page = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            for t in page.extract_tables():
                if t and len(t) > 0:
                    tables_by_page.append((page_num + 1, t))
    
    if not tables_by_page:
        return None, [], {"tables": 0, "classes": []}
    
    # Определяем тип по первой подходящей таблице
    doc_type = None
    for _, t in tables_by_page:
        dt = _detect_table_type(t)
        if dt:
            doc_type = dt
            break
    if not doc_type:
        doc_type = "frp" if any(max(len(r) for r in t) >= 4 for _, t in tables_by_page) else "codifier"
    
    # Объединение по классу: таблицы одного класса склеиваются
    merged = {}  # class -> list of dfs
    current_class = "_"
    for page_num, table in tables_by_page:
        dt = _detect_table_type(table)
        if dt != doc_type:
            continue
        df = _table_to_df(table)
        if df.empty or len(df) < 2:
            continue
        cls = _get_class_from_frp_table(table) if doc_type == "frp" else _get_class_from_codifier_table(table)
        if cls:
            current_class = cls
        if current_class not in merged:
            merged[current_class] = []
        merged[current_class].append(df)
    
    # Склеиваем df одного класса
    result = []
    for cls in sorted(merged.keys(), key=lambda x: int(x) if str(x).isdigit() else 999):
        dfs = merged[cls]
        combined = pd.concat(dfs, ignore_index=True)
        result.append((cls, combined))
    
    stats = {"tables": len(tables_by_page), "classes": list(merged.keys())}
    return doc_type, result, stats


# --- Извлечение из ФРП (текст doc/txt) ---

def _read_doc_or_txt(file_content, filename):
    """Читает текст из .txt или .docx (.doc не поддерживается — используйте .docx или .txt)."""
    name = (filename or "").lower()
    if name.endswith('.txt'):
        return file_content.decode('utf-8', errors='replace')
    if name.endswith('.docx'):
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        return "\n".join(p.text for p in doc.paragraphs)
    if name.endswith('.doc'):
        raise ValueError("Формат .doc не поддерживается. Сохраните файл как .docx или скопируйте в .txt")
    return file_content.decode('utf-8', errors='replace')


def extract_frp_from_text(text, subject, program):
    """
    Извлечение из текста ФРП (после заголовков «содержание обучения» и «предметные результаты»).
    Правила: N класс, Раздел (без слова), Тема (без слова). Разбивка по точкам или абзацам.
    """
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    lines = text.split('\n')
    
    content_start = None
    skills_start = None
    for i, line in enumerate(lines):
        low = line.strip().lower()
        if 'содержание обучения' in low:
            content_start = i + 1
        if 'предметные результаты' in low:
            skills_start = i + 1
            break
    
    def parse_block(start_idx, end_idx, is_content):
        items = []
        current_class = ""
        current_section = ""
        current_topic = ""
        buffer = []
        
        def flush_buffer():
            nonlocal buffer
            if not buffer:
                return
            combined = " ".join(buffer).strip()
            if not combined:
                buffer = []
                return
            if "." in combined:
                # Разбиваем только если после точки идет пробел и заглавная буква
                parts = re.split(r'\.\s+(?=[А-ЯЁA-Z])', combined)
                for p in parts:
                    p = p.strip()
                    if not p:
                        continue
                    if not p.endswith('.'):
                        p += "."
                    items.append({
                        'code': '', 'text': p,
                        'class': current_class, 'section': current_section,
                        'topic': current_topic, 'subject': subject,
                        'program': program, 'sources': ['фрп_текст']
                    })
            else:
                for para in buffer:
                    para = para.strip()
                    if para:
                        items.append({
                            'code': '', 'text': para,
                            'class': current_class, 'section': current_section,
                            'topic': current_topic, 'subject': subject,
                            'program': program, 'sources': ['фрп_текст']
                        })
            buffer = []
        
        for i in range(start_idx, end_idx if end_idx is not None else len(lines)):
            line = lines[i]
            stripped = line.strip()
            if not stripped:
                flush_buffer()
                continue
            
            low = stripped.lower()
            if low.startswith('федеральная программа') or low.startswith('федеральная рабочая программа'):
                continue
            if re.match(r'^\d+\s*$', stripped):
                continue
            
            class_match = re.match(r'^(\d+)\s*класс\s*\.?\s*$', low, re.I)
            if class_match:
                flush_buffer()
                current_class = class_match.group(1)
                continue
            
            if re.match(r'^раздел\b', low):
                flush_buffer()
                current_topic = ""
                rest = re.sub(r'^раздел\s*[.:;\-–—\s]*', '', stripped, flags=re.I).strip()
                rest = re.sub(r'^[.:;\-–—\s]+', '', rest).strip()
                if rest:
                    current_section = rest
                continue
            
            if re.match(r'^тема\b', low):
                flush_buffer()
                rest = re.sub(r'^тема\s*[.:;\-–—\s]*', '', stripped, flags=re.I).strip()
                rest = re.sub(r'^[.:;\-–—\s]+', '', rest).strip()
                if rest:
                    current_topic = rest
                continue
            
            buffer.append(stripped)
        
        flush_buffer()
        return items
    
    content_items = []
    skills_items = []
    
    if content_start is not None:
        end = skills_start if skills_start is not None else len(lines)
        content_items = parse_block(content_start, end, True)
    
    if skills_start is not None:
        skills_items = parse_block(skills_start, None, False)
    
    return content_items, skills_items


def json_to_excel_sorted(data_dict, columns_title='Содержание'):
    """Преобразует JSON в Excel с листами по предметам. Сортировка: класс → раздел → тема → порядок по ID."""
    # Собираем строки с предметом
    rows = []
    for key, item in data_dict.items():
        subject = item.get('subject', '') or ''
        rows.append({
            'subject': subject,
            'Класс': item.get('class', ''),
            'Раздел': item.get('section', ''),
            'Тема': item.get('topic', ''),
            columns_title: item.get('text', ''),
            '_sort_key': key
        })
    
    df = pd.DataFrame(rows)
    df['subject'] = df['subject'].fillna('').astype(str).str.strip()
    df.loc[df['subject'] == '', 'subject'] = 'Общее'  # пустые — в лист "Общее"
    
    subjects = sorted(df['subject'].unique())
    
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        for subject in subjects:
            sub_df = df[df['subject'] == subject].copy()
            sub_df = sub_df.drop(columns=['subject'])
            sub_df = sub_df.sort_values(['Класс', 'Раздел', 'Тема', '_sort_key'], na_position='last')
            sub_df = sub_df.drop(columns=['_sort_key'])
            sheet_name = str(subject)[:31].replace('/', '-').replace('\\', '-').replace('*', '').replace('?', '').replace('[', '').replace(']', '').replace(':', '-')
            sub_df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    buf.seek(0)
    return buf.getvalue()


# --- Функции для работы с LLM ---

def _load_openrouter_models_from_file() -> List[Dict]:
    """
    Читает список моделей из new_models.txt.

    Формат файла:
      MODELS = [
        ("vendor/model", input_price_per_token_usd, output_price_per_token_usd),
      ]

    В app.py цены хранятся в USD за 1M токенов (как и раньше),
    поэтому конвертируем: price_per_1m = price_per_token * 1_000_000.
    """
    import ast

    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "new_models.txt")
    try:
        raw = Path(path).read_text(encoding="utf-8")
    except Exception:
        return []

    m = re.search(r"MODELS\s*=\s*(\[[\s\S]*?\])\s*$", raw.strip())
    if not m:
        return []

    try:
        items = ast.literal_eval(m.group(1))
    except Exception:
        return []

    models: List[Dict] = []
    for model_id, in_per_tok, out_per_tok in items:
        model_id = str(model_id).strip()
        if not model_id:
            continue

        short = model_id.split("/", 1)[1] if "/" in model_id else model_id
        models.append(
            {
                "key": model_id,
                "display": short,
                "model_id": model_id,
                "provider": "openrouter",
                "in_price": float(in_per_tok) * 1_000_000,
                "out_price": float(out_per_tok) * 1_000_000,
            }
        )

    return models


# Список доступных моделей
AVAILABLE_MODELS = [
    {
        'key':       'claude_direct',
        'display':   'Клод. Просто клод',
        'model_id':  None,  # определяется из claude_working_model
        'provider':  'anthropic',
        'in_price':  3.0,
        'out_price': 15.0,
    },
]

AVAILABLE_MODELS.extend(_load_openrouter_models_from_file())

_MODEL_BY_KEY = {m['key']: m for m in AVAILABLE_MODELS}


def get_claude_api_key():
    """Получает API ключ Claude из секретов Streamlit или session_state."""
    return (st.secrets.get("CLAUDE_API_KEY", "")
            or st.session_state.get('claude_api_key', ''))

def test_claude_api_key(api_key: str, verify_ssl: bool = True) -> Dict:
    """
    Тестирует API ключ Claude и возвращает информацию о доступных моделях и параметрах.
    Возвращает словарь с результатами проверки.
    """
    if not api_key:
        return {'valid': False, 'error': 'API ключ не введен'}
    
    # Список версий API для проверки
    api_versions = [
        "2024-10-22",
        "2024-06-01", 
        "2023-06-01"
    ]
    
    # Список моделей для проверки
    test_models = [
        "claude-3-5-sonnet-20241022",
        "claude-sonnet-4-20250514",
        "claude-3-opus-20240229",
        "claude-3-sonnet-20240229"
    ]
    
    # Тестовый запрос (минимальный)
    test_messages = [{
        "role": "user",
        "content": "Hi"
    }]
    
    url = "https://api.anthropic.com/v1/messages"
    
    last_error = None
    
    for version in api_versions:
        for model in test_models:
            test_data = {
                "model": model,
                "max_tokens": 10,
                "messages": test_messages
            }
            
            headers = {
                "x-api-key": api_key,
                "anthropic-version": version,
                "content-type": "application/json"
            }
            
            try:
                response = requests.post(
                    url, 
                    headers=headers, 
                    json=test_data, 
                    timeout=30, 
                    verify=verify_ssl
                )
                
                if response.status_code == 200:
                    result = response.json()
                    return {
                        'valid': True,
                        'api_version': version,
                        'model': model,
                        'response': result,
                        'headers': headers,
                        'data': test_data
                    }
                elif response.status_code == 401:
                    return {
                        'valid': False,
                        'error': 'Неверный API ключ (401 Unauthorized)',
                        'api_version': version,
                        'model': model
                    }
                elif response.status_code == 403:
                    # Сохраняем информацию об ошибке, но продолжаем пробовать
                    last_error = {
                        'status_code': response.status_code,
                        'error_text': response.text[:200] if response.text else 'No error text',
                        'api_version': version,
                        'model': model
                    }
                    continue
                else:
                    # Сохраняем информацию об ошибке, но продолжаем пробовать
                    last_error = {
                        'status_code': response.status_code,
                        'error_text': response.text[:200] if response.text else 'No error text',
                        'api_version': version,
                        'model': model
                    }
            except requests.exceptions.SSLError:
                if verify_ssl:
                    # Пробуем без SSL
                    try:
                        response = requests.post(
                            url, 
                            headers=headers, 
                            json=test_data, 
                            timeout=30, 
                            verify=False
                        )
                        if response.status_code == 200:
                            result = response.json()
                            return {
                                'valid': True,
                                'api_version': version,
                                'model': model,
                                'response': result,
                                'headers': headers,
                                'data': test_data,
                                'ssl_warning': True
                            }
                    except Exception as e:
                        continue
                continue
            except Exception as e:
                # Пробуем следующую комбинацию
                continue
    
    # Если ничего не сработало, возвращаем последнюю ошибку или общую
    return {
        'valid': False,
        'error': 'Не удалось подключиться ни с одной комбинацией версии API и модели',
        'last_error': last_error
    }

def call_claude_api(messages: List[Dict], api_key: str = None, model: str = None, api_version: str = None, verify_ssl: bool = True) -> Optional[str]:
    """Вызывает API выбранной модели (Anthropic или OpenRouter).
    Если api_key передан явно, использует его (режим LLM-структурирования).
    Иначе берёт ключ из st.secrets в зависимости от выбранной модели."""
    selected_key = st.session_state.get('selected_model_key', 'claude_direct')
    model_cfg = _MODEL_BY_KEY.get(selected_key, _MODEL_BY_KEY['claude_direct'])
    provider = model_cfg['provider']

    # --- Anthropic (прямой Claude) ---
    if provider == 'anthropic':
        _key = api_key or st.secrets.get("CLAUDE_API_KEY", "") or st.session_state.get('claude_api_key', '')
        if not _key:
            st.error("Не задан ключ CLAUDE_API_KEY в секретах Streamlit.")
            return None
        _model = model or st.session_state.get('claude_working_model', "claude-sonnet-4-20250514")
        _api_version = api_version or st.session_state.get('claude_working_api_version', "2023-06-01")
        url = "https://api.anthropic.com/v1/messages"
        headers = {
            "x-api-key": _key,
            "anthropic-version": _api_version,
            "content-type": "application/json",
        }
        data = {"model": _model, "max_tokens": 4096, "messages": messages}

        def _extract_anthropic(result):
            usage = result.get('usage', {})
            st.session_state['_last_claude_usage'] = {
                'input_tokens':  usage.get('input_tokens', 0),
                'output_tokens': usage.get('output_tokens', 0),
                'model_key': selected_key,
            }
            return result.get('content', [{}])[0].get('text', '')

        try:
            response = requests.post(url, headers=headers, json=data, timeout=30, verify=verify_ssl)
            response.raise_for_status()
            return _extract_anthropic(response.json())
        except requests.exceptions.SSLError:
            if verify_ssl:
                try:
                    st.warning("⚠️ SSL ошибка. Повторяю без проверки...")
                    response = requests.post(url, headers=headers, json=data, timeout=30, verify=False)
                    response.raise_for_status()
                    return _extract_anthropic(response.json())
                except Exception as e2:
                    st.error(f"Ошибка Claude API (без SSL): {e2}")
                    return None
            else:
                st.error("Ошибка SSL при вызове Claude API.")
                return None
        except Exception as e:
            st.error(f"Ошибка Claude API: {e}")
            return None

    # --- OpenRouter ---
    else:
        _or_key = st.secrets.get("OPENROUTER_API_KEY", "") or st.session_state.get('openrouter_api_key', '')
        if not _or_key:
            st.error("Не задан ключ OPENROUTER_API_KEY в секретах Streamlit.")
            return None
        _model_id = model_cfg['model_id']
        url = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {_or_key}",
            "Content-Type": "application/json",
        }
        data = {"model": _model_id, "messages": messages}

        def _extract_openrouter(result):
            usage = result.get('usage', {})
            st.session_state['_last_claude_usage'] = {
                'input_tokens':  usage.get('prompt_tokens', 0),
                'output_tokens': usage.get('completion_tokens', 0),
                'model_key': selected_key,
            }
            choices = result.get('choices', [])
            if choices:
                return choices[0].get('message', {}).get('content', '')
            return ''

        try:
            response = requests.post(url, headers=headers, json=data, timeout=30)
            response.raise_for_status()
            return _extract_openrouter(response.json())
        except Exception as e:
            st.error(f"Ошибка OpenRouter API ({_model_id}): {e}")
            return None


def _get_model_prices(model_key: str) -> dict:
    cfg = _MODEL_BY_KEY.get(model_key)
    if cfg:
        return {'input': cfg['in_price'], 'output': cfg['out_price']}
    return {'input': 3.0, 'output': 15.0}


def _accumulate_cost():
    """Считывает _last_claude_usage и прибавляет к счётчикам затрат."""
    usage = st.session_state.get('_last_claude_usage')
    if not usage:
        return
    model_key = usage.get('model_key', st.session_state.get('selected_model_key', 'claude_direct'))
    prices = _get_model_prices(model_key)
    in_tok  = usage.get('input_tokens', 0)
    out_tok = usage.get('output_tokens', 0)
    cost_usd = (in_tok * prices['input'] + out_tok * prices['output']) / 1_000_000
    st.session_state['db_cost_input_tokens']  = st.session_state.get('db_cost_input_tokens', 0) + in_tok
    st.session_state['db_cost_output_tokens'] = st.session_state.get('db_cost_output_tokens', 0) + out_tok
    st.session_state['db_cost_usd']           = st.session_state.get('db_cost_usd', 0.0) + cost_usd
    st.session_state['_last_claude_usage']    = None


def _accumulate_cost_for_tag_run(run_id: Optional[int]):
    """Считывает _last_claude_usage и прибавляет к счётчикам затрат текущего прогона тегирования."""
    usage = st.session_state.get('_last_claude_usage')
    if not usage or not run_id:
        return

    model_key = usage.get('model_key', st.session_state.get('selected_model_key', 'claude_direct'))
    prices = _get_model_prices(model_key)
    in_tok = int(usage.get('input_tokens', 0) or 0)
    out_tok = int(usage.get('output_tokens', 0) or 0)
    cost_usd = (in_tok * prices['input'] + out_tok * prices['output']) / 1_000_000

    # Храним суммы в session_state по run_id (чтобы копилось в течение прогона).
    costs = st.session_state.get('tag_costs_by_run', {})
    key = str(int(run_id))
    prev = costs.get(key, {'input_tokens': 0, 'output_tokens': 0, 'usd': 0.0})
    prev['input_tokens'] = int(prev.get('input_tokens', 0)) + in_tok
    prev['output_tokens'] = int(prev.get('output_tokens', 0)) + out_tok
    prev['usd'] = float(prev.get('usd', 0.0)) + float(cost_usd)
    costs[key] = prev
    st.session_state['tag_costs_by_run'] = costs

    st.session_state['_last_claude_usage'] = None

def group_content_by_structure(data_dict: Dict) -> Dict:
    """Группирует записи содержания по предмету -> класс -> раздел -> тема."""
    grouped = {}
    
    for key, record in data_dict.items():
        if 'content' not in key.lower():
            continue
        
        subject = record.get('subject', '').strip() or 'без предмета'
        class_num = str(record.get('class', '')).strip() or '0'
        section = record.get('section', '').strip() or 'без раздела'
        topic = record.get('topic', '').strip() or 'без темы'
        text = record.get('text', '').strip()
        
        if not text:
            continue
        
        if subject not in grouped:
            grouped[subject] = {}
        if class_num not in grouped[subject]:
            grouped[subject][class_num] = {}
        if section not in grouped[subject][class_num]:
            grouped[subject][class_num][section] = {}
        if topic not in grouped[subject][class_num][section]:
            grouped[subject][class_num][section][topic] = []
        
        grouped[subject][class_num][section][topic].append(text)
    
    return grouped

def get_frp_sections_and_topics(data_dict: Dict) -> Dict:
    """Собирает разделы и темы из записей, где есть источник фрп_таблица."""
    frp_structure = {}
    
    for key, record in data_dict.items():
        sources = record.get('sources', [])
        if not sources or 'фрп_таблица' not in sources:
            continue
        
        subject = record.get('subject', '').strip()
        section = record.get('section', '').strip()
        topic = record.get('topic', '').strip()
        
        if not subject or not section:
            continue
        
        if subject not in frp_structure:
            frp_structure[subject] = {}
        if section not in frp_structure[subject]:
            frp_structure[subject][section] = set()
        if topic:
            frp_structure[subject][section].add(topic)
    
    # Преобразуем set в list для JSON сериализации
    result = {}
    for subject, sections in frp_structure.items():
        result[subject] = {}
        for section, topics in sections.items():
            result[subject][section] = sorted(list(topics))
    
    return result

def format_content_text(grouped: Dict, frp_structure: Dict) -> str:
    """Формирует текст для передачи модели."""
    lines = []
    
    # Добавляем информацию о разделах и темах ФРП
    if frp_structure:
        lines.append("По фрп имеются следующие разделы и темы:")
        for subject, sections in frp_structure.items():
            for section, topics in sections.items():
                lines.append(f"раздел: {section}")
                for topic in topics:
                    lines.append(f"  {topic}")
        lines.append("")
    
    # Группируем по предмету и классу
    for subject in sorted(grouped.keys()):
        for class_num in sorted(grouped[subject].keys(), key=lambda x: int(x) if str(x).isdigit() else 0):
            lines.append(f"предмет: {subject}")
            lines.append(f"класс: {class_num}")
            
            sections = grouped[subject][class_num]
            
            # Сначала записи с разделом и темой
            for section in sorted(sections.keys()):
                if section == 'без раздела':
                    continue
                topics = sections[section]
                for topic in sorted(topics.keys()):
                    if topic == 'без темы':
                        continue
                    lines.append(f"раздел: {section}")
                    lines.append(f"тема: {topic}")
                    
                    texts = topics[topic]
                    # Обрабатываем каждую запись отдельно: добавляем точку, если её нет
                    processed_texts = []
                    for text in texts:
                        text = text.strip()
                        if text:
                            # Если не заканчивается точкой/восклицательным/вопросительным, добавляем точку
                            if not text.rstrip().endswith(('.', '!', '?')):
                                text = text.rstrip() + '.'
                            processed_texts.append(text)
                    content_text = ' '.join(processed_texts)
                    lines.append(content_text)
                    lines.append("")
            
            # Записи с разделом, но без темы
            for section in sorted(sections.keys()):
                if section == 'без раздела':
                    continue
                if 'без темы' in sections[section]:
                    lines.append(f"раздел: {section}")
                    lines.append("тема: без темы")
                    texts = sections[section]['без темы']
                    # Обрабатываем каждую запись отдельно: добавляем точку, если её нет
                    processed_texts = []
                    for text in texts:
                        text = text.strip()
                        if text:
                            # Если не заканчивается точкой/восклицательным/вопросительным, добавляем точку
                            if not text.rstrip().endswith(('.', '!', '?')):
                                text = text.rstrip() + '.'
                            processed_texts.append(text)
                    content_text = ' '.join(processed_texts)
                    lines.append(content_text)
                    lines.append("")
            
            # Записи без раздела (в конец)
            if 'без раздела' in sections:
                lines.append("раздел: без раздела")
                all_no_section = []
                for topic_texts in sections['без раздела'].values():
                    all_no_section.extend(topic_texts)
                if all_no_section:
                    # Обрабатываем каждую запись отдельно: добавляем точку, если её нет
                    processed_texts = []
                    for text in all_no_section:
                        text = text.strip()
                        if text:
                            # Если не заканчивается точкой/восклицательным/вопросительным, добавляем точку
                            if not text.rstrip().endswith(('.', '!', '?')):
                                text = text.rstrip() + '.'
                            processed_texts.append(text)
                    content_text = ' '.join(processed_texts)
                    lines.append(content_text)
                    lines.append("")
    
    return '\n'.join(lines)

def parse_llm_response(response_text: str, subject: str, class_num: str) -> List[Dict]:
    """
    Парсит ответ от LLM и возвращает список записей.
    Обрабатывает: массивы [...], объекты {...}, markdown-блоки ```json```,
    неполные/обрезанные JSON, лишний текст до и после.
    """
    def _extract_items(data):
        """Извлекает список dict-записей из распарсенного объекта."""
        items = []
        if isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    items.append(item)
        elif isinstance(data, dict):
            # Словарь вида {"key": {...}, ...}
            for value in data.values():
                if isinstance(value, dict):
                    items.append(value)
                elif isinstance(value, list):
                    for v in value:
                        if isinstance(v, dict):
                            items.append(v)
        return items

    def _add_meta(items):
        for item in items:
            item['subject'] = subject
            item['class'] = class_num
        return items

    def _try_parse(text):
        """Пробует json.loads; при ошибке пытается дописать недостающие закрывающие символы."""
        text = text.strip()
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
        # Дописываем недостающие ']' или '}'
        for closing in (']', ']}', '}]', '}}', '"]', '"}', '"}]'):
            try:
                return json.loads(text + closing)
            except json.JSONDecodeError:
                pass
        # Пробуем срезать текст до последней полной записи
        # Ищем последнее }  перед которым можно закрыть массив
        last_brace = text.rfind('},')
        if last_brace == -1:
            last_brace = text.rfind('}')
        if last_brace > 0:
            candidate = text[:last_brace + 1]
            # Дополняем до массива или объекта
            for wrap in ('', ']', '}'):
                try:
                    return json.loads(candidate + wrap)
                except json.JSONDecodeError:
                    pass
        return None

    records = []
    text = response_text.strip()

    # 1. Убираем markdown-обёртку ```json ... ``` или просто ``` ... ```
    if '```' in text:
        # Берём содержимое между первым и последним ```
        parts = text.split('```')
        # parts[0] — до первого ```, parts[1] — внутри, parts[2] — после
        if len(parts) >= 3:
            inner = parts[1]
            # Убираем метку языка в первой строке (json, python и т.п.)
            lines = inner.split('\n')
            if lines and lines[0].strip().lower() in ('json', 'python', ''):
                inner = '\n'.join(lines[1:])
            text = inner.strip()
        # Если блоков несколько — берём самый длинный
        elif len(parts) >= 2:
            text = max(parts, key=len).strip()

    # 2. Сначала ищем массив [...]
    arr_match = re.search(r'\[[\s\S]*\]', text)
    if arr_match:
        data = _try_parse(arr_match.group())
        if data is not None:
            items = _extract_items(data)
            if items:
                return _add_meta(items)

    # 3. Ищем объект {...}
    obj_match = re.search(r'\{[\s\S]*\}', text)
    if obj_match:
        data = _try_parse(obj_match.group())
        if data is not None:
            items = _extract_items(data)
            if items:
                return _add_meta(items)

    # 4. Пробуем весь текст целиком (на случай если нет лишних символов)
    data = _try_parse(text)
    if data is not None:
        items = _extract_items(data)
        if items:
            return _add_meta(items)

    return records

# --- UI ---

st.title("📚 Извлечение и преобразование данных")
st.markdown("*ФРП, кодификатор, JSON → Excel*")

# Общий sidebar с выбором модели
with st.sidebar:
    st.header("⚙️ Настройки LLM")

    st.markdown("""
<style>
section[data-testid="stSidebar"] div[role="radiogroup"] {
    display: flex;
    flex-direction: column;
    gap: 2px;
}
section[data-testid="stSidebar"] div[role="radiogroup"] > label {
    padding: 5px 10px;
    border-radius: 6px;
    cursor: pointer;
    white-space: normal !important;
    line-height: 1.4;
    font-size: 0.83em;
    transition: background 0.1s ease;
}
section[data-testid="stSidebar"] div[role="radiogroup"] > label:hover {
    background: rgba(40, 167, 80, 0.12);
}
section[data-testid="stSidebar"] div[role="radiogroup"] > label:has(input:checked) {
    background: rgba(40, 167, 80, 0.25);
    font-weight: 600;
    color: #1a8a3a;
}
section[data-testid="stSidebar"] div[role="radiogroup"] > label > div[data-baseweb="radio"] {
    display: none;
}
</style>
""", unsafe_allow_html=True)

    _model_options = [m['key'] for m in AVAILABLE_MODELS]
    _model_labels  = {m['key']: m['display'] for m in AVAILABLE_MODELS}
    _cur_model = st.session_state.get('selected_model_key', 'claude_direct')
    if _cur_model not in _model_options:
        _cur_model = 'claude_direct'

    st.radio(
        "Модель",
        options=_model_options,
        index=_model_options.index(_cur_model),
        format_func=lambda k: _model_labels[k],
        key='selected_model_key',
        help="OpenRouter — модели на все случаи жизни. «Клод. Просто клод» — прямой доступ к Anthropic API.",
    )

    # --- Затраты на LLM (текущая сессия работы с БД) ---
    _in_tok  = st.session_state.get('db_cost_input_tokens', 0)
    _out_tok = st.session_state.get('db_cost_output_tokens', 0)
    _usd     = st.session_state.get('db_cost_usd', 0.0)
    if _in_tok > 0 or _out_tok > 0:
        st.markdown("---")
        st.markdown("**💰 Затраты (текущая сессия)**")
        st.caption(f"Входящие токены: {_in_tok:,}")
        st.caption(f"Исходящие токены: {_out_tok:,}")
        st.metric("Стоимость, USD", f"${_usd:.4f}")
        st.metric("Стоимость, ₽", f"{_usd * 90:.2f} ₽")
        st.caption("Сбрасывается при сохранении в базу")

    # --- Затраты на LLM (текущий прогон тегирования) ---
    _tag_run_id = st.session_state.get('tag_run_id')
    _tag_costs = st.session_state.get('tag_costs_by_run', {})
    _tag_key = str(int(_tag_run_id)) if _tag_run_id else None
    if _tag_key and isinstance(_tag_costs, dict) and _tag_key in _tag_costs:
        _tc = _tag_costs.get(_tag_key) or {}
        _tin = int(_tc.get('input_tokens', 0) or 0)
        _tout = int(_tc.get('output_tokens', 0) or 0)
        _tusd = float(_tc.get('usd', 0.0) or 0.0)
        if _tin > 0 or _tout > 0:
            st.markdown("---")
            st.markdown(f"**🏷️💰 Затраты (прогон run_id={_tag_key})**")
            st.caption(f"Входящие токены: {_tin:,}")
            st.caption(f"Исходящие токены: {_tout:,}")
            st.metric("Стоимость, USD", f"${_tusd:.4f}")
            st.metric("Стоимость, ₽", f"{_tusd * 90:.2f} ₽")

# Инициализация session_state
for k, v in [
    ('mode', 'frp_table'),
    ('extracted', False),
    ('intermediate_skills', {}),
    ('intermediate_content', {}),
    ('sections_df', None),
    ('original_pairs', []),
    ('final_skills_json', None),
    ('final_content_json', None),
    ('excel_skills_bytes', None),
    ('excel_content_bytes', None),
    ('available_jsons', []),
    ('last_extraction_mode', None),
    ('pdf_extracted', False),
    ('pdf_doc_type', None),
    ('pdf_merged', []),
    ('merge_jsons', []),  # список {name, data, type} для объединения
    ('compare_frp_table', None),  # ФРП таблица для сравнения
    ('compare_frp_text', None),   # ФРП текст для сравнения
    ('compare_codifier', None),   # Кодификатор для сравнения
    ('compare_etalon_data', None),  # Эталон для сопоставления (не изменяется)
    ('compare_report', []),       # Журнал отчёта
    ('compare_stats', {}),        # Статистика
    ('compare_pending_decisions', {}),  # [устаревшее] оставлено для совместимости
    ('compare_current_class', None),    # [устаревшее] оставлено для совместимости
    ('compare_fails', {}),             # Записи без предмета/класса/текста
    ('compare_for_choice', {}),        # Сличение двух и выбора (высокий порог сходства)
    ('compare_for_section_topic', {}), # Подбор раздела и темы
    ('compare_for_section_only', {}),  # Выбор раздела (совпадений нет)
    ('compare_simple_mode', False),     # Режим «простое сравнение»
    ('compare_merged_result', None),    # Результат объединения
    ('claude_api_key', ''),             # API ключ Claude (запасной, если нет в secrets)
    ('claude_verify_ssl', True),        # оставлено для совместимости
    ('claude_working_api_version', '2023-06-01'),
    ('claude_working_model', 'claude-sonnet-4-20250514'),
    ('selected_model_key', 'claude_direct'),  # выбранная модель для работы
    ('llm_content_data', None),         # Загруженные данные содержания для LLM
    ('llm_grouped_data', None),         # Сгруппированные данные
    ('llm_frp_structure', None),        # Структура ФРП разделов и тем
    ('llm_formatted_text', None),       # Отформатированный текст для модели
    ('llm_results', {}),                # Результаты обработки по парам предмет+класс
    ('llm_final_json', None),          # Финальный объединенный JSON
    # --- db_input mode ---
    ('db_frp_df', None),               # DataFrame frp_topics из БД
    ('db_fixed', False),               # Зафиксирован ли выбор ФРП
    ('db_fixed_topic_id', None),       # id выбранной темы frp_topics
    ('db_fixed_label', ''),            # Текст "Работаем с ..."
    ('db_mode_type', None),            # 'skills' или 'content'
    ('db_items', []),                  # список обрабатываемых элементов
    ('db_uid_counter', 0),             # счётчик уникальных id
    ('db_show_confirm', False),        # показывать диалог подтверждения
    ('db_save_result', None),          # результат сохранения
    ('db_add_frp_open', False),        # открыта форма добавления новой темы ФРП
    # --- стоимость вызовов Claude ---
    ('db_cost_input_tokens', 0),
    ('db_cost_output_tokens', 0),
    ('db_cost_usd', 0.0),
    ('_last_claude_usage', None),
    # --- батч "Доработать всё" ---
    ('db_batch_running', False),
    ('db_batch_pos', 0),
    ('db_batch_stop', False),
    # --- первичное тегирование ---
    ('tag_run_id', None),
    ('tag_subject_id', None),
    ('tag_program', 'базовый'),
    ('tag_only_nonempty', True),
    ('tag_topic_df', None),
    ('tag_topic_pos', 0),
    ('tag_stop', False),
    ('tag_last_result', None),
    ('tag_last_topic_id', None),
    ('tag_costs_by_run', {}),
    # --- просмотр базы ---
    ('vdb_type', None),
    ('vdb_reassign', False),
    ('vdb_df', None),
    ('vdb_grp_map', {}),
]:
    if k not in st.session_state:
        st.session_state[k] = v

# Выбор режима
st.header("Выберите режим")

mode = st.radio(
    "Режим работы:",
    options=[
        'frp_table',      # Из ФРП (таблица)
        'codifier',       # Из кодификатора
        'pdf',            # Из PDF
        'frp_text',       # Из ФРП (текст)
        'json_to_excel',  # JSON → Excel-таблицы
        'json_merge',     # Объединение JSON
        'json_compare',   # Слияние с сравнением JSON
        'llm_structure',  # Структурирование с помощью LLM
        'db_input',       # Добавление в базу данных
        'tagging_init',   # Первичное извлечение тегов
        'view_db',        # Просмотр базы данных
    ],
    format_func=lambda x: {
        'frp_table': 'Извлечение: ФРП (таблица Excel)',
        'codifier': 'Извлечение: из кодификатора',
        'pdf': 'Извлечение: PDF (ФРП или кодификатор)',
        'frp_text': 'Извлечение: ФРП (текст)',
        'json_to_excel': 'Преобразование: JSON → Excel-таблицы',
        'json_merge': 'Объединение нескольких JSON в один',
        'json_compare': 'Слияние и сравнение JSON файлов',
        'llm_structure': '🤖 Структурирование с помощью LLM',
        'db_input': '💾 Добавление в базу данных',
        'tagging_init': '🏷️ Тегирование: первичное извлечение',
        'view_db': '📋 Просмотр базы данных',
    }[x],
    horizontal=True,
    key='mode_selector'
)

st.markdown("---")

# ============ РЕЖИМ: ФРП (таблица) ============
if mode == 'frp_table':
    st.header("1️⃣ Извлечение из ФРП (таблица)")

    uploaded_file = st.file_uploader("Загрузите Excel файл ФРП", type=['xlsx', 'xls'], key='frp_upload')
    program_value = st.radio("Программа:", ['базовый', 'профильный'], horizontal=True, key='frp_program')

    if uploaded_file and not st.session_state.extracted:
        if st.button("Извлечь данные", type="primary"):
            with st.spinner("Извлечение данных..."):
                file_content = uploaded_file.read()
                xl_file = pd.ExcelFile(io.BytesIO(file_content))
                
                all_skills = []
                all_content = []
                for sheet_name in xl_file.sheet_names:
                    skills, content = extract_from_sheet(None, sheet_name, program_value, file_content)
                    all_skills.extend(skills)
                    all_content.extend(content)
                
                intermediate_skills = {f"skill_{i:04d}": s for i, s in enumerate(all_skills, 1)}
                intermediate_content = {f"content_{i:04d}": c for i, c in enumerate(all_content, 1)}
                
                sections_data = []
                seen = set()
                for item in list(intermediate_skills.values()) + list(intermediate_content.values()):
                    key = (item['section'] or '', item['topic'] or '')
                    if key not in seen:
                        sections_data.append({'Раздел': item['section'] or '', 'Тема': item['topic'] or ''})
                        seen.add(key)
                
                st.session_state.intermediate_skills = intermediate_skills
                st.session_state.intermediate_content = intermediate_content
                st.session_state.sections_df = pd.DataFrame(sections_data)
                st.session_state.original_pairs = [(r['Раздел'], r['Тема']) for _, r in st.session_state.sections_df.iterrows()]
                st.session_state.extracted = True
                st.session_state.last_extraction_mode = 'frp_table'
                st.rerun()

    if st.session_state.extracted:
        st.success(f"✅ Извлечено: {len(st.session_state.intermediate_skills)} навыков, {len(st.session_state.intermediate_content)} содержания")
        
        # Шаг 2: Редактирование
        st.header("2️⃣ Редактирование разделов и тем")
        st.caption("Отредактируйте таблицу и нажмите «Применить изменения»")
        
        edited_df = st.data_editor(
            st.session_state.sections_df,
            use_container_width=True,
            key="sections_editor",
            column_config={
                "Раздел": st.column_config.TextColumn("Раздел", width="large"),
                "Тема": st.column_config.TextColumn("Тема", width="large")
            },
            num_rows="fixed"
        )
        
        col1, col2, col3 = st.columns([1, 1, 3])
        with col1:
            if st.button("Применить изменения", type="primary"):
                mapping = {}
                for i in range(len(st.session_state.original_pairs)):
                    old_s, old_t = st.session_state.original_pairs[i]
                    if i < len(edited_df):
                        new_s = str(edited_df.iloc[i]['Раздел']).strip() if pd.notna(edited_df.iloc[i]['Раздел']) else ''
                        new_t = str(edited_df.iloc[i]['Тема']).strip() if pd.notna(edited_df.iloc[i]['Тема']) else ''
                        mapping[(old_s, old_t)] = (new_s, new_t)
                
                for skill in st.session_state.intermediate_skills.values():
                    key = (skill['section'], skill['topic'])
                    if key in mapping:
                        skill['section'], skill['topic'] = mapping[key]
                
                for content in st.session_state.intermediate_content.values():
                    key = (content['section'], content['topic'])
                    if key in mapping:
                        content['section'], content['topic'] = mapping[key]
                
                st.session_state.sections_df = edited_df.copy()
                st.session_state.original_pairs = [(edited_df.iloc[i]['Раздел'] or '', edited_df.iloc[i]['Тема'] or '') 
                    for i in range(len(edited_df))]
                st.success("✅ Изменения применены!")
                st.rerun()
        
        # Шаг 3: Сохранение в JSON
        st.header("3️⃣ Сохранение в JSON")
        
        if 'final_skills_json' not in st.session_state:
            st.session_state.final_skills_json = None
            st.session_state.final_content_json = None
        
        if st.button("Разбить на предложения"):
            with st.spinner("Обработка..."):
                final_skills = {}
                counter = 1
                for skill in st.session_state.intermediate_skills.values():
                    for sentence in split_text(skill['text']):
                        new_skill = skill.copy()
                        new_skill['text'] = sentence
                        final_skills[f"skill_{counter:04d}"] = new_skill
                        counter += 1
                
                final_content = {}
                counter = 1
                for content in st.session_state.intermediate_content.values():
                    for sentence in split_text(content['text']):
                        new_content = content.copy()
                        new_content['text'] = sentence
                        final_content[f"content_{counter:04d}"] = new_content
                        counter += 1
                
                st.session_state.final_skills_json = json.dumps(final_skills, ensure_ascii=False, indent=2)
                st.session_state.final_content_json = json.dumps(final_content, ensure_ascii=False, indent=2)
                st.session_state.final_counts = (len(final_skills), len(final_content))
                st.rerun()
        
        if st.session_state.final_skills_json:
            st.download_button("📥 Скачать frp_skills.json", st.session_state.final_skills_json.encode('utf-8'), 
                file_name="frp_skills.json", mime="application/json")
            st.download_button("📥 Скачать frp_content.json", st.session_state.final_content_json.encode('utf-8'), 
                file_name="frp_content.json", mime="application/json", key="dl_content")
            if 'final_counts' in st.session_state:
                st.info(f"Навыков: {st.session_state.final_counts[0]}, Содержания: {st.session_state.final_counts[1]}")
        
        if st.button("🔄 Начать заново", key='frp_reset'):
            st.session_state.extracted = False
            st.session_state.intermediate_skills = {}
            st.session_state.intermediate_content = {}
            st.session_state.sections_df = None
            st.session_state.original_pairs = []
            st.session_state.final_skills_json = None
            st.session_state.final_content_json = None
            st.session_state.excel_skills_bytes = None
            st.session_state.excel_content_bytes = None
            st.rerun()

# ============ РЕЖИМ: Кодификатор ============
elif mode == 'codifier':
    st.header("Извлечение из кодификатора")

    codifier_file = st.file_uploader("Загрузите Excel файл кодификатора", type=['xlsx', 'xls'], key='codifier_upload')
    subject_input = st.text_input("Предмет:", value="математика", key='codifier_subject')
    program_cod = st.radio("Программа:", ['базовый', 'профильный'], horizontal=True, key='codifier_program')

    if codifier_file:
        if st.button("🚀 Обработать кодификатор", type="primary", key='codifier_process'):
            with st.spinner("Обработка..."):
                file_content = codifier_file.read()
                xl_file = pd.ExcelFile(io.BytesIO(file_content))

                skills_sheet = next((n for n in xl_file.sheet_names if 'результат' in n.lower()), None)
                content_sheet = next((n for n in xl_file.sheet_names if 'содержан' in n.lower()), None)

                skills_dict = {}
                content_dict = {}

                if skills_sheet:
                    df_skills = pd.read_excel(io.BytesIO(file_content), sheet_name=skills_sheet, header=None)
                    skills_list = parse_codifier_sheet(df_skills, subject=subject_input.strip(), program=program_cod)
                    skills_dict = create_separate_elements(skills_list, 'skill')

                if content_sheet:
                    df_content = pd.read_excel(io.BytesIO(file_content), sheet_name=content_sheet, header=None)
                    content_list = parse_codifier_sheet(df_content, subject=subject_input.strip(), program=program_cod)
                    content_dict = create_separate_elements(content_list, 'content')

                st.session_state.final_skills_json = json.dumps(skills_dict, ensure_ascii=False, indent=2)
                st.session_state.final_content_json = json.dumps(content_dict, ensure_ascii=False, indent=2)
                st.session_state.final_counts = (len(skills_dict), len(content_dict))
                st.session_state.last_extraction_mode = 'codifier'
                st.rerun()

    if st.session_state.get('final_counts') and mode == 'codifier' and st.session_state.get('last_extraction_mode') == 'codifier':
        sk, ct = st.session_state.final_counts
        st.success(f"✅ Извлечено: {sk} навыков, {ct} содержания")
        col1, col2 = st.columns(2)
        with col1:
            if st.session_state.get('final_skills_json'):
                st.download_button("📥 Скачать навыки", st.session_state.final_skills_json.encode('utf-8'),
                    file_name="codifier_skills.json", mime="application/json", key="dl_cod_skills")
        with col2:
            if st.session_state.get('final_content_json'):
                st.download_button("📥 Скачать содержание", st.session_state.final_content_json.encode('utf-8'),
                    file_name="codifier_content.json", mime="application/json", key="dl_cod_content")
        st.caption("Данные также доступны в режиме «JSON → таблицы» для создания Excel.")

# ============ РЕЖИМ: PDF ============
elif mode == 'pdf':
    st.header("Извлечение из PDF")
    st.caption("Загрузка PDF → определение типа (ФРП/кодификатор) → извлечение таблиц с объединением по классу.")
    
    pdf_file = st.file_uploader("Загрузите PDF файл", type=['pdf'], key='pdf_upload')
    program_pdf = st.radio("Программа:", ['базовый', 'профильный'], horizontal=True, key='pdf_program')
    subject_pdf = st.text_input("Предмет (для кодификатора или если один предмет в ФРП):", value="математика", key='pdf_subject')
    
    if pdf_file:
        if st.button("Извлечь таблицы из PDF", type="primary", key='pdf_extract_btn'):
            try:
                import pdfplumber
            except ImportError:
                st.error("Установите pdfplumber: pip install pdfplumber")
            else:
                with st.spinner("Обработка PDF..."):
                    pdf_bytes = pdf_file.read()
                    doc_type, merged_dfs, stats = extract_and_merge_pdf_tables(pdf_bytes)
                    st.session_state.pdf_doc_type = doc_type
                    st.session_state.pdf_merged = merged_dfs
                    st.session_state.pdf_extracted = True
                    st.session_state.pdf_stats = stats
                    st.rerun()
    
    if st.session_state.get('pdf_extracted'):
        dt = st.session_state.pdf_doc_type
        stats = st.session_state.get('pdf_stats', {})
        merged = st.session_state.get('pdf_merged', [])
        
        st.success(f"Тип: **{dt.upper()}** | Таблиц извлечено: {stats.get('tables', 0)} | Классов/секций: {len(stats.get('classes', []))} — {', '.join(map(str, stats.get('classes', [])))}")
        
        col_save, col_process = st.columns(2)
        with col_save:
            if st.button("Сохранить в Excel", key='pdf_save_excel'):
                buf = io.BytesIO()
                with pd.ExcelWriter(buf, engine='openpyxl') as wr:
                    for name, df in merged:
                        sn = str(name)[:31]
                        df.to_excel(wr, sheet_name=sn, index=False)
                buf.seek(0)
                st.download_button("📥 Скачать Excel", buf.getvalue(), file_name="pdf_extracted.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key='dl_pdf_excel')
        
        with col_process:
            if st.button("Обработать и извлечь в JSON", type="primary", key='pdf_process_btn'):
                all_skills, all_content = [], []
                subj = subject_pdf.strip() or 'документ'
                prog = program_pdf
                for cls_or_name, df in merged:
                    if dt == 'frp':
                        sk, ct = extract_frp_from_df(df, subj, prog)
                        all_skills.extend(sk)
                        all_content.extend(ct)
                    else:
                        items = parse_codifier_sheet(df, subject=subj, program=prog)
                        flat = str(df.iloc[:3].values).lower() if len(df) > 0 else ""
                        if "содержан" in flat and "результат" not in flat:
                            all_content.extend(items)
                        else:
                            all_skills.extend(items)
                if dt == 'codifier':
                    sk_dict = create_separate_elements(all_skills, 'skill') if all_skills else {}
                    ct_dict = create_separate_elements(all_content, 'content') if all_content else {}
                
                if dt == 'frp':
                    inter_sk = {f"skill_{i:04d}": s for i, s in enumerate(all_skills, 1)}
                    inter_ct = {f"content_{i:04d}": c for i, c in enumerate(all_content, 1)}
                    sec_data = []
                    seen = set()
                    for it in list(inter_sk.values()) + list(inter_ct.values()):
                        k = (it.get('section', '') or '', it.get('topic', '') or '')
                        if k not in seen:
                            sec_data.append({'Раздел': k[0], 'Тема': k[1]})
                            seen.add(k)
                    st.session_state.intermediate_skills = inter_sk
                    st.session_state.intermediate_content = inter_ct
                    st.session_state.sections_df = pd.DataFrame(sec_data)
                    st.session_state.original_pairs = [(r['Раздел'], r['Тема']) for _, r in st.session_state.sections_df.iterrows()]
                    st.session_state.extracted = True
                    st.session_state.last_extraction_mode = 'frp_table'
                    st.session_state.final_skills_json = None
                    st.session_state.final_content_json = None
                    st.info("Данные готовы. Переключитесь на режим «ФРП (таблица Excel)» для редактирования разделов/тем и сохранения в JSON.")
                else:
                    st.session_state.final_skills_json = json.dumps(sk_dict, ensure_ascii=False, indent=2)
                    st.session_state.final_content_json = json.dumps(ct_dict, ensure_ascii=False, indent=2)
                    st.session_state.final_counts = (len(sk_dict), len(ct_dict))
                    st.session_state.last_extraction_mode = 'codifier'
                    st.success(f"Извлечено: {len(sk_dict)} навыков, {len(ct_dict)} содержания. Скачайте JSON в режиме «Кодификатор» или «JSON → таблицы».")
                st.rerun()

# ============ РЕЖИМ: ФРП (текст) ============
elif mode == 'frp_text':
    st.header("Извлечение из ФРП (текст)")
    st.caption("Загрузка DOC/DOCX/TXT → поиск «содержание обучения» и «предметные результаты» → извлечение в JSON.")
    
    frp_text_file = st.file_uploader("Загрузите DOC, DOCX или TXT", type=['doc', 'docx', 'txt'], key='frp_text_upload')
    subject_text = st.text_input("Предмет (обязательно):", value="", placeholder="напр. русский язык, математика", key='frp_text_subject')
    program_text = st.radio("Программа:", ['базовый', 'профильный'], horizontal=True, key='frp_text_program')
    
    if st.button("Обработать", type="primary", key='frp_text_process'):
        if not frp_text_file:
            st.warning("Сначала загрузите файл.")
        elif not subject_text.strip():
            st.warning("Заполните поле «Предмет» перед обработкой.")
        else:
            try:
                raw = frp_text_file.read()
                text = _read_doc_or_txt(raw, frp_text_file.name)
                content_items, skills_items = extract_frp_from_text(text, subject_text.strip(), program_text)
                content_dict = {f"content_{i:04d}": c for i, c in enumerate(content_items, 1)}
                skills_dict = {f"skill_{i:04d}": s for i, s in enumerate(skills_items, 1)}
                st.session_state.final_skills_json = json.dumps(skills_dict, ensure_ascii=False, indent=2)
                st.session_state.final_content_json = json.dumps(content_dict, ensure_ascii=False, indent=2)
                st.session_state.final_counts = (len(skills_dict), len(content_dict))
                st.session_state.frp_text_prefix = True
                st.success(f"Извлечено: {len(skills_dict)} навыков, {len(content_dict)} содержания.")
                st.rerun()
            except Exception as e:
                st.error(f"Ошибка: {e}")
                import traceback
                st.code(traceback.format_exc())
    
    if st.session_state.get('frp_text_prefix') and st.session_state.get('final_skills_json'):
        st.download_button("📥 Скачать frp_text_skills.json", st.session_state.final_skills_json.encode('utf-8'),
            file_name="frp_text_skills.json", mime="application/json", key='dl_frp_text_skills')
        st.download_button("📥 Скачать frp_text_content.json", st.session_state.final_content_json.encode('utf-8'),
            file_name="frp_text_content.json", mime="application/json", key='dl_frp_text_content')
        if st.session_state.get('final_counts'):
            st.info(f"Навыков: {st.session_state.final_counts[0]}, Содержания: {st.session_state.final_counts[1]}")

# ============ РЕЖИМ: JSON → Excel ============
elif mode == 'json_to_excel':
    st.header("Преобразование JSON → Excel-таблицы")

    # Собираем доступные JSON из session
    available = []
    if st.session_state.get('final_skills_json'):
        data = json.loads(st.session_state.final_skills_json)
        available.append({'type': 'навыки', 'data': data, 'name': 'Из текущей сессии (навыки)', 'info': get_json_info(data)})
    if st.session_state.get('final_content_json'):
        data = json.loads(st.session_state.final_content_json)
        available.append({'type': 'содержание', 'data': data, 'name': 'Из текущей сессии (содержание)', 'info': get_json_info(data)})

    # Добавляем загруженные ранее (из available_jsons)
    for entry in st.session_state.get('available_jsons', []):
        available.append(entry)

    # Показываем доступные
    if available:
        st.subheader("Доступные JSON")
        for i, entry in enumerate(available):
            info = entry.get('info') or get_json_info(entry.get('data', {}))
            with st.expander(f"📄 {entry.get('name', f'JSON {i+1}')} — {info.get('type', '?')}, {info.get('count', 0)} записей"):
                st.write("**Предметы:**", ", ".join(info.get('subjects', ['—'])))
                st.write("**Классы:**", ", ".join(str(c) for c in info.get('classes', ['—'])))
        st.caption("Для добавления ещё — загрузите файл ниже.")
    else:
        st.caption("Нет доступных JSON. Загрузите файл или выполните извлечение в другом режиме.")

    # Загрузка дополнительных JSON
    extra_upload = st.file_uploader("Загрузить ещё JSON", type=['json'], key='json_extra_upload')
    extra_type = st.radio("Тип файла:", ['навыки', 'содержание'], horizontal=True, key='extra_json_type')
    if extra_upload and st.button("Добавить к списку", key='add_json_btn'):
        try:
            data = json.loads(extra_upload.read().decode('utf-8'))
            name = extra_upload.name or "Загруженный файл"
            existing_names = [a.get('name') for a in st.session_state.get('available_jsons', [])]
            if name in existing_names:
                st.warning("Файл с таким именем уже добавлен.")
            else:
                entry = {'type': extra_type, 'data': data, 'name': name, 'info': get_json_info(data)}
                st.session_state.available_jsons = st.session_state.get('available_jsons', []) + [entry]
                st.success("Файл добавлен в список.")
                st.rerun()
        except Exception as e:
            st.error(f"Ошибка: {e}")

    # Выбор и конвертация
    st.subheader("Создание Excel")
    skills_options = [a for a in available if a.get('type') == 'навыки']
    content_options = [a for a in available if a.get('type') == 'содержание']

    excel_col1, excel_col2 = st.columns(2)
    with excel_col1:
        st.markdown("**Навыки → Excel**")
        if skills_options:
            sel_skills = st.selectbox("Выберите JSON", range(len(skills_options)), format_func=lambda i: skills_options[i].get('name', f'Вариант {i+1}'), key='sel_skills')
            if st.button("Создать Excel — навыки", key='excel_skills_btn'):
                st.session_state.excel_skills_bytes = json_to_excel_sorted(skills_options[sel_skills]['data'], 'Навык')
                st.rerun()
            if st.session_state.excel_skills_bytes:
                st.download_button("📥 Скачать frp_skills.xlsx", st.session_state.excel_skills_bytes,
                    file_name="frp_skills.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key='dl_skills_xlsx')
        else:
            st.caption("Нет JSON с навыками.")

    with excel_col2:
        st.markdown("**Содержание → Excel**")
        if content_options:
            sel_content = st.selectbox("Выберите JSON", range(len(content_options)), format_func=lambda i: content_options[i].get('name', f'Вариант {i+1}'), key='sel_content')
            if st.button("Создать Excel — содержание", key='excel_content_btn'):
                st.session_state.excel_content_bytes = json_to_excel_sorted(content_options[sel_content]['data'], 'Содержание')
                st.rerun()
            if st.session_state.excel_content_bytes:
                st.download_button("📥 Скачать frp_content.xlsx", st.session_state.excel_content_bytes,
                    file_name="frp_content.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key='dl_content_xlsx')
        else:
            st.caption("Нет JSON с содержанием.")

# ============ РЕЖИМ: Объединение JSON ============
elif mode == 'json_merge':
    st.header("Объединение нескольких JSON в один")
    st.caption("Загружайте JSON одного типа (только навыки или только содержание). Сквозная нумерация при объединении.")
    
    if 'merge_jsons' not in st.session_state:
        st.session_state.merge_jsons = []
    
    merge_list = st.session_state.merge_jsons
    
    # Список загруженных
    if merge_list:
        types = set(e.get('type') for e in merge_list)
        if len(types) > 1:
            st.error("⚠️ Смешивать навыки и содержание нельзя! Загружайте файлы только одного типа. Очистите список и начните заново.")
        else:
            st.subheader("Загружено для объединения")
            for i, e in enumerate(merge_list):
                info = e.get('info') or get_json_info(e.get('data', {}))
                st.text(f"  {i+1}. {e.get('name', '?')} — {info.get('count', 0)} записей")
    
    # Загрузка (или «Загрузить ещё» если уже есть файлы)
    merge_upload = st.file_uploader("Загрузить ещё JSON" if merge_list else "Загрузить JSON", type=['json'], key='merge_upload')
    
    if merge_upload:
        if st.button("Добавить к списку", key='merge_add_btn'):
            try:
                data = json.loads(merge_upload.read().decode('utf-8'))
                detected = get_json_info(data)
                det_type = 'навыки' if detected.get('type') == 'навыки' else 'содержание'
                if merge_list and det_type != (merge_list[0].get('type') or ''):
                    st.warning("Этот файл содержит «содержание», а в списке — навыки (или наоборот). Смешивать нельзя!")
                else:
                    entry = {'type': det_type, 'data': data, 'name': merge_upload.name or "Файл", 'info': detected}
                    st.session_state.merge_jsons = merge_list + [entry]
                    st.success("Файл добавлен.")
                    st.rerun()
            except Exception as e:
                st.error(f"Ошибка: {e}")
    
    if merge_list:
        if st.button("Объединить", type="primary", key='merge_do_btn'):
            types = set(e.get('type') for e in merge_list)
            if len(types) > 1:
                st.error("В списке смешаны навыки и содержание. Удалите лишние файлы.")
            else:
                merged = {}
                counter = 1
                prefix = 'skill' if 'навыки' in (merge_list[0].get('type') or '') else 'content'
                for entry in merge_list:
                    for _key, item in sorted(entry['data'].items(), key=lambda x: x[0]):
                        merged[f"{prefix}_{counter:04d}"] = item
                        counter += 1
                st.session_state.merged_json_result = json.dumps(merged, ensure_ascii=False, indent=2)
                st.session_state.merged_json_type = prefix
                st.session_state.merged_json_count = len(merged)
                st.rerun()
        
        if st.button("Очистить список", key='merge_clear_btn'):
            st.session_state.merge_jsons = []
            st.session_state.merged_json_result = None
            st.rerun()
    
    if st.session_state.get('merged_json_result'):
        jr = st.session_state.merged_json_result
        cnt = st.session_state.get('merged_json_count', 0)
        st.success(f"Объединено: {cnt} записей")
        fname = "merged_skills.json" if st.session_state.get('merged_json_type') == 'skill' else "merged_content.json"
        st.download_button("📥 Скачать объединённый JSON", jr.encode('utf-8'), file_name=fname, mime="application/json", key='dl_merged')

# ============ РЕЖИМ: Слияние и сравнение JSON ============
elif mode == 'json_compare':
    st.header("🔗 Слияние и сравнение JSON файлов")
    st.caption("Загрузка ФРП (таблица), ФРП (текст) и кодификатора → автоматическое сравнение и объединение с интерактивным выбором")
    
    # Загрузка файлов
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("ФРП таблица")
        frp_table_file = st.file_uploader("Загрузить ФРП таблица", type=['json'], key='compare_frp_table_upload')
        if frp_table_file:
            try:
                data = json.loads(frp_table_file.read().decode('utf-8'))
                is_valid, message = validate_json_source(data, 'фрп_планирование')
                if is_valid:
                    # Добавляем префикс к ключам для различения источников
                    prefixed_data = add_prefix_to_keys(data, 'frp_table_')
                    st.session_state.compare_frp_table = {'name': frp_table_file.name, 'data': prefixed_data}
                    st.success(f"✅ {len(prefixed_data)} записей")
                else:
                    st.warning(f"⚠️ {message}")
            except Exception as e:
                st.error(f"Ошибка: {e}")
    
    with col2:
        st.subheader("ФРП текст")
        frp_text_file = st.file_uploader("Загрузить ФРП текст", type=['json'], key='compare_frp_text_upload')
        if frp_text_file:
            try:
                data = json.loads(frp_text_file.read().decode('utf-8'))
                is_valid, message = validate_json_source(data, 'фрп_текст')
                if is_valid:
                    # Добавляем префикс к ключам для различения источников
                    prefixed_data = add_prefix_to_keys(data, 'frp_text_')
                    st.session_state.compare_frp_text = {'name': frp_text_file.name, 'data': prefixed_data}
                    st.success(f"✅ {len(prefixed_data)} записей")
                else:
                    st.warning(f"⚠️ {message}")
            except Exception as e:
                st.error(f"Ошибка: {e}")
    
    with col3:
        st.subheader("Кодификатор")
        codifier_file = st.file_uploader("Загрузить кодификатор", type=['json'], key='compare_codifier_upload')
        if codifier_file:
            try:
                data = json.loads(codifier_file.read().decode('utf-8'))
                is_valid, message = validate_json_source(data, 'кодификатор')
                if is_valid:
                    # Добавляем префикс к ключам для различения источников
                    prefixed_data = add_prefix_to_keys(data, 'codifier_')
                    st.session_state.compare_codifier = {'name': codifier_file.name, 'data': prefixed_data}
                    st.success(f"✅ {len(prefixed_data)} записей")
                else:
                    st.warning(f"⚠️ {message}")
            except Exception as e:
                st.error(f"Ошибка: {e}")
    
    # Проверка наличия файлов
    frp_table_loaded = st.session_state.compare_frp_table is not None
    frp_text_loaded = st.session_state.compare_frp_text is not None
    codifier_loaded = st.session_state.compare_codifier is not None
    
    if not frp_table_loaded:
        st.info("ℹ️ Загрузите ФРП таблицу для начала сравнения. Без этого файла сравнение невозможно.")
    elif not frp_text_loaded and not codifier_loaded:
        st.info("ℹ️ Загрузите хотя бы один дополнительный файл (ФРП текст или кодификатор) для сравнения.")
    else:
        # Кнопки начала сравнения
        btn_col1, btn_col2 = st.columns(2)
        with btn_col1:
            start_full = st.button("🚀 Начать сравнение и слияние", type="primary", key='compare_start_btn')
        with btn_col2:
            start_simple = st.button("⚡ Простое сравнение", key='compare_simple_btn', help="Без подбора раздела/темы по похожим — всё сразу в пилот или выбор раздела")
        
        if start_full or start_simple:
            st.session_state.compare_simple_mode = bool(start_simple)
            # Инициализация отчёта и статистики
            st.session_state.compare_report = []
            st.session_state.compare_stats = {
                'auto_merged': 0,
                'user_selected': 0,
                'both_saved': 0,
                'section_assigned_auto': 0,
                'section_assigned_user': 0
            }
            st.session_state.compare_pending_decisions = {}
            st.session_state.compare_current_class = None
            st.session_state.compare_fails = {}
            st.session_state.compare_for_choice = {}
            st.session_state.compare_for_section_topic = {}
            st.session_state.compare_for_section_only = {}
            st.session_state.compare_merged_result = None
            
            # Сохраняем исходное количество записей в исследуемых файлах
            st.session_state.compare_initial_counts = {}
            if frp_text_loaded:
                st.session_state.compare_initial_counts['frp_text'] = len(st.session_state.compare_frp_text['data'])
            if codifier_loaded:
                st.session_state.compare_initial_counts['codifier'] = len(st.session_state.compare_codifier['data'])
            
            # Эталон — исходный словарь, не изменяется. Рабочая копия — для добавления/объединения.
            etalon = copy.deepcopy(st.session_state.compare_frp_table['data'])
            st.session_state.compare_etalon_data = etalon
            st.session_state.compare_base_data = copy.deepcopy(etalon)
            # Определяем порядок сравнения
            if frp_table_loaded and frp_text_loaded and codifier_loaded:
                st.session_state.compare_iteration = 1
                st.session_state.compare_compare_data = st.session_state.compare_frp_text['data']
                st.session_state.compare_next_data = st.session_state.compare_codifier['data']
            elif frp_table_loaded and frp_text_loaded:
                st.session_state.compare_iteration = 1
                st.session_state.compare_compare_data = st.session_state.compare_frp_text['data']
                st.session_state.compare_next_data = None
            elif frp_table_loaded and codifier_loaded:
                st.session_state.compare_iteration = 1
                st.session_state.compare_compare_data = st.session_state.compare_codifier['data']
                st.session_state.compare_next_data = None
            
            st.rerun()
        
        if 'compare_iteration' in st.session_state:
            
            # CSS: текст навыков чёрный (не серый)
            st.markdown("""
                <style>
                div[data-testid="stVerticalBlock"] textarea[disabled] { color: #000 !important; -webkit-text-fill-color: #000 !important; }
                .stCaption { color: #000 !important; }
                </style>
                """, unsafe_allow_html=True)
            
            # Фейлы — в самое начало (все возможные)
            if st.session_state.compare_fails:
                with st.expander(f"⚠️ Фейлы — записи без предмета/класса/текста ({len(st.session_state.compare_fails)})", expanded=True):
                    for fk, frec in sorted(st.session_state.compare_fails.items()):
                        subj = (frec.get('subject') or '').strip() or '(пусто)'
                        cls = str(frec.get('class', '')) or '(пусто)'
                        txt = (frec.get('text') or '') or '(пусто)'
                        st.markdown(f"**{fk}** | предмет: {subj} | класс: {cls}")
                        st.write(f"Текст: {txt}")
            
            # Режим
            if st.session_state.get('compare_simple_mode'):
                st.info("⚡ **Простое сравнение**: совпадения и выбор из двух — как есть; остальное сразу в пилот при наличии раздела или выбор раздела.")
            # Проверка: compare_compare_data должен быть пуст после обработки
            remaining = st.session_state.get('compare_compare_data', {})
            if remaining:
                st.warning(f"⚠️ В compare_compare_data осталось {len(remaining)} записей — не должно остаться. Ключи: {list(remaining.keys())[:10]}{'…' if len(remaining) > 10 else ''}")
            else:
                st.success("✅ compare_compare_data пуст — все записи обработаны")
            
            # Записи без сходства: что с ними сделано
            if st.session_state.compare_report:
                no_match_ok_reports = [r for r in st.session_state.compare_report if r.get('action') == 'no_match_section_ok']
                pending_section = len(st.session_state.compare_for_section_only)
                if no_match_ok_reports or pending_section:
                    st.write("**Записи без сходства с эталоном:**")
                if no_match_ok_reports:
                    with st.expander(f"✅ Без совпадений, раздел совпал — автосохранено: {len(no_match_ok_reports)}", expanded=False):
                        for r in no_match_ok_reports:
                            txt = r.get('text', '')
                            st.write(f"**Текст:** {txt[:100] + ('…' if len(txt) > 100 else '')}")
                            st.write(f"Раздел: {r.get('section', '')} | {r.get('note', '')}")
                if pending_section > 0:
                    st.write(f"Ожидают выбора раздела: {pending_section} (см. блок «Записи без совпадений» ниже)")
                auto_merged_reports = [
                    r for r in st.session_state.compare_report 
                    if r.get('action') in ['auto_merge', 'auto_merge_new']
                ]
                if auto_merged_reports:
                    with st.expander(f"✅ Автоматически объединено записей: {len(auto_merged_reports)}", expanded=False):
                        for report_item in auto_merged_reports:
                            if report_item.get('action') == 'auto_merge':
                                st.write(f"**Объединено:**")
                                st.write(f"- **Эталонная:** {report_item.get('base_text', '')}")
                                st.write(f"- **Исследуемая:** {report_item.get('compare_text', '')}")
                                st.write(f"- Похожесть: {report_item.get('similarity', 0):.1%}")
                                if report_item.get('section'):
                                    st.write(f"- Раздел: {report_item.get('section')} | Тема: {report_item.get('topic')}")
            
            # Работа пользователя: фейлы, затем три словаря для выбора
            has_work = (st.session_state.compare_fails or st.session_state.compare_for_choice or
                        st.session_state.compare_for_section_topic or st.session_state.compare_for_section_only)
            if has_work:
                # Кнопка «Принять всё как есть» — добавляет все необработанные записи в базу как есть
                any_pending = (st.session_state.compare_for_choice or st.session_state.compare_for_section_topic or
                               st.session_state.compare_for_section_only)
                # Две кнопки: «Принять все изменения» (по текущему выбору) и «Принять всё как есть» (игнорируя выбор)
                if any_pending:
                    btn_col_a, btn_col_b = st.columns(2)
                    with btn_col_a:
                        accept_changes = st.button("✅ Принять все изменения", key="accept_all_changes",
                            help="Применить текущий выбор (радио, селекты) ко всем отображаемым записям")
                    with btn_col_b:
                        accept_as_is = st.button("✅ Принять всё как есть", key="accept_all_as_is",
                            help="Добавить все записи в базу без учёта выбора (как есть)")
                
                if accept_changes:
                    # Применяем текущий выбор из UI ко всем записям
                    base_data = st.session_state.compare_base_data
                    keys_to_del = []
                    for compare_key, decision in list(st.session_state.compare_for_choice.items()):
                        choice_key = f"choice_{compare_key}"
                        if choice_key in st.session_state:
                            selected = st.session_state[choice_key]
                            base_key = decision.get('base_key')
                            base_record = decision.get('base_record', {})
                            compare_record = decision.get('compare_record', {})
                            if selected == 'base' and base_key and base_key in base_data:
                                merged_rec = base_data[base_key].copy()
                                merged_rec['sources'] = list(set(merged_rec.get('sources', []) + compare_record.get('sources', [])))
                                base_data[base_key] = merged_rec
                            elif selected == 'compare' and base_key and base_key in base_data:
                                rec = compare_record.copy()
                                rec['section'] = base_record.get('section', '')
                                rec['topic'] = base_record.get('topic', '')
                                rec['sources'] = list(set(base_data[base_key].get('sources', []) + rec.get('sources', [])))
                                base_data[base_key] = rec
                            elif selected == 'both' and base_key:
                                rec = compare_record.copy()
                                rec['section'] = base_record.get('section', '')
                                rec['topic'] = base_record.get('topic', '')
                                # Объединяем источники
                                base_sources = base_data[base_key].get('sources', [])
                                compare_sources = compare_record.get('sources', [])
                                rec['sources'] = list(set(base_sources + compare_sources))
                                prefix = 'frp_text_' if ('фрп_текст' in (compare_record.get('sources') or [])) else 'codifier_'
                                max_num = 0
                                for k in base_data:
                                    m = re.search(r'(?:skill_|content_)(\d+)', k)
                                    if m: max_num = max(max_num, int(m.group(1)))
                                base_data[f"{prefix}skill_{max_num+1:04d}"] = rec
                            keys_to_del.append(('choice', compare_key))
                    # section_topic — восстанавливаем опции из similar_records, читаем индекс из radio
                    for compare_key, decision in list(st.session_state.compare_for_section_topic.items()):
                        rec = decision.get('record', {}).copy()
                        section_key = f"section_{compare_key}"
                        topic_key = f"topic_{compare_key}"
                        similar = decision.get('similar_records', [])
                        sections_count = {}
                        topics_count = {}
                        for sim in similar:
                            s = (sim.get('record', {}).get('section') or '').strip()
                            t = (sim.get('record', {}).get('topic') or '').strip()
                            if s: sections_count[s] = sections_count.get(s, 0) + 1
                            if t: topics_count[t] = topics_count.get(t, 0) + 1
                        unique_sections = [x[0] for x in sorted(sections_count.items(), key=lambda z: -z[1])]
                        unique_topics = [x[0] for x in sorted(topics_count.items(), key=lambda z: -z[1])]
                        section_options = [''] + unique_sections
                        topic_options = [''] + unique_topics
                        current_s = rec.get('section', '').strip()
                        current_t = rec.get('topic', '').strip()
                        if current_s and current_s not in unique_sections:
                            unique_sections = [current_s] + unique_sections
                            section_options = [''] + unique_sections
                        if current_t and current_t not in unique_topics:
                            unique_topics = [current_t] + unique_topics
                            topic_options = [''] + unique_topics
                        if section_key in st.session_state:
                            idx = st.session_state[section_key]
                            if isinstance(idx, int) and 0 <= idx < len(section_options):
                                rec['section'] = section_options[idx]
                        if topic_key in st.session_state:
                            idx = st.session_state[topic_key]
                            if isinstance(idx, int) and 0 <= idx < len(topic_options):
                                rec['topic'] = topic_options[idx]
                        prefix = 'frp_text_' if ('фрп_текст' in (rec.get('sources') or [])) else 'codifier_'
                        max_num = max((int(m.group(1)) for k in base_data for m in [re.search(r'(?:skill_|content_)(\d+)', k)] if m), default=0)
                        base_data[f"{prefix}skill_{max_num+1:04d}"] = rec
                        keys_to_del.append(('section_topic', compare_key))
                    # section_only — selectbox возвращает выбранную строку напрямую
                    for compare_key, decision in list(st.session_state.compare_for_section_only.items()):
                        rec = decision.get('record', {}).copy()
                        sel_key = f"new_section_{compare_key}"
                        if sel_key in st.session_state:
                            sel_val = st.session_state[sel_key]
                            if isinstance(sel_val, str):
                                rec['section'] = sel_val
                        prefix = 'frp_text_' if ('фрп_текст' in (rec.get('sources') or [])) else 'codifier_'
                        max_num = max((int(m.group(1)) for k in base_data for m in [re.search(r'(?:skill_|content_)(\d+)', k)] if m), default=0)
                        base_data[f"{prefix}skill_{max_num+1:04d}"] = rec
                        keys_to_del.append(('section_only', compare_key))
                    for kind, k in keys_to_del:
                        if kind == 'choice' and k in st.session_state.compare_for_choice:
                            del st.session_state.compare_for_choice[k]
                        elif kind == 'section_topic' and k in st.session_state.compare_for_section_topic:
                            del st.session_state.compare_for_section_topic[k]
                        elif kind == 'section_only' and k in st.session_state.compare_for_section_only:
                            del st.session_state.compare_for_section_only[k]
                    if keys_to_del:
                        st.session_state.compare_stats['section_assigned_user'] = st.session_state.compare_stats.get('section_assigned_user', 0) + len(keys_to_del)
                        if _check_and_transition_next_iteration():
                            st.success("✅ Все сравнения завершены!")
                        st.rerun()
                
                elif accept_as_is:
                    base_data = st.session_state.compare_base_data
                    prefix = 'merged_'
                    import re
                    max_num = max((int(m.group(1)) for k in base_data for m in [re.search(r'(?:skill_|content_)(\d+)', k)] if m), default=0)
                    to_add = []
                    for compare_key, decision in list(st.session_state.compare_for_choice.items()):
                        rec = decision.get('compare_record', decision.get('record'))
                        if rec:
                            to_add.append(('choice', compare_key, rec))
                    for compare_key, decision in list(st.session_state.compare_for_section_topic.items()):
                        rec = decision.get('record', {})
                        if rec:
                            to_add.append(('section_topic', compare_key, rec))
                    for compare_key, decision in list(st.session_state.compare_for_section_only.items()):
                        rec = decision.get('record', {})
                        if rec:
                            to_add.append(('section_only', compare_key, rec))
                    for kind, compare_key, rec in to_add:
                        max_num += 1
                        base_data[f"{prefix}skill_{max_num:04d}"] = rec.copy()
                        if kind == 'choice' and compare_key in st.session_state.compare_for_choice:
                            del st.session_state.compare_for_choice[compare_key]
                        elif kind == 'section_topic' and compare_key in st.session_state.compare_for_section_topic:
                            del st.session_state.compare_for_section_topic[compare_key]
                        elif kind == 'section_only' and compare_key in st.session_state.compare_for_section_only:
                            del st.session_state.compare_for_section_only[compare_key]
                    if to_add:
                        st.session_state.compare_stats['section_assigned_user'] = st.session_state.compare_stats.get('section_assigned_user', 0) + len(to_add)
                        if _check_and_transition_next_iteration():
                            st.success("✅ Все сравнения завершены!")
                        st.rerun()
                
                # Сличение двух и выбора (высокий порог сходства)
                if st.session_state.compare_for_choice:
                    st.subheader("Сличение двух и выбора (высокий порог сходства)")
                for compare_key, decision in list(st.session_state.compare_for_choice.items()):
                    if decision.get('type') == 'choice':
                            base_text = decision['base_record'].get('text', '')
                            compare_text = decision['compare_record'].get('text', '')
                            similarity = decision.get('similarity', 0)
                            base_section = decision['base_record'].get('section', 'не указан')
                            base_topic = decision['base_record'].get('topic', 'не указана')
                            
                            # Извлекаем номер записи из ключа
                            record_num = ''
                            if compare_key:
                                # Пытаемся извлечь номер из ключа вида "frp_text_skill_0001" или "codifier_skill_0001"
                                parts = compare_key.split('_')
                                if len(parts) >= 2 and parts[-1].isdigit():
                                    record_num = f" (№{parts[-1]})"
                                elif '_skill_' in compare_key or '_content_' in compare_key:
                                    # Извлекаем номер после skill_ или content_
                                    import re
                                    match = re.search(r'(?:skill_|content_)(\d+)', compare_key)
                                    if match:
                                        record_num = f" (№{match.group(1)})"
                            
                            st.write(f"**Исследуемая запись{record_num} совпала с эталонной (похожесть: {similarity:.1%})**")
                            st.write(f"**Эталонная запись:** Раздел: {base_section} | Тема: {base_topic}")
                            
                            base_prefix = "✅ Вариант 1 (эталонная): "
                            compare_prefix = "✅ Вариант 2 (исследуем): "
                            base_label = f"{base_prefix}{base_text}"
                            compare_label = f"{compare_prefix}{compare_text}"
                            
                            # Радиобаттоны для выбора с выровненными текстами
                            choice_key = f"choice_{compare_key}"
                            selected = st.radio(
                                "Выберите вариант:",
                                options=['base', 'compare', 'both'],
                                format_func=lambda x: {
                                    'base': base_label,
                                    'compare': compare_label,
                                    'both': "✅ Сохранить оба варианта"
                                }[x],
                                key=choice_key
                            )
                            
                            col_save1, col_save2 = st.columns(2)
                            with col_save1:
                                if st.button("💾 Сохранить выбранное", key=f"save_{choice_key}"):
                                    # Сохраняем выбранный вариант
                                    base_data = st.session_state.compare_base_data
                                    base_key = decision.get('base_key')
                                    
                                    if selected == 'base' or selected == 'both':
                                        # Обновляем базовую запись
                                        if base_key:
                                            merged_record = base_data[base_key].copy()
                                            compare_sources = decision['compare_record'].get('sources', [])
                                            base_sources = merged_record.get('sources', [])
                                            merged_record['sources'] = list(set(base_sources + compare_sources))
                                            base_data[base_key] = merged_record
                                    
                                    if selected == 'compare' or selected == 'both':
                                        # Добавляем сравниваемую запись
                                        # Раздел/тему всегда берём из эталонной записи
                                        base_section = decision['base_record'].get('section', '')
                                        base_topic = decision['base_record'].get('topic', '')
                                        
                                        if selected == 'compare':
                                            # Заменяем базовую
                                            if base_key:
                                                compare_record = decision['compare_record'].copy()
                                                compare_record['section'] = base_section
                                                compare_record['topic'] = base_topic
                                                base_sources = base_data[base_key].get('sources', [])
                                                compare_sources = compare_record.get('sources', [])
                                                compare_record['sources'] = list(set(base_sources + compare_sources))
                                                base_data[base_key] = compare_record
                                        else:
                                            # Добавляем как новую
                                            # Определяем префикс на основе источника
                                            compare_sources = decision['compare_record'].get('sources', [])
                                            if compare_sources and 'фрп_текст' in compare_sources:
                                                prefix = 'frp_text_'
                                            elif compare_sources and 'кодификатор' in compare_sources:
                                                prefix = 'codifier_'
                                            else:
                                                prefix = 'merged_'
                                            
                                            # Находим максимальный номер среди существующих ключей с таким префиксом
                                            max_num = 0
                                            for k in base_data.keys():
                                                if k.startswith(prefix) and ('skill_' in k or 'content_' in k):
                                                    try:
                                                        parts = k.split('_')
                                                        if len(parts) >= 3 and parts[-1].isdigit():
                                                            max_num = max(max_num, int(parts[-1]))
                                                    except:
                                                        pass
                                            
                                            new_key = f"{prefix}skill_{max_num + 1:04d}"
                                            compare_record = decision['compare_record'].copy()
                                            compare_record['section'] = base_section
                                            compare_record['topic'] = base_topic
                                            # Объединяем источники с базовой записью
                                            if base_key:
                                                base_sources = base_data[base_key].get('sources', [])
                                                compare_sources_list = compare_record.get('sources', [])
                                                compare_record['sources'] = list(set(base_sources + compare_sources_list))
                                            base_data[new_key] = compare_record
                                    
                                    # Обновляем статистику и журнал
                                    if selected == 'both':
                                        st.session_state.compare_stats['both_saved'] += 1
                                        st.session_state.compare_report.append({
                                            'action': 'both_saved',
                                            'base_text': base_text,
                                            'compare_text': compare_text
                                        })
                                    else:
                                        st.session_state.compare_stats['user_selected'] += 1
                                        st.session_state.compare_report.append({
                                            'action': 'user_selected',
                                            'selected': selected,
                                            'base_text': base_text,
                                            'compare_text': compare_text
                                        })
                                    
                                    # Удаляем из словаря — запись обработана
                                    if compare_key in st.session_state.compare_for_choice:
                                        del st.session_state.compare_for_choice[compare_key]
                                    if _check_and_transition_next_iteration():
                                        st.success("✅ Все сравнения завершены!" if st.session_state.compare_merged_result else "Переход к следующему источнику...")
                                    st.rerun()
                            
                            with col_save2:
                                pass
                
                # 3. Подбор раздела и темы (совпадения есть, но не очень близкие)
                if st.session_state.compare_for_section_topic:
                    st.subheader("Подбор раздела и темы")
                for compare_key, decision in list(st.session_state.compare_for_section_topic.items()):
                            record_text = decision['record'].get('text', '')
                            similar_records = decision.get('similar_records', [])
                            
                            record_num = ''
                            if compare_key:
                                # Пытаемся извлечь номер из ключа вида "frp_text_skill_0001" или "codifier_skill_0001"
                                parts = compare_key.split('_')
                                if len(parts) >= 2 and parts[-1].isdigit():
                                    record_num = f" (№{parts[-1]})"
                                elif '_skill_' in compare_key or '_content_' in compare_key:
                                    # Извлекаем номер после skill_ или content_
                                    import re
                                    match = re.search(r'(?:skill_|content_)(\d+)', compare_key)
                                    if match:
                                        record_num = f" (№{match.group(1)})"
                            
                            st.write(f"**Исследуемая запись{record_num}:**")
                            st.text_area("Текст записи:", value=record_text, height=60, key=f"text_section_{compare_key}", disabled=True)
                            
                            # Показываем 3 самые ближайшие эталонные записи
                            if similar_records:
                                st.write("**Три самые ближайшие эталонные записи:**")
                                for i, sim_rec in enumerate(similar_records, 1):
                                    rec = sim_rec.get('record', {})
                                    rec_text = rec.get('text', '')
                                    similarity = sim_rec.get('similarity', 0)
                                    section = rec.get('section', 'не указан')
                                    topic = rec.get('topic', 'не указана')
                                    
                                    st.write(f"{i}. **{rec_text}**")
                                    st.write(f"   Раздел: {section} | Тема: {topic} | Похожесть: {similarity:.1%}")
                            # Анализируем разделы и темы похожих записей с подсчётом частоты и хранением текстов записей
                            sections_count = {}
                            topics_count = {}
                            topics_records = {}  # Словарь: тема -> список текстов записей
                            for sim_rec in similar_records:
                                rec = sim_rec.get('record', {})
                                sec = rec.get('section', '').strip()
                                top = rec.get('topic', '').strip()
                                rec_text = rec.get('text', '').strip()
                                if sec:
                                    sections_count[sec] = sections_count.get(sec, 0) + 1
                                if top:
                                    topics_count[top] = topics_count.get(top, 0) + 1
                                    # Сохраняем текст записи для этой темы
                                    if top not in topics_records:
                                        topics_records[top] = []
                                    if rec_text:
                                        topics_records[top].append(rec_text)
                            
                            current_section = decision['record'].get('section', '').strip()
                            current_topic = decision['record'].get('topic', '').strip()
                            
                            st.write("**Текущая разметка записи:**")
                            st.write(f"- Раздел: {current_section if current_section else 'не размечено'}")
                            st.write(f"- Тема: {current_topic if current_topic else 'не размечено'}")
                            
                            if sections_count or topics_count:
                                st.write("**Разметка похожих записей:**")
                                # Сортируем по частоте (от большей к меньшей)
                                unique_sections = sorted(sections_count.items(), key=lambda x: x[1], reverse=True)
                                unique_sections = [s[0] for s in unique_sections]
                                unique_topics = sorted(topics_count.items(), key=lambda x: x[1], reverse=True)
                                unique_topics = [t[0] for t in unique_topics]
                                
                                # Показываем частоты для информации с полными текстами записей
                                sections_info = [f"{s} ({sections_count[s]})" for s in unique_sections]
                                topics_info = []
                                for t in unique_topics:
                                    count = topics_count[t]
                                    # Формируем список текстов записей для этой темы
                                    records_texts = topics_records.get(t, [])
                                    if records_texts:
                                        # Показываем все записи полностью
                                        texts_str = "; ".join([f'"{text}"' for text in records_texts])
                                        topics_info.append(f"{t} ({count}): {texts_str}")
                                    else:
                                        topics_info.append(f"{t} ({count})")
                                
                                st.write(f"- Разделы похожих: {', '.join(sections_info) if sections_info else 'нет'}")
                                st.write(f"- Темы похожих:")
                                for topic_info in topics_info:
                                    st.write(f"  • {topic_info}")
                                
                                # Проверяем, все ли одинаковые
                                if len(unique_sections) == 1 and len(unique_topics) == 1:
                                    # Все три эталонные записи имеют одинаковые раздел и тему - автоматически присваиваем
                                    assigned_section = unique_sections[0]
                                    assigned_topic = unique_topics[0]
                                    
                                    st.success(f"✅ Все три ближайшие эталонные записи имеют одинаковые раздел и тему. Присвоено: Раздел: {assigned_section} | Тема: {assigned_topic}")
                                    
                                    decision['record']['section'] = assigned_section
                                    decision['record']['topic'] = assigned_topic
                                    
                                    # Определяем префикс на основе источника
                                    compare_sources = decision['record'].get('sources', [])
                                    if compare_sources and 'фрп_текст' in compare_sources:
                                        prefix = 'frp_text_'
                                    elif compare_sources and 'кодификатор' in compare_sources:
                                        prefix = 'codifier_'
                                    else:
                                        prefix = 'merged_'
                                    
                                    # Находим максимальный номер среди существующих ключей с таким префиксом
                                    max_num = 0
                                    for k in st.session_state.compare_base_data.keys():
                                        if k.startswith(prefix) and ('skill_' in k or 'content_' in k):
                                            try:
                                                parts = k.split('_')
                                                if len(parts) >= 3 and parts[-1].isdigit():
                                                    max_num = max(max_num, int(parts[-1]))
                                            except:
                                                pass
                                    
                                    # Добавляем в merged_data
                                    new_key = f"{prefix}skill_{max_num + 1:04d}"
                                    st.session_state.compare_base_data[new_key] = decision['record']
                                    
                                    st.session_state.compare_stats['section_assigned_auto'] += 1
                                    st.session_state.compare_report.append({
                                        'action': 'section_assigned_auto',
                                        'text': record_text,
                                        'section': assigned_section,
                                        'topic': assigned_topic
                                    })
                                    if compare_key in st.session_state.compare_for_section_topic:
                                        del st.session_state.compare_for_section_topic[compare_key]
                                    if _check_and_transition_next_iteration():
                                        st.success("✅ Все сравнения завершены!" if st.session_state.compare_merged_result else "Переход к следующему источнику...")
                                    st.rerun()
                                else:
                                    # Разные - запрашиваем у пользователя
                                    # Добавляем исходный раздел/тему, если их нет в списке
                                    if current_section and current_section not in unique_sections:
                                        unique_sections = [current_section] + unique_sections
                                        sections_count[current_section] = 0  # для метки "(исходный)"
                                    if current_topic and current_topic not in unique_topics:
                                        unique_topics = [current_topic] + unique_topics
                                        topics_count[current_topic] = 0  # для метки "(исходный)"
                                    # Используем радиобаттоны для выбора раздела
                                    if unique_sections:
                                        st.write("**Выберите раздел:**")
                                        section_options = [''] + unique_sections
                                        section_labels = ['Не указывать раздел'] + [
                                            f"{s} ({sections_count.get(s, 0)}) [исходный]" if s == current_section
                                            else f"{s} ({sections_count.get(s, 0)})"
                                            for s in unique_sections
                                        ]
                                        
                                        section_key = f"section_{compare_key}"
                                        # Проверяем, есть ли уже значение в session_state
                                        if section_key in st.session_state:
                                            default_section_index = st.session_state[section_key]
                                            # Проверяем, что индекс в допустимом диапазоне
                                            if default_section_index >= len(section_options):
                                                default_section_index = 0
                                        else:
                                            # Используем index только если значения нет в session_state
                                            default_section_index = 0 if not current_section else (section_options.index(current_section) if current_section in section_options else 0)
                                        
                                        # Дополнительная проверка на всякий случай
                                        default_section_index = max(0, min(default_section_index, len(section_options) - 1))
                                        
                                        selected_section_idx = st.radio(
                                            "",
                                            options=range(len(section_options)),
                                            format_func=lambda i: section_labels[i],
                                            index=default_section_index,
                                            key=section_key
                                        )
                                        selected_section = section_options[selected_section_idx]
                                    else:
                                        selected_section = ''
                                    
                                    # Используем радиобаттоны для выбора темы (только названия, без текстов)
                                    if unique_topics:
                                        st.write("**Выберите тему:**")
                                        topic_options = [''] + unique_topics
                                        topic_labels = ['Не указывать тему'] + [
                                            f"{t} ({topics_count.get(t, 0)}) [исходный]" if t == current_topic
                                            else f"{t} ({topics_count.get(t, 0)})"
                                            for t in unique_topics
                                        ]
                                        
                                        topic_key = f"topic_{compare_key}"
                                        # Проверяем, есть ли уже значение в session_state
                                        if topic_key in st.session_state:
                                            default_topic_index = st.session_state[topic_key]
                                            # Проверяем, что индекс в допустимом диапазоне
                                            if default_topic_index >= len(topic_options):
                                                default_topic_index = 0
                                        else:
                                            # Используем index только если значения нет в session_state
                                            default_topic_index = 0 if not current_topic else (topic_options.index(current_topic) if current_topic in topic_options else 0)
                                        
                                        # Дополнительная проверка на всякий случай
                                        default_topic_index = max(0, min(default_topic_index, len(topic_options) - 1))
                                        
                                        selected_topic_idx = st.radio(
                                            "",
                                            options=range(len(topic_options)),
                                            format_func=lambda i: topic_labels[i],
                                            index=default_topic_index,
                                            key=topic_key
                                        )
                                        selected_topic = topic_options[selected_topic_idx]
                                    else:
                                        selected_topic = ''
                                    
                                    if st.button("💾 Сохранить разметку", key=f"save_section_{compare_key}"):
                                        decision['record']['section'] = selected_section
                                        decision['record']['topic'] = selected_topic
                                        
                                        # Определяем префикс на основе источника
                                        compare_sources = decision['record'].get('sources', [])
                                        if compare_sources and 'фрп_текст' in compare_sources:
                                            prefix = 'frp_text_'
                                        elif compare_sources and 'кодификатор' in compare_sources:
                                            prefix = 'codifier_'
                                        else:
                                            prefix = 'merged_'
                                        
                                        # Находим максимальный номер среди существующих ключей с таким префиксом
                                        max_num = 0
                                        for k in st.session_state.compare_base_data.keys():
                                            if k.startswith(prefix) and ('skill_' in k or 'content_' in k):
                                                try:
                                                    parts = k.split('_')
                                                    if len(parts) >= 3 and parts[-1].isdigit():
                                                        max_num = max(max_num, int(parts[-1]))
                                                except:
                                                    pass
                                        
                                        # Добавляем в merged_data
                                        new_key = f"{prefix}skill_{max_num + 1:04d}"
                                        st.session_state.compare_base_data[new_key] = decision['record']
                                        
                                        st.session_state.compare_stats['section_assigned_user'] += 1
                                        st.session_state.compare_report.append({
                                            'action': 'section_assigned_user',
                                            'text': record_text,
                                            'section': selected_section,
                                            'topic': selected_topic
                                        })
                                        
                                        if compare_key in st.session_state.compare_for_section_topic:
                                            del st.session_state.compare_for_section_topic[compare_key]
                                        if _check_and_transition_next_iteration():
                                            st.success("✅ Все сравнения завершены!" if st.session_state.compare_merged_result else "Переход к следующему источнику...")
                                        st.rerun()
                
                # 4. Выбор раздела (совпадений нет, раздел неизвестен)
                if st.session_state.compare_for_section_only:
                    st.subheader("Записи без совпадений, требуется выбрать раздел")
                for compare_key, decision in list(st.session_state.compare_for_section_only.items()):
                            record_text = decision['record'].get('text', '')
                            subject = decision.get('subject', '')
                            class_num = decision.get('class', '0')
                            record_num = ''
                            if compare_key:
                                parts = compare_key.split('_')
                                if len(parts) >= 2 and parts[-1].isdigit():
                                    record_num = f" (№{parts[-1]})"
                                elif '_skill_' in compare_key or '_content_' in compare_key:
                                    import re
                                    match = re.search(r'(?:skill_|content_)(\d+)', compare_key)
                                    if match:
                                        record_num = f" (№{match.group(1)})"
                            st.write(f"**Запись{record_num}:**")
                            st.text_area("Текст:", value=record_text, height=80, key=f"text_new_{compare_key}", disabled=True)
                            all_sections = decision.get('base_sections', [])
                            if not all_sections:
                                base_source = st.session_state.get('compare_etalon_data') or st.session_state.compare_base_data
                                base_records_list = [
                                    r for r in base_source.values()
                                    if (r.get('subject', '').strip().lower() == subject.strip().lower() and
                                        (class_num == '0' or (str(r.get('class', '')).strip() or '0') == class_num))
                                ]
                                sections_count = {}
                                for rec in base_records_list:
                                    s = (rec.get('section') or '').strip()
                                    if s:
                                        sections_count[s] = sections_count.get(s, 0) + 1
                                all_sections = sorted(sections_count.keys(), key=lambda x: (-sections_count.get(x, 0), x))
                                if not all_sections:
                                    for rec in base_source.values():
                                        s = (rec.get('section') or '').strip()
                                        if s:
                                            sections_count[s] = sections_count.get(s, 0) + 1
                                    all_sections = sorted(sections_count.keys(), key=lambda x: (-sections_count.get(x, 0), x))
                            current_section = decision['record'].get('section', '').strip()
                            current_topic = decision['record'].get('topic', '').strip()
                            st.write(f"Текущий раздел: {current_section or 'не указан'}" + (f" | Тема: {current_topic}" if current_topic else ""))
                            if all_sections:
                                section_options = [''] + all_sections
                                if current_section and current_section not in section_options:
                                    section_options = ['', current_section] + [s for s in all_sections]
                                default_idx = section_options.index(current_section) if current_section in section_options else 0
                                selected_section = st.selectbox(
                                    "Выберите раздел:",
                                    options=section_options,
                                    index=default_idx,
                                    key=f"new_section_{compare_key}",
                                    format_func=lambda x, cs=current_section: "Не указывать" if not x else (f"{x} (исходный)" if x == cs else x)
                                )
                            else:
                                selected_section = ''
                                st.info("Нет доступных разделов")
                            if st.button("💾 Сохранить", key=f"save_new_{compare_key}"):
                                decision['record']['section'] = selected_section
                                compare_sources = decision['record'].get('sources', [])
                                prefix = 'frp_text_' if (compare_sources and 'фрп_текст' in compare_sources) else 'codifier_' if (compare_sources and 'кодификатор' in compare_sources) else 'merged_'
                                max_num = 0
                                for k in st.session_state.compare_base_data.keys():
                                    if k.startswith(prefix) and ('skill_' in k or 'content_' in k):
                                        try:
                                            parts = k.split('_')
                                            if len(parts) >= 3 and parts[-1].isdigit():
                                                max_num = max(max_num, int(parts[-1]))
                                        except: pass
                                new_key = f"{prefix}skill_{max_num + 1:04d}"
                                st.session_state.compare_base_data[new_key] = decision['record']
                                st.session_state.compare_stats['section_assigned_user'] += 1
                                st.session_state.compare_report.append({
                                    'action': 'new_record_assigned',
                                    'compare_key': compare_key,
                                    'text': record_text,
                                    'section': selected_section,
                                    'topic': decision['record'].get('topic', '')
                                })
                                if compare_key in st.session_state.compare_for_section_only:
                                    del st.session_state.compare_for_section_only[compare_key]
                                if _check_and_transition_next_iteration():
                                    st.success("✅ Все сравнения завершены!" if st.session_state.compare_merged_result else "Переход к следующему источнику...")
                                st.rerun()
                
                # Когда все три словаря пусты (только фейлы могут остаться) — кнопка для перехода
                if not (st.session_state.compare_for_choice or st.session_state.compare_for_section_topic or st.session_state.compare_for_section_only):
                    if st.button("➡️ Перейти к следующей итерации или завершить", key="proceed_next_iter"):
                        if _check_and_transition_next_iteration():
                            st.success("✅ Сравнение завершено!" if st.session_state.compare_merged_result else "Переход к следующему источнику...")
                        st.rerun()
            
            else:
                # Если уже есть результат и нечего обрабатывать — не запускаем обработку заново (избегаем бесконечного цикла)
                already_done = (
                    st.session_state.get('compare_merged_result') is not None
                    and not st.session_state.get('compare_compare_data')
                    and not st.session_state.get('compare_next_data')
                )
                # Запускаем первую итерацию сравнения (только когда есть что обрабатывать)
                if not already_done and 'compare_base_data' in st.session_state:
                    base_data = st.session_state.compare_base_data
                    # Копируем, чтобы не портить исходные данные
                    compare_data = copy.deepcopy(st.session_state.compare_compare_data)
                    st.session_state.compare_compare_data = compare_data
                    
                    # Извлекаем записи без предмета/класса/текста
                    fails = extract_fails_and_clean(compare_data)
                    st.session_state.compare_fails.update(fails)
                    
                    etalon = st.session_state.get('compare_etalon_data') or base_data
                    merged_data, for_choice, for_section_topic, for_section_only = process_comparison_iteration(
                        base_data, compare_data,
                        st.session_state.compare_report,
                        st.session_state.compare_stats,
                        etalon_data=etalon,
                        simple_mode=st.session_state.get('compare_simple_mode', False)
                    )
                    
                    st.session_state.compare_base_data = merged_data
                    st.session_state.compare_for_choice = for_choice
                    st.session_state.compare_for_section_topic = for_section_topic
                    st.session_state.compare_for_section_only = for_section_only
                    
                    has_pending = for_choice or for_section_topic or for_section_only
                    if has_pending:
                        st.rerun()
                    else:
                        # Нет решений - переходим к следующей итерации или завершаем
                        if st.session_state.compare_next_data:
                            st.info("Первая итерация завершена без решений. Переходим к сравнению с кодификатором...")
                            st.session_state.compare_iteration = 2
                            st.session_state.compare_compare_data = copy.deepcopy(st.session_state.compare_next_data)
                            st.session_state.compare_next_data = None
                            
                            compare_data2 = st.session_state.compare_compare_data
                            fails2 = extract_fails_and_clean(compare_data2)
                            st.session_state.compare_fails.update(fails2)
                            
                            st.session_state.compare_etalon_data = copy.deepcopy(merged_data)
                            base_data2 = copy.deepcopy(merged_data)
                            merged_data2, for_choice2, for_section_topic2, for_section_only2 = process_comparison_iteration(
                                base_data2, compare_data2,
                                st.session_state.compare_report,
                                st.session_state.compare_stats,
                                etalon_data=st.session_state.compare_etalon_data,
                                simple_mode=st.session_state.get('compare_simple_mode', False)
                            )
                            
                            st.session_state.compare_base_data = merged_data2
                            st.session_state.compare_for_choice = for_choice2
                            st.session_state.compare_for_section_topic = for_section_topic2
                            st.session_state.compare_for_section_only = for_section_only2
                            
                            if for_choice2 or for_section_topic2 or for_section_only2:
                                st.rerun()
                            else:
                                st.session_state.compare_merged_result = merged_data2
                                st.success("✅ Все сравнения завершены!")
                            st.rerun()
                        else:
                            st.session_state.compare_merged_result = merged_data
                            st.success("✅ Сравнение завершено!")
                            st.rerun()
            
            if st.session_state.compare_stats:
                stats = st.session_state.compare_stats
                st.subheader("Статистика")
                col1, col2, col3, col4, col5 = st.columns(5)
                with col1:
                    st.metric("Автообъединено", stats.get('auto_merged', 0))
                with col2:
                    st.metric("Выбрано пользователем", stats.get('user_selected', 0))
                with col3:
                    st.metric("Сохранено оба", stats.get('both_saved', 0))
                with col4:
                    st.metric("Раздел авто", stats.get('section_assigned_auto', 0))
                with col5:
                    st.metric("Раздел вручную", stats.get('section_assigned_user', 0))
            
            # Кнопка «Сохранить объединённое» — всегда доступна, при нажатии собирает всё в один JSON
            st.subheader("Результаты")
            if st.button("💾 Сохранить объединённое", key="save_merged_btn"):
                # 1. Берём всё обработанное (base_data)
                base_data = st.session_state.get('compare_base_data', {})
                all_records = list(base_data.values())
                # 2. Добавляем необработанное: for_choice, for_section_topic, for_section_only
                for compare_key, decision in st.session_state.get('compare_for_choice', {}).items():
                    rec = decision.get('compare_record', decision.get('record'))
                    if rec:
                        all_records.append(rec.copy())
                for compare_key, decision in st.session_state.get('compare_for_section_topic', {}).items():
                    rec = decision.get('record', {})
                    if rec:
                        all_records.append(rec.copy())
                for compare_key, decision in st.session_state.get('compare_for_section_only', {}).items():
                    rec = decision.get('record', {})
                    if rec:
                        all_records.append(rec.copy())
                # 3. Добавляем фейлы как есть
                for fk, frec in st.session_state.get('compare_fails', {}).items():
                    all_records.append(frec.copy())
                # Сквозная нумерация
                merged = {f"skill_{i+1:04d}": rec for i, rec in enumerate(all_records)}
                st.session_state.compare_merged_result = merged
                st.rerun()
            
            if st.session_state.compare_merged_result:
                merged_json = json.dumps(st.session_state.compare_merged_result, ensure_ascii=False, indent=2)
                st.download_button(
                    "📥 Скачать объединённый JSON",
                    merged_json.encode('utf-8'),
                    file_name="merged_compared.json",
                    mime="application/json",
                    key='dl_merged_compared'
                )
                report_text = "\n".join([
                    f"{i+1}. {json.dumps(item, ensure_ascii=False)}"
                    for i, item in enumerate(st.session_state.compare_report)
                ])
                st.download_button(
                    "📄 Скачать отчёт",
                    report_text.encode('utf-8'),
                    file_name="comparison_report.txt",
                    mime="text/plain",
                    key='dl_report'
                )

# ============ РЕЖИМ: Структурирование с помощью LLM ============
elif mode == 'llm_structure':
    st.header("🤖 Структурирование содержания с помощью LLM")
    st.caption("Загрузите JSON файл с элементами содержания. Модель предложит логичное разделение на разделы и темы.")

    # Инициализация session_state для LLM режима
    if 'llm_content_data' not in st.session_state:
        st.session_state.llm_content_data = None
    if 'llm_grouped_data' not in st.session_state:
        st.session_state.llm_grouped_data = None
    if 'llm_frp_structure' not in st.session_state:
        st.session_state.llm_frp_structure = None
    if 'llm_formatted_text' not in st.session_state:
        st.session_state.llm_formatted_text = None
    if 'llm_results' not in st.session_state:
        st.session_state.llm_results = {}  # {pair_key: [records]}
    if 'llm_raw_responses' not in st.session_state:
        st.session_state.llm_raw_responses = {}  # {pair_key: str}
    if 'llm_prompt_template' not in st.session_state:
        st.session_state.llm_prompt_template = """Изучи представленные элементы содержания из ФРП (федеральная рабочая программа) и предложи логичное разделение на разделы и темы внутри разделов, опираясь на все доступные элементы содержания. Твои разделы и темы должны полностью покрывать все элементы, однако не обязательно повторять названия тех разделов и тем, которые представлены сейчас. Ты можешь объединять или разделять разделы и темы, как считаешь логичным. В пределах темы элементы содержания обязательно должны быть логично связаны; существенно разные элементы содержания не должны оказаться в одной теме.
Составь список разделов и тем внутри каждого раздела, оформи этот список в формате JSON со следующими полями:
- "section": название раздела, который ты предлагаешь
- "frp_section": название раздела из ФРП, внутри которого должен быть расположен этот раздел, или просто сам раздел ФРП, если ты считаешь, что название было подходящим (если соответствие невозможно, оставляй пустым)
- "topic": название темы внутри этого раздела, который ты предлагаешь
- "frp_topic": похожая/такая же/более охватывающая тема из ФРП (если есть, иначе пустая строка)

Верни результат в формате JSON массива объектов без посторонних символов и своих комментариев."""
    if 'llm_custom_prompt' not in st.session_state:
        st.session_state.llm_custom_prompt = None

    def _build_pair_text(subject, class_num):
        """Формирует текст запроса для одной пары предмет+класс."""
        pair_text_lines = []
        pair_text_lines.append(f"предмет: {subject}")
        pair_text_lines.append(f"класс: {class_num}")

        sections = st.session_state.llm_grouped_data[subject][class_num]

        # Записи с разделом и темой
        for section in sorted(sections.keys()):
            if section == 'без раздела':
                continue
            topics = sections[section]
            for topic in sorted(topics.keys()):
                if topic == 'без темы':
                    continue
                pair_text_lines.append(f"раздел: {section}")
                pair_text_lines.append(f"тема: {topic}")
                processed_texts = []
                for text in topics[topic]:
                    text = text.strip()
                    if text:
                        if not text.rstrip().endswith(('.', '!', '?')):
                            text = text.rstrip() + '.'
                        processed_texts.append(text)
                pair_text_lines.append(' '.join(processed_texts))
                pair_text_lines.append("")

        # Записи с разделом, но без темы
        for section in sorted(sections.keys()):
            if section == 'без раздела':
                continue
            if 'без темы' in sections[section]:
                pair_text_lines.append(f"раздел: {section}")
                pair_text_lines.append("тема: без темы")
                processed_texts = []
                for text in sections[section]['без темы']:
                    text = text.strip()
                    if text:
                        if not text.rstrip().endswith(('.', '!', '?')):
                            text = text.rstrip() + '.'
                        processed_texts.append(text)
                pair_text_lines.append(' '.join(processed_texts))
                pair_text_lines.append("")

        # Записи без раздела
        if 'без раздела' in sections:
            pair_text_lines.append("раздел: без раздела")
            all_no_section = []
            for topic_texts in sections['без раздела'].values():
                all_no_section.extend(topic_texts)
            if all_no_section:
                processed_texts = []
                for text in all_no_section:
                    text = text.strip()
                    if text:
                        if not text.rstrip().endswith(('.', '!', '?')):
                            text = text.rstrip() + '.'
                        processed_texts.append(text)
                pair_text_lines.append(' '.join(processed_texts))
                pair_text_lines.append("")

        pair_text = '\n'.join(pair_text_lines)

        # Добавляем информацию о ФРП разделах в начало
        if st.session_state.llm_frp_structure:
            frp_info_lines = ["По фрп имеются следующие разделы и темы:"]
            if subject in st.session_state.llm_frp_structure:
                frp_sections_src = st.session_state.llm_frp_structure[subject]
            else:
                frp_sections_src = {}
                for frp_sections in st.session_state.llm_frp_structure.values():
                    frp_sections_src.update(frp_sections)
            for sec, tops in frp_sections_src.items():
                frp_info_lines.append(f"раздел: {sec}")
                for top in tops:
                    frp_info_lines.append(f"  {top}")
            pair_text = '\n'.join(frp_info_lines) + '\n\n' + pair_text

        return pair_text

    def _records_from_editor(edited_df, original_records, subject, class_num):
        """Возвращает список записей из отредактированного DataFrame."""
        updated = []
        for i, (_, row) in enumerate(edited_df.iterrows()):
            if i < len(original_records):
                base = original_records[i].copy()
            else:
                base = {'subject': subject, 'class': class_num, 'text': '', 'sources': ['llm_structure']}
            base['section'] = str(row.get('Раздел', '')).strip()
            base['frp_section'] = str(row.get('Раздел ФРП', '')).strip()
            base['topic'] = str(row.get('Тема', '')).strip()
            base['frp_topic'] = str(row.get('Тема ФРП', '')).strip()
            updated.append(base)
        return updated

    # --- Загрузка файла ---
    uploaded_file = st.file_uploader("Загрузите JSON файл с элементами содержания", type=['json'], key='llm_upload')

    if uploaded_file:
        try:
            data = json.loads(uploaded_file.read().decode('utf-8'))
            sample_key = next(iter(data.keys()), '')
            if 'content' not in sample_key.lower():
                st.error("❌ Ошибка: Загруженный файл не содержит элементы содержания (content). Проверьте формат файла.")
            else:
                st.session_state.llm_content_data = data
                st.success(f"✅ Файл загружен: {len(data)} записей")

                if st.button("📊 Подготовить данные для анализа", type="primary", key='llm_prepare'):
                    with st.spinner("Группировка данных..."):
                        grouped = group_content_by_structure(data)
                        frp_structure = get_frp_sections_and_topics(data)
                        formatted_text = format_content_text(grouped, frp_structure)
                        st.session_state.llm_grouped_data = grouped
                        st.session_state.llm_frp_structure = frp_structure
                        st.session_state.llm_formatted_text = formatted_text
                        st.success("✅ Данные подготовлены!")
                        st.rerun()
        except json.JSONDecodeError as e:
            st.error(f"❌ Ошибка при чтении JSON файла: {e}")
        except Exception as e:
            st.error(f"❌ Ошибка: {e}")
            import traceback
            st.code(traceback.format_exc())

    # --- Подготовленный текст и настройки ---
    if st.session_state.llm_formatted_text:
        st.subheader("Подготовленный текст для анализа")
        with st.expander("📝 Просмотр текста", expanded=False):
            st.text_area("Текст", value=st.session_state.llm_formatted_text, height=300, disabled=True, key='llm_text_view')

        # Редактирование промпта
        if st.button("✏️ Редактировать промпт", key='llm_edit_prompt'):
            st.session_state.llm_show_prompt_editor = True

        if st.session_state.get('llm_show_prompt_editor', False):
            st.subheader("Редактирование промпта")
            edited_prompt = st.text_area(
                "Промпт для модели",
                value=st.session_state.llm_custom_prompt or st.session_state.llm_prompt_template,
                height=200,
                key='llm_prompt_editor'
            )
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Сохранить", key='llm_save_prompt'):
                    st.session_state.llm_custom_prompt = edited_prompt
                    st.session_state.llm_show_prompt_editor = False
                    st.success("Промпт сохранен!")
                    st.rerun()
            with col2:
                if st.button("↩️ Вернуться к исходному", key='llm_reset_prompt'):
                    st.session_state.llm_custom_prompt = None
                    st.session_state.llm_show_prompt_editor = False
                    st.rerun()

        # --- Проверка API ключа ---
        api_key = get_claude_api_key()
        if not api_key:
            st.warning("⚠️ Для работы с LLM необходимо добавить ключ CLAUDE_API_KEY в секреты Streamlit.")
        else:
            if st.session_state.llm_grouped_data:
                # Список пар предмет+класс
                subject_class_pairs = []
                for subject in sorted(st.session_state.llm_grouped_data.keys()):
                    for class_num in sorted(
                        st.session_state.llm_grouped_data[subject].keys(),
                        key=lambda x: int(x) if str(x).isdigit() else 0
                    ):
                        subject_class_pairs.append((subject, class_num))

                if subject_class_pairs:
                    total_pairs = len(subject_class_pairs)
                    done_pairs = sum(1 for s, c in subject_class_pairs if f"{s}_{c}" in st.session_state.llm_results)
                    unprocessed = [(s, c) for s, c in subject_class_pairs if f"{s}_{c}" not in st.session_state.llm_results]

                    st.subheader("Обработка данных")
                    st.write(f"Групп предмет+класс: **{total_pairs}**, обработано: **{done_pairs}**")

                    # --- Кнопка «Обработать все» ---
                    if unprocessed:
                        if st.button("🚀 Обработать все группы с помощью LLM", type="primary", key='llm_run_all'):
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            errors_list = []

                            prompt = st.session_state.llm_custom_prompt or st.session_state.llm_prompt_template
                            verify_ssl = st.session_state.get('claude_verify_ssl', True)
                            model = st.session_state.get('claude_working_model', 'claude-sonnet-4-20250514')
                            api_version = st.session_state.get('claude_working_api_version', '2023-06-01')

                            for i, (subject, class_num) in enumerate(subject_class_pairs):
                                pair_key = f"{subject}_{class_num}"
                                if pair_key in st.session_state.llm_results:
                                    progress_bar.progress((i + 1) / total_pairs)
                                    continue

                                status_text.write(
                                    f"⏳ Обрабатываю: **{subject}**, класс **{class_num}** "
                                    f"({i + 1}/{total_pairs})..."
                                )

                                pair_text = _build_pair_text(subject, class_num)
                                full_prompt = prompt + "\n\n" + pair_text
                                messages = [{"role": "user", "content": full_prompt}]

                                response = call_claude_api(
                                    messages, api_key,
                                    model=model, api_version=api_version, verify_ssl=verify_ssl
                                )

                                if response:
                                    st.session_state.llm_raw_responses[pair_key] = response
                                    records = parse_llm_response(response, subject, class_num)
                                    if records:
                                        st.session_state.llm_results[pair_key] = records
                                    else:
                                        errors_list.append(
                                            f"Не удалось распарсить ответ для {subject}, {class_num}"
                                        )
                                        st.session_state.llm_results[pair_key] = []
                                else:
                                    st.session_state.llm_raw_responses[pair_key] = "(нет ответа от API)"
                                    errors_list.append(f"Ошибка API для {subject}, {class_num}")

                                progress_bar.progress((i + 1) / total_pairs)

                            status_text.empty()
                            if errors_list:
                                for err in errors_list:
                                    st.error(err)
                            else:
                                st.success("✅ Все группы успешно обработаны!")
                            st.rerun()
                    else:
                        st.success(f"✅ Все группы обработаны ({total_pairs})")
                        if st.button("🔄 Сбросить и обработать заново", key='llm_rerun_all'):
                            st.session_state.llm_results = {}
                            st.session_state.llm_raw_responses = {}
                            st.rerun()

                    # --- Отображение результатов ---
                    if st.session_state.llm_results:
                        st.markdown("---")
                        st.subheader("📋 Результаты структурирования")

                        # Словарь {pair_key: (edited_df, original_records, subject, class_num)}
                        # заполняется по мере рендера таблиц; используется кнопками внизу
                        all_edited = {}

                        for idx, (subject, class_num) in enumerate(subject_class_pairs):
                            pair_key = f"{subject}_{class_num}"
                            if pair_key not in st.session_state.llm_results:
                                continue

                            records = st.session_state.llm_results[pair_key]

                            st.subheader(f"📌 {subject} — {class_num} класс")

                            if not records:
                                st.warning("Нет данных (ошибка при обработке этой группы).")
                                raw = st.session_state.llm_raw_responses.get(pair_key)
                                if raw:
                                    with st.expander("🔍 Сырой ответ модели", expanded=True):
                                        st.text_area("", value=raw, height=250, disabled=True,
                                                     key=f'llm_raw_err_{pair_key}')
                                continue

                            df_data = []
                            for i, rec in enumerate(records):
                                df_data.append({
                                    '№': i + 1,
                                    'Раздел': rec.get('section', ''),
                                    'Тема': rec.get('topic', ''),
                                    'Раздел ФРП': rec.get('frp_section', ''),
                                    'Тема ФРП': rec.get('frp_topic', ''),
                                })

                            df = pd.DataFrame(df_data)

                            edited_df = st.data_editor(
                                df,
                                use_container_width=True,
                                key=f'llm_editor_{pair_key}',
                                num_rows="dynamic",
                                column_config={
                                    '№': st.column_config.NumberColumn('№', width='small', disabled=True),
                                    'Раздел': st.column_config.TextColumn('Раздел', width='medium'),
                                    'Тема': st.column_config.TextColumn('Тема', width='medium'),
                                    'Раздел ФРП': st.column_config.TextColumn('Раздел ФРП', width='medium'),
                                    'Тема ФРП': st.column_config.TextColumn('Тема ФРП', width='medium'),
                                }
                            )

                            all_edited[pair_key] = (edited_df, records, subject, class_num)

                            if st.button("💾 Сохранить изменения", key=f'llm_save_{pair_key}'):
                                updated = _records_from_editor(edited_df, records, subject, class_num)
                                st.session_state.llm_results[pair_key] = updated
                                st.success("✅ Изменения сохранены!")
                                st.rerun()

                            raw = st.session_state.llm_raw_responses.get(pair_key)
                            if raw:
                                with st.expander("🔍 Сырой ответ модели", expanded=False):
                                    st.text_area("", value=raw, height=250, disabled=True,
                                                 key=f'llm_raw_{pair_key}')

                        # --- Кнопки внизу ---
                        if all_edited:
                            st.markdown("---")
                            bot_col1, bot_col2 = st.columns(2)

                            with bot_col1:
                                if st.button(
                                    "💾 Сохранить все изменения",
                                    type="primary",
                                    key='llm_save_all_btn'
                                ):
                                    for pk, (edf, recs, subj, cls) in all_edited.items():
                                        st.session_state.llm_results[pk] = _records_from_editor(
                                            edf, recs, subj, cls
                                        )
                                    st.success("✅ Все изменения сохранены!")
                                    st.rerun()

                            with bot_col2:
                                # Формируем JSON из текущего состояния редакторов (render-time)
                                current_json = {}
                                counter = 1
                                for pk, (edf, _, subj, cls) in all_edited.items():
                                    for _, row in edf.iterrows():
                                        section_val = str(row.get('Раздел', '')).strip()
                                        topic_val = str(row.get('Тема', '')).strip()
                                        current_json[f"content_{counter:04d}"] = {
                                            'subject': subj,
                                            'class': cls,
                                            'section': section_val,
                                            'frp_section': str(row.get('Раздел ФРП', '')).strip(),
                                            'topic': topic_val,
                                            'frp_topic': str(row.get('Тема ФРП', '')).strip(),
                                            'text': '',
                                            'sources': ['llm_structure'],
                                        }
                                        counter += 1

                                current_json_str = json.dumps(current_json, ensure_ascii=False, indent=2)
                                st.download_button(
                                    "📥 Сохранить файл",
                                    current_json_str.encode('utf-8'),
                                    file_name="llm_structured_content.json",
                                    mime="application/json",
                                    key='llm_download_json'
                                )

# ============ РЕЖИМ: БАЗА ДАННЫХ ============
elif mode == 'db_input':
    st.header("💾 Добавление в базу данных")

    # Проверяем подключение к БД
    _db_url = os.environ.get('DATABASE_URL', '')
    if not _db_url:
        st.error("DATABASE_URL не задан. Создайте файл .env с DATABASE_URL=postgresql://...")
        st.stop()

    # Загружаем справочник frp_topics
    if st.session_state.db_frp_df is None:
        with st.spinner("Загружаю темы ФРП из базы..."):
            st.session_state.db_frp_df = load_frp_topics_cached()
    frp_df: pd.DataFrame = st.session_state.db_frp_df

    if frp_df.empty:
        st.error("Не удалось загрузить таблицу frp_topics. Проверьте подключение к БД.")
        st.stop()

    # ── Раздел 1: Выбор темы ФРП ─────────────────────────────────────────────
    st.subheader("Тема ФРП")

    _col1, _col2, _col3, _col4 = st.columns(4)

    # Предмет
    subjects_df = load_subjects_cached()
    if not subjects_df.empty:
        subjects = sorted(subjects_df['name'].unique())
        subject_id_map = (
            subjects_df[['name', 'id']]
            .drop_duplicates()
            .set_index('name')['id']
            .to_dict()
        )
    else:
        # fallback: если subjects ещё не создана в БД
        subjects = sorted(frp_df['subject'].unique())
        subject_id_map = {}

    sel_subj = _col1.selectbox("Предмет", [''] + subjects, key='db_sel_subject')

    # Класс — фильтрованный по предмету
    if sel_subj:
        classes_df = frp_df[frp_df['subject'] == sel_subj]
    else:
        classes_df = frp_df
    classes = sorted(classes_df['grade_class'].unique(), key=lambda x: int(x) if x.isdigit() else 99)
    sel_class = _col2.selectbox("Класс", [''] + classes, key='db_sel_class')

    # Раздел — фильтрованный по предмету + классу
    if sel_subj and sel_class:
        sections_df = frp_df[(frp_df['subject'] == sel_subj) & (frp_df['grade_class'] == sel_class)]
    elif sel_subj:
        sections_df = frp_df[frp_df['subject'] == sel_subj]
    elif sel_class:
        sections_df = frp_df[frp_df['grade_class'] == sel_class]
    else:
        sections_df = frp_df
    sections = sorted(sections_df['section'].unique())
    sel_section = _col3.selectbox("Раздел", [''] + sections, key='db_sel_section')

    # Тема — фильтрованная по всем трём
    topic_filter = frp_df.copy()
    if sel_subj:
        topic_filter = topic_filter[topic_filter['subject'] == sel_subj]
    if sel_class:
        topic_filter = topic_filter[topic_filter['grade_class'] == sel_class]
    if sel_section:
        topic_filter = topic_filter[topic_filter['section'] == sel_section]
    topics = sorted(topic_filter['topic'].unique())
    sel_topic = _col4.selectbox("Тема", [''] + topics, key='db_sel_topic')

    # Получаем id выбранной темы
    def _get_frp_id(subj, cls, sect, top):
        row = frp_df[
            (frp_df['subject'] == subj) &
            (frp_df['grade_class'] == cls) &
            (frp_df['section'] == sect) &
            (frp_df['topic'] == top)
        ]
        if not row.empty:
            return int(row.iloc[0]['id'])
        return None

    # Авто-фиксация темы при каждом рендере по текущему выбору в дропдаунах
    if sel_subj and sel_class and sel_section and sel_topic:
        _auto_id = _get_frp_id(sel_subj, sel_class, sel_section, sel_topic)
        st.session_state.db_fixed          = True
        st.session_state.db_fixed_topic_id = _auto_id
        st.session_state.db_fixed_label    = (
            f"предмет: **{sel_subj}**, класс: **{sel_class}**, "
            f"раздел: **{sel_section}**, тема: **{sel_topic}**"
        )
    else:
        st.session_state.db_fixed          = False
        st.session_state.db_fixed_topic_id = None
        st.session_state.db_fixed_label    = ''

    # Кнопка "Добавить тему ФРП"
    _btn_col1, _btn_col2 = st.columns([2, 10])
    with _btn_col1:
        if st.button("➕ Добавить тему ФРП", key='db_add_frp_btn', use_container_width=True):
            st.session_state.db_add_frp_open = not st.session_state.db_add_frp_open

    # Форма добавления новой темы ФРП
    if st.session_state.db_add_frp_open:
        with st.container(border=True):
            st.markdown("**Новая тема ФРП**")
            _f1, _f2, _f3, _f4 = st.columns(4)
            new_subj    = _f1.selectbox("Предмет",  subjects, index=(subjects.index(sel_subj) if sel_subj in subjects else 0), key='db_new_subj')
            new_class   = _f2.text_input("Класс",    value=sel_class or '',   key='db_new_class')
            new_section = _f3.text_input("Раздел",   value=sel_section or '', key='db_new_section')
            new_topic   = _f4.text_input("Тема",     value=sel_topic or '',   key='db_new_topic')
            if st.button("💾 Сохранить новую тему", key='db_save_new_frp'):
                _ns = normalize_db_text(new_subj)
                _nc = normalize_db_text(new_class)
                _nse = normalize_db_text(new_section)
                _nt = normalize_db_text(new_topic)
                if _ns and _nc and _nse and _nt:
                    _existing = frp_df[
                        (frp_df['subject'] == _ns) & (frp_df['grade_class'] == _nc) &
                        (frp_df['section'] == _nse) & (frp_df['topic'] == _nt)
                    ]
                    if not _existing.empty:
                        st.warning("Такая комбинация уже есть в базе.")
                    else:
                        _conn = get_db_conn()
                        if _conn:
                            try:
                                _subject_id = subject_id_map.get(_ns)
                                if not _subject_id:
                                    st.error("Не найден subject_id для выбранного предмета. Проверьте таблицу subjects.")
                                    raise RuntimeError("subject_id not found")
                                _cur = _conn.cursor()
                                _cur.execute(
                                    "INSERT INTO frp_topics (grade_class, subject_id, subject_name, section, topic, program) "
                                    "VALUES (%s,%s,%s,%s,%s,%s) RETURNING id",
                                    (_nc, int(_subject_id), _ns, _nse, _nt, 'базовый')
                                )
                                _new_id = _cur.fetchone()[0]
                                _conn.commit()
                                _cur.close()
                                _conn.close()
                                st.success(f"Тема добавлена (id={_new_id}).")
                                load_frp_topics_cached.clear()
                                st.session_state.db_frp_df = None
                                st.session_state.db_add_frp_open = False
                                st.rerun()
                            except Exception as _e:
                                st.error(f"Ошибка сохранения: {_e}")
                        else:
                            st.error("Нет подключения к БД.")
                else:
                    st.warning("Заполните все поля.")

    # Сообщение о фиксации
    if st.session_state.db_fixed:
        st.success(f"Работаем с: {st.session_state.db_fixed_label}")
        _topic_id = st.session_state.db_fixed_topic_id
        if _topic_id:
            _conn_cnt = get_db_conn()
            if _conn_cnt:
                try:
                    with _conn_cnt.cursor() as _cur_cnt:
                        _cur_cnt.execute(
                            "SELECT COUNT(*) FROM skill_defs WHERE frp_topic_id = %s",
                            (_topic_id,)
                        )
                        _skills_cnt = _cur_cnt.fetchone()[0]
                        _cur_cnt.execute(
                            "SELECT COUNT(*) FROM content_element_defs WHERE frp_topic_id = %s",
                            (_topic_id,)
                        )
                        _content_cnt = _cur_cnt.fetchone()[0]
                    st.caption(
                        f"В базе по этой теме: навыков — **{_skills_cnt}**, "
                        f"элементов содержания — **{_content_cnt}**"
                    )
                except Exception:
                    pass
                finally:
                    _conn_cnt.close()

    st.markdown("---")

    # ── Раздел 2: Тип + ввод текста ──────────────────────────────────────────
    _tab_col1, _tab_col2, _tab_col3 = st.columns([2, 2, 8])
    with _tab_col1:
        _skills_active = st.session_state.db_mode_type == 'skills'
        if st.button(
            "📚 Навыки" if not _skills_active else "📚 **Навыки** ✓",
            key='db_tab_skills',
            type='primary' if _skills_active else 'secondary',
            use_container_width=True
        ):
            st.session_state.db_mode_type = 'skills'
            st.rerun()
    with _tab_col2:
        _content_active = st.session_state.db_mode_type == 'content'
        if st.button(
            "📄 Содержание" if not _content_active else "📄 **Содержание** ✓",
            key='db_tab_content',
            type='primary' if _content_active else 'secondary',
            use_container_width=True
        ):
            st.session_state.db_mode_type = 'content'
            st.rerun()

    if st.session_state.db_mode_type is None:
        st.info("Выберите тип данных: Навыки или Содержание.")

    input_text = st.text_area(
        "Вставьте текст для обработки",
        height=180,
        key='db_input_text',
        placeholder="Вставьте текст — он будет разбит на отдельные строки..."
    )

    _proc_col, _all_col, _spacer = st.columns([2, 3, 7])
    with _proc_col:
        _proc_disabled = (not st.session_state.db_fixed
                          or st.session_state.db_mode_type is None)
        if st.button("⚙️ Обработать", key='db_process_btn', type='primary',
                     use_container_width=True, disabled=_proc_disabled):
            if not input_text.strip():
                st.warning("Введите текст.")
            elif st.session_state.db_mode_type is None:
                st.warning("Сначала выберите тип: Навыки или Содержание.")
            else:
                _sentences = split_into_sentences(input_text)
                _new_items = []
                for _s in _sentences:
                    st.session_state.db_uid_counter += 1
                    _new_items.append({
                        'uid': st.session_state.db_uid_counter,
                        'text': _s,
                        'original_frp': _s,
                        'sub_items': [],
                        'llm_done': False,
                    })
                st.session_state.db_items = _new_items
                st.session_state.db_save_result = None
                st.rerun()

    with _all_col:
        _has_unprocessed = bool(st.session_state.db_items) and any(
            not it.get('llm_done') for it in st.session_state.db_items
        )
        if st.button(
            "✨ Доработать всё",
            key='db_atomize_all_btn',
            disabled=not _has_unprocessed,
            use_container_width=True,
        ):
            st.session_state.db_batch_running = True
            st.session_state.db_batch_pos = 0
            st.session_state.db_batch_stop = False
            st.rerun()

    st.markdown("---")

    # ── Раздел 3: Обработанные элементы (фрагмент — не перерисовывает всю страницу) ──
    @st.fragment
    def _items_editor():
        if not st.session_state.db_items:
            return

        _atomize_prompt = load_atomize_prompt(st.session_state.db_mode_type or 'skills')
        _type_label = 'навыки' if st.session_state.db_mode_type == 'skills' else 'элементы содержания'

        # --- Батч "Доработать всё": один элемент за фрагмент-рерун ---
        if st.session_state.get('db_batch_running'):
            _bp  = st.session_state.db_batch_pos
            _bt  = len(st.session_state.db_items)
            _bap = load_atomize_prompt(st.session_state.db_mode_type or 'skills')

            _pb_col, _stop_col = st.columns([6, 1])
            with _pb_col:
                st.progress(
                    _bp / max(_bt, 1),
                    text=f"Обрабатываю {min(_bp + 1, _bt)} / {_bt}…  *(нажмите ⛔ чтобы остановить)*"
                )
            with _stop_col:
                if st.button("⛔ Стоп", key='db_batch_stop_btn', use_container_width=True):
                    st.session_state.db_batch_running = False
                    st.session_state.db_batch_stop = True
                    st.rerun()

            if _bp < _bt:
                _bitem = st.session_state.db_items[_bp]
                if not _bitem.get('llm_done'):
                    _bcur = st.session_state.get(f'db_item_text_{_bitem["uid"]}', _bitem['text'])
                    with st.spinner(f"Элемент {_bp + 1} / {_bt}: {_bcur[:60]}…"):
                        _bmsgs = [{"role": "user", "content": f"{_bap}\n\nТекст: {_bcur}"}]
                        _bresp = call_claude_api(_bmsgs)
                        _accumulate_cost()
                        if _bresp:
                            try:
                                _bm = re.search(r'\{.*\}', _bresp, re.DOTALL)
                                if _bm:
                                    _bp_parsed = json.loads(_bm.group())
                                    _bsub = []
                                    for _ba in _bp_parsed.get('atomic_skills', []):
                                        st.session_state.db_uid_counter += 1
                                        _bsub.append({'uid': st.session_state.db_uid_counter, 'text': _ba})
                                    st.session_state.db_items[_bp]['sub_items'] = _bsub
                                    st.session_state.db_items[_bp]['llm_done'] = True
                                    st.session_state.db_items[_bp]['original_frp'] = _bcur
                            except Exception as _be:
                                st.warning(f"Элемент {_bp + 1}: не удалось разобрать ответ ({_be})")
                st.session_state.db_batch_pos = _bp + 1
                if _bp + 1 >= _bt:
                    st.session_state.db_batch_running = False
                st.rerun()
            else:
                st.session_state.db_batch_running = False
                st.rerun()

            return  # пока идёт батч — не рисуем остальное

        _items_to_delete = []

        for _i, _item in enumerate(st.session_state.db_items):
            _uid = _item['uid']

            with st.container(border=True):
                _row_col1, _row_col2 = st.columns([10, 1])
                with _row_col1:
                    # Значение хранится в st.session_state[key] — не пишем обратно в db_items
                    st.text_area(
                        f"Элемент {_i + 1}",
                        value=_item['text'],
                        height=80,
                        key=f'db_item_text_{_uid}',
                        label_visibility='collapsed'
                    )
                with _row_col2:
                    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
                    if st.button("🗑️", key=f'db_del_item_{_uid}', help="Удалить"):
                        _items_to_delete.append(_uid)

                # Кнопки "Доработать" / "Вручную" — показываем только пока не начали редактировать
                if not _item.get('llm_done'):
                    _llm_col1, _llm_col2, _llm_col3 = st.columns([3, 3, 6])
                    with _llm_col1:
                        if st.button("✨ Доработать", key=f'db_atomize_{_uid}'):
                            _cur_text = st.session_state.get(f'db_item_text_{_uid}', _item['text'])
                            with st.spinner("Запрос к модели..."):
                                _messages = [{"role": "user", "content": f"{_atomize_prompt}\n\nТекст: {_cur_text}"}]
                                _response = call_claude_api(_messages)
                                _accumulate_cost()
                                if _response:
                                    try:
                                        _json_match = re.search(r'\{.*\}', _response, re.DOTALL)
                                        if _json_match:
                                            _parsed = json.loads(_json_match.group())
                                            _atomic = _parsed.get('atomic_skills', [])
                                            _sub = []
                                            for _a in _atomic:
                                                st.session_state.db_uid_counter += 1
                                                _sub.append({'uid': st.session_state.db_uid_counter, 'text': _a})
                                            st.session_state.db_items[_i]['sub_items'] = _sub
                                            st.session_state.db_items[_i]['llm_done'] = True
                                            st.session_state.db_items[_i]['original_frp'] = _cur_text
                                            st.rerun()
                                    except Exception as _e:
                                        st.error(f"Ошибка разбора ответа LLM: {_e}")
                                else:
                                    st.error("LLM не вернул ответ.")
                    with _llm_col2:
                        if st.button("✏️ Вручную", key=f'db_manual_{_uid}'):
                            _cur_text = st.session_state.get(f'db_item_text_{_uid}', _item['text'])
                            st.session_state.db_items[_i]['llm_done'] = True
                            st.session_state.db_items[_i]['sub_items'] = []
                            st.session_state.db_items[_i]['original_frp'] = _cur_text
                            st.rerun()

                # Подэлементы — показываются после "Доработать" или "Вручную"
                if _item.get('llm_done'):
                    if _item.get('sub_items'):
                        st.markdown("**Атомарные элементы:**")
                    _subs_to_delete = []
                    for _j, _sub in enumerate(_item['sub_items']):
                        _sub_uid = _sub['uid']
                        _sc1, _sc2 = st.columns([10, 1])
                        with _sc1:
                            # Значение хранится в st.session_state[key] — не пишем обратно
                            st.text_input(
                                f"sub_{_uid}_{_j}",
                                value=_sub['text'],
                                key=f'db_sub_{_sub_uid}',
                                label_visibility='collapsed'
                            )
                        with _sc2:
                            if st.button("✕", key=f'db_del_sub_{_sub_uid}', help="Удалить"):
                                _subs_to_delete.append(_sub_uid)

                    if _subs_to_delete:
                        st.session_state.db_items[_i]['sub_items'] = [
                            s for s in st.session_state.db_items[_i]['sub_items']
                            if s['uid'] not in _subs_to_delete
                        ]
                        st.rerun()

                    if st.button("＋ Добавить элемент", key=f'db_add_sub_{_uid}'):
                        st.session_state.db_uid_counter += 1
                        st.session_state.db_items[_i]['sub_items'].append({
                            'uid': st.session_state.db_uid_counter, 'text': ''
                        })
                        st.rerun()

        # Удаляем отмеченные элементы
        if _items_to_delete:
            st.session_state.db_items = [it for it in st.session_state.db_items if it['uid'] not in _items_to_delete]
            st.rerun()

        st.markdown("---")

        # ── Кнопка "Сохранить в базу" ─────────────────────────────────────────
        if st.button(f"💾 Сохранить в базу ({_type_label})", type='primary', key='db_save_btn'):
            if not st.session_state.db_fixed:
                st.warning("Зафиксируйте тему ФРП перед сохранением.")
            elif st.session_state.db_mode_type is None:
                st.warning("Выберите тип данных.")
            else:
                st.session_state.db_show_confirm = True
                st.rerun()

        # Диалог подтверждения
        if st.session_state.db_show_confirm:
            with st.container(border=True):
                st.markdown(f"**Сохраняем в {_type_label}?**")
                st.caption(st.session_state.get('db_fixed_label', ''))
                _ok_col, _cancel_col = st.columns(2)
                with _ok_col:
                    if st.button("✅ Ок", key='db_confirm_ok'):
                        # Собираем записи для сохранения
                        _records = []
                        for _item in st.session_state.db_items:
                            _cur_item_text = st.session_state.get(
                                f'db_item_text_{_item["uid"]}', _item['text']
                            )
                            if _item.get('llm_done'):
                                # После доработки (LLM или вручную) — сохраняем подэлементы
                                for _sub in _item.get('sub_items', []):
                                    _cur_sub_text = st.session_state.get(
                                        f'db_sub_{_sub["uid"]}', _sub['text']
                                    )
                                    _t = normalize_db_text(_cur_sub_text)
                                    if _t:
                                        _records.append({
                                            'label': _t,
                                            'frp_label': normalize_db_text(_item.get('original_frp', _cur_item_text))
                                        })
                            else:
                                # Без доработки — сохраняем основной текст напрямую
                                _t = normalize_db_text(_cur_item_text)
                                if _t:
                                    _records.append({
                                        'label': _t,
                                        'frp_label': normalize_db_text(_cur_item_text)
                                    })

                        # Сохраняем в БД
                        _table = 'skill_defs' if st.session_state.db_mode_type == 'skills' else 'content_element_defs'
                        _frp_id = st.session_state.db_fixed_topic_id
                        _conn = get_db_conn()
                        if not _conn:
                            st.error("Нет подключения к БД.")
                        else:
                            try:
                                _cur = _conn.cursor()
                                _inserted = 0
                                for _rec in _records:
                                    _cur.execute(
                                        f"INSERT INTO {_table} (label_normalized, label_display, frp_label, frp_topic_id) "
                                        f"VALUES (%s, %s, %s, %s) ON CONFLICT (label_normalized) DO NOTHING",
                                        (_rec['label'], _rec['label'], _rec['frp_label'], _frp_id)
                                    )
                                    if _cur.rowcount == 1:
                                        _inserted += 1
                                _conn.commit()
                                _cur.close()
                                _conn.close()
                                st.session_state.db_save_result = _inserted
                                st.session_state.db_show_confirm = False
                                st.session_state.db_items = []
                                st.session_state['db_input_text'] = ''
                                st.session_state.db_cost_input_tokens = 0
                                st.session_state.db_cost_output_tokens = 0
                                st.session_state.db_cost_usd = 0.0
                                st.rerun(scope="app")
                            except Exception as _e:
                                st.error(f"Ошибка сохранения: {_e}")

                with _cancel_col:
                    if st.button("❌ Отмена", key='db_confirm_cancel'):
                        st.session_state.db_show_confirm = False
                        st.rerun()

        # Результат сохранения — остаётся до следующего нажатия "Обработать"
        if st.session_state.db_save_result is not None:
            st.success(f"Готово! Добавлено {st.session_state.db_save_result} новых записей.")

    _items_editor()

# ============ РЕЖИМ: ПЕРВИЧНОЕ ТЕГИРОВАНИЕ ============
elif mode == 'tagging_init':
    st.header("🏷️ Первичное извлечение тегов по темам")

    _db_url = os.environ.get('DATABASE_URL', '')
    if not _db_url:
        st.error("DATABASE_URL не задан. Создайте файл .env с DATABASE_URL=postgresql://...")
        st.stop()

    # Миграция 005 нужна для хранения промежуточных результатов
    st.caption("Сырые теги сохраняются в стейджинг-таблицы (до нормализации и создания канонических tags).")

    # Проверим, что применена миграция 005_tag_staging.sql (новая схема: тема -> термины)
    _chk = get_db_conn()
    if _chk is None:
        st.error("Нет подключения к БД.")
        st.stop()
    try:
        with _chk.cursor() as _cur:
            _cur.execute(
                """
                SELECT 1
                FROM information_schema.tables
                WHERE table_schema='public' AND table_name='tag_topic_terms'
                """
            )
            _ok = bool(_cur.fetchone())
        _chk.close()
        if not _ok:
            st.error("Не найдены стейджинг-таблицы для тегов. Выполните миграцию `db/migrations/005_tag_staging.sql` в Neon SQL Editor.")
            st.stop()
    except Exception:
        try:
            _chk.close()
        except Exception:
            pass
        st.error("Не удалось проверить наличие стейджинг-таблиц. Проверьте подключение к БД.")
        st.stop()

    _subjects_df = load_subjects_cached()
    if _subjects_df.empty:
        st.error("Таблица subjects пуста или недоступна. Проверьте миграции.")
        st.stop()

    _subj_names = sorted(_subjects_df['name'].unique())
    _subj_name_to_id = (
        _subjects_df[['name', 'id']]
        .drop_duplicates()
        .set_index('name')['id']
        .to_dict()
    )

    _c1, _c2, _c3, _c4 = st.columns([3, 2, 2, 5])
    with _c1:
        _sel_subj = st.selectbox("Предмет", _subj_names, key='tag_sel_subject_name')
    with _c2:
        _sel_program = st.selectbox("Программа", ['базовый', 'профильный'], key='tag_sel_program')
    with _c3:
        _only_nonempty = st.checkbox("Только темы с записями", value=True, key='tag_sel_only_nonempty')
    with _c4:
        st.caption("Обработка идёт по одной теме за раз: можно смотреть результат и останавливать/продолжать.")

    _sel_subject_id = int(_subj_name_to_id.get(_sel_subj))

    def _load_topics_df(subject_id: int, program: str) -> pd.DataFrame:
        conn = get_db_conn()
        if conn is None:
            return pd.DataFrame()
        try:
            df = pd.read_sql(
                """
                SELECT f.id,
                       f.grade_class,
                       f.section,
                       f.topic,
                       f.program,
                       (SELECT COUNT(*) FROM skill_defs sd WHERE sd.frp_topic_id = f.id) AS skills_cnt,
                       (SELECT COUNT(*) FROM content_element_defs cd WHERE cd.frp_topic_id = f.id) AS content_cnt
                FROM frp_topics f
                WHERE f.subject_id = %s
                  AND f.program = %s
                ORDER BY
                  CASE WHEN f.grade_class ~ '^[0-9]+$' THEN f.grade_class::int ELSE 99 END,
                  f.section,
                  f.topic,
                  f.id
                """,
                conn,
                params=(int(subject_id), str(program)),
            )
            conn.close()
            if df.empty:
                return df
            df['total_cnt'] = df['skills_cnt'].fillna(0).astype(int) + df['content_cnt'].fillna(0).astype(int)
            if _only_nonempty:
                df = df[df['total_cnt'] > 0].reset_index(drop=True)
            return df
        except Exception:
            conn.close()
            return pd.DataFrame()

    if st.button("🔄 Обновить список тем", key='tag_refresh_topics'):
        st.session_state.tag_topic_df = None

    if st.session_state.tag_topic_df is None:
        st.session_state.tag_topic_df = _load_topics_df(_sel_subject_id, _sel_program)

    _topics_df = st.session_state.tag_topic_df
    if _topics_df is None or _topics_df.empty:
        st.warning("Нет тем для выбранного предмета/программы (или нет записей в skill_defs/content_element_defs).")
        st.stop()

    st.dataframe(
        _topics_df[['id', 'grade_class', 'section', 'topic', 'program', 'skills_cnt', 'content_cnt', 'total_cnt']],
        use_container_width=True,
        hide_index=True,
    )

    st.markdown("---")

    _run_col1, _run_col2, _run_col3 = st.columns([3, 3, 6])
    with _run_col1:
        if st.button("🆕 Начать новый прогон", key='tag_new_run_btn', use_container_width=True):
            _conn = get_db_conn()
            if _conn is None:
                st.error("Нет подключения к БД.")
            else:
                try:
                    with _conn.cursor() as _cur:
                        _cur.execute(
                            "INSERT INTO tag_extraction_runs(subject_id, program, status) VALUES (%s, %s, %s) RETURNING id",
                            (_sel_subject_id, _sel_program, 'running'),
                        )
                        _rid = int(_cur.fetchone()[0])
                    _conn.commit()
                    _conn.close()
                    st.session_state.tag_run_id = _rid
                    st.session_state.tag_subject_id = _sel_subject_id
                    st.session_state.tag_program = _sel_program
                    st.session_state.tag_topic_pos = 0
                    st.session_state.tag_stop = False
                    st.session_state.tag_last_result = None
                    st.session_state.tag_last_topic_id = None
                    st.success(f"Создан прогон run_id={_rid}")
                except Exception as _e:
                    try:
                        _conn.close()
                    except Exception:
                        pass
                    st.error(f"Не удалось создать прогон: {_e}")

    with _run_col2:
        _run_id = st.session_state.get('tag_run_id')
        st.metric("Текущий run_id", str(_run_id) if _run_id else "—")

    with _run_col3:
        st.caption("Если закрыли страницу — прогон можно продолжить: run_id хранится в сессии браузера. Позже добавим выбор/резюмирование по run_id из БД.")

    if not st.session_state.get('tag_run_id'):
        st.info("Создайте новый прогон, чтобы сохранять результаты извлечения.")
        st.stop()

    _run_id = int(st.session_state.tag_run_id)

    _pos = int(st.session_state.get('tag_topic_pos', 0))
    _total = len(_topics_df)
    _pos = max(0, min(_pos, max(_total - 1, 0)))
    st.session_state.tag_topic_pos = _pos

    _cur_topic = _topics_df.iloc[_pos].to_dict()
    _cur_topic_id = int(_cur_topic['id'])

    def _get_saved_terms_count(run_id: int, frp_topic_id: int) -> int:
        conn = get_db_conn()
        if conn is None:
            return 0
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT COUNT(*) FROM tag_topic_terms WHERE run_id=%s AND frp_topic_id=%s",
                    (int(run_id), int(frp_topic_id)),
                )
                cnt = int(cur.fetchone()[0])
            conn.close()
            return cnt
        except Exception:
            try:
                conn.close()
            except Exception:
                pass
            return 0

    _saved_terms_cnt = _get_saved_terms_count(_run_id, _cur_topic_id)

    _pcol, _mcol, _bcol1, _bcol2, _bcol3 = st.columns([6, 2, 2, 2, 2])
    with _pcol:
        st.progress((_pos) / max(_total, 1), text=f"Тема {_pos + 1} / {_total}: id={_cur_topic_id} — {_cur_topic.get('section','')} / {_cur_topic.get('topic','')}")
    with _mcol:
        st.metric("Сохранено терминов", _saved_terms_cnt)
    with _bcol1:
        if st.button("⬅️ Назад", key='tag_prev_topic', disabled=_pos <= 0, use_container_width=True):
            st.session_state.tag_topic_pos = _pos - 1
            st.session_state.tag_last_result = None
            st.rerun()
    with _bcol2:
        if st.button("➡️ Вперёд", key='tag_next_topic', disabled=_pos >= _total - 1, use_container_width=True):
            st.session_state.tag_topic_pos = _pos + 1
            st.session_state.tag_last_result = None
            st.rerun()
    with _bcol3:
        if st.button("⛔ Стоп", key='tag_stop_btn', use_container_width=True):
            st.session_state.tag_stop = True
            st.warning("Остановлено. Вы можете продолжить позже, нажмите «Извлечь теги» на нужной теме.")

    def _extract_json_payload(text: str) -> Optional[Dict]:
        """
        Извлекает JSON из ответа модели и отрезает лишний текст.
        Умеет:
        - снимать ```json fences
        - доставать самый большой {...} или [...] блок
        - пытаться дописать закрывающие скобки/кавычки (как в parse_llm_response)
        """
        if not text:
            return None
        t = str(text).strip()

        # markdown fences
        if "```" in t:
            parts = t.split("```")
            if len(parts) >= 3:
                inner = parts[1]
                lines = inner.split("\n")
                if lines and lines[0].strip().lower() in ("json", "javascript", ""):
                    inner = "\n".join(lines[1:])
                t = inner.strip()
            else:
                t = max(parts, key=len).strip()

        def _try_parse(s: str):
            s = s.strip()
            try:
                return json.loads(s)
            except Exception:
                pass
            for closing in ("]", "]}", "}]",
                            "}", "}}",
                            "\"]", "\"}", "\"}]"):
                try:
                    return json.loads(s + closing)
                except Exception:
                    pass
            last_brace = s.rfind("},")
            if last_brace == -1:
                last_brace = s.rfind("}")
            if last_brace > 0:
                candidate = s[: last_brace + 1]
                for wrap in ("", "]", "}"):
                    try:
                        return json.loads(candidate + wrap)
                    except Exception:
                        pass
            return None

        # предпочитаем объект {...} (по нашему контракту), но если его нет — пробуем массив
        obj_matches = re.findall(r"\{[\s\S]*\}", t)
        arr_matches = re.findall(r"\[[\s\S]*\]", t)

        candidates = []
        if obj_matches:
            candidates.extend(obj_matches)
        if arr_matches:
            candidates.extend(arr_matches)
        if not candidates:
            parsed = _try_parse(t)
            return parsed if isinstance(parsed, dict) else None

        # берём самый большой блок (обычно это “правильный” JSON)
        candidates.sort(key=len, reverse=True)
        for cand in candidates[:3]:
            parsed = _try_parse(cand)
            if isinstance(parsed, dict):
                return parsed
        return None

    def _load_topic_records_dedup(topic_id: int) -> Tuple[List[Dict], Dict[str, List[Tuple[str, int]]], Dict[Tuple[str, int], str], int]:
        """
        Возвращает:
        - records_for_llm: уникальные по frp_label (по normalize_db_text), чтобы модель не анализировала одно и то же много раз
        - norm_to_sources: norm_text -> список (source_table, id) всех записей с этим текстом
        - proto_to_norm: (source_table, id) прототипа -> norm_text (для обратного маппинга ответа модели)
        - total_raw_records: сколько было записей ДО дедупликации (для UI)
        """
        conn = get_db_conn()
        if conn is None:
            return [], {}, {}, 0
        try:
            skills = pd.read_sql(
                "SELECT id, frp_label FROM skill_defs WHERE frp_topic_id = %s AND frp_label IS NOT NULL AND btrim(frp_label) <> '' ORDER BY id",
                conn,
                params=(topic_id,),
            )
            content = pd.read_sql(
                "SELECT id, frp_label FROM content_element_defs WHERE frp_topic_id = %s AND frp_label IS NOT NULL AND btrim(frp_label) <> '' ORDER BY id",
                conn,
                params=(topic_id,),
            )
            conn.close()
        except Exception:
            try:
                conn.close()
            except Exception:
                pass
            return [], {}, {}, 0

        raw_rows: List[Tuple[str, int, str]] = []
        for _, r in skills.iterrows():
            raw_rows.append(("skill_defs", int(r["id"]), str(r["frp_label"])))
        for _, r in content.iterrows():
            raw_rows.append(("content_element_defs", int(r["id"]), str(r["frp_label"])))

        total_raw = len(raw_rows)

        # norm_text -> prototype (source_table, id, text)
        prototypes: Dict[str, Tuple[str, int, str]] = {}
        norm_to_sources: Dict[str, List[Tuple[str, int]]] = {}

        for src_table, src_id, txt in raw_rows:
            norm = normalize_db_text(txt)
            if not norm:
                continue
            norm_to_sources.setdefault(norm, []).append((src_table, src_id))
            if norm not in prototypes:
                prototypes[norm] = (src_table, src_id, txt.strip())

        records_for_llm: List[Dict] = []
        proto_to_norm: Dict[Tuple[str, int], str] = {}
        for norm, (src_table, src_id, txt) in prototypes.items():
            records_for_llm.append({"source_table": src_table, "id": int(src_id), "text": txt})
            proto_to_norm[(src_table, int(src_id))] = norm

        return records_for_llm, norm_to_sources, proto_to_norm, total_raw

    def _save_topic_terms(run_id: int, subject_id: int, frp_topic_id: int, terms: List[str]) -> None:
        conn = get_db_conn()
        if conn is None:
            raise RuntimeError("Нет подключения к БД")
        with conn.cursor() as cur:
            for term in terms:
                t = normalize_db_text(term)
                if not t:
                    continue
                cur.execute(
                    """
                    INSERT INTO tag_topic_terms(run_id, subject_id, frp_topic_id, term)
                    VALUES (%s, %s, %s, %s)
                    ON CONFLICT (run_id, frp_topic_id, term_norm) DO NOTHING
                    """,
                    (run_id, subject_id, frp_topic_id, t),
                )
        conn.commit()
        conn.close()

    if st.session_state.get('tag_stop'):
        st.info("Остановлено. Снимите «Стоп» просто продолжив обработку вручную на нужной теме.")

    _act1, _act2, _act3 = st.columns([3, 3, 6])
    with _act1:
        if st.button("✨ Извлечь теги (эта тема)", key='tag_extract_btn', use_container_width=True):
            records, norm_to_sources, proto_to_norm, total_raw = _load_topic_records_dedup(_cur_topic_id)
            if not records:
                st.warning("В этой теме нет записей skill_defs/content_element_defs с frp_label.")
            else:
                st.info(
                    f"Отправляю в модель: **{len(records)}** уникальных формулировок frp_label "
                    f"(всего записей в теме: **{total_raw}**)."
                )
                payload = {
                    "subject": _sel_subj,
                    "frp_topic": {
                        "id": _cur_topic_id,
                        "grade_class": str(_cur_topic.get("grade_class", "")),
                        "program": str(_cur_topic.get("program", "")),
                        "section": str(_cur_topic.get("section", "")),
                        "topic": str(_cur_topic.get("topic", "")),
                    },
                    "records": records,
                }
                prompt = load_tag_prompt(1)
                msgs = [{
                    "role": "user",
                    # без indent, чтобы не раздувать запрос
                    "content": f"{prompt}\n\nВХОДНЫЕ ДАННЫЕ (JSON):\n{json.dumps(payload, ensure_ascii=False, separators=(',', ':'))}"
                }]
                with st.spinner("Запрос к модели..."):
                    resp = call_claude_api(msgs)
                    _accumulate_cost_for_tag_run(_run_id)
                parsed = _extract_json_payload(resp or "")
                if not parsed or "items" not in parsed:
                    st.error("Не удалось разобрать ответ модели. Попробуйте ещё раз или поменяйте модель.")
                else:
                    items = parsed.get("items", [])
                    # лёгкая валидация
                    ok_items: List[Dict] = []
                    expected_proto_keys = set(proto_to_norm.keys())
                    matched_proto_keys = set()
                    unexpected_items = 0
                    for it in items:
                        if not isinstance(it, dict):
                            continue
                        if it.get("source_table") not in ("skill_defs", "content_element_defs"):
                            continue
                        if "id" not in it:
                            continue
                        tags = it.get("tags", [])
                        if not isinstance(tags, list) or not tags:
                            continue

                        src_table = str(it["source_table"])
                        src_id = int(it["id"])
                        norm = proto_to_norm.get((src_table, src_id))
                        if not norm:
                            # если модель вернула неожиданный id/таблицу — пропускаем
                            unexpected_items += 1
                            continue
                        matched_proto_keys.add((src_table, src_id))

                        cleaned_tags = [normalize_db_text(x) for x in tags if normalize_db_text(x)]
                        if not cleaned_tags:
                            continue

                        # расширяем на ВСЕ записи с таким же frp_label (дедуп по norm)
                        for dst_table, dst_id in norm_to_sources.get(norm, []):
                            ok_items.append({
                                "source_table": dst_table,
                                "id": int(dst_id),
                                "tags": cleaned_tags,
                            })

                    if not ok_items:
                        st.error("Ответ модели не содержит корректных items.")
                    else:
                        try:
                            # Требование: сохраняем ТОЛЬКО "тема -> термины" (каждый термин отдельной строкой)
                            topic_terms_set = set()
                            for it in ok_items:
                                for tg in it.get("tags", []):
                                    nt = normalize_db_text(tg)
                                    if nt:
                                        topic_terms_set.add(nt)
                            topic_terms = sorted(topic_terms_set)
                            _save_topic_terms(_run_id, _sel_subject_id, _cur_topic_id, topic_terms)
                            st.session_state.tag_last_result = [{"term": t} for t in topic_terms]
                            st.session_state.tag_last_topic_id = _cur_topic_id
                            st.success(
                                f"Сохранено терминов по теме: {len(topic_terms)}. "
                                f"Уникальных frp_label отправлено в модель: {len(records)}."
                            )
                            missing = len(expected_proto_keys - matched_proto_keys)
                            covered = len(matched_proto_keys)
                            if missing > 0 or unexpected_items > 0:
                                st.warning(
                                    f"Покрытие ответа: {covered}/{len(expected_proto_keys)} уникальных frp_label. "
                                    f"Не покрыто: {missing}. Неожиданных items: {unexpected_items}."
                                )
                        except Exception as _e:
                            st.error(f"Ошибка сохранения в БД: {_e}")

    with _act2:
        if st.button("🧹 Очистить результаты (эта тема)", key='tag_clear_topic_btn', use_container_width=True):
            _conn = get_db_conn()
            if _conn is None:
                st.error("Нет подключения к БД.")
            else:
                try:
                    with _conn.cursor() as _cur:
                        _cur.execute(
                            "DELETE FROM tag_topic_terms WHERE run_id=%s AND frp_topic_id=%s",
                            (_run_id, _cur_topic_id),
                        )
                    _conn.commit()
                    _conn.close()
                    st.session_state.tag_last_result = None
                    st.success("Удалены термины для этой темы (только для текущего run_id).")
                except Exception as _e:
                    try:
                        _conn.close()
                    except Exception:
                        pass
                    st.error(f"Ошибка очистки: {_e}")

    with _act3:
        st.caption("Совет: если тема слишком большая и ответ не влезает — позже добавим автоматическое разбиение на чанки.")

    if st.session_state.get('tag_last_result') and st.session_state.get('tag_last_topic_id') == _cur_topic_id:
        st.markdown("**Последний результат (эта тема):**")
        st.dataframe(pd.DataFrame(st.session_state.tag_last_result), use_container_width=True, hide_index=True)


# ============ РЕЖИМ: ПРОСМОТР БАЗЫ ДАННЫХ ============
elif mode == 'view_db':
    st.header("📋 Просмотр базы данных")

    _vdb_url = os.environ.get('DATABASE_URL', '')
    if not _vdb_url:
        st.error("DATABASE_URL не задан.")
        st.stop()

    # ── Выбор типа данных ─────────────────────────────────────────────────────
    _vt1, _vt2, _ = st.columns([2, 2, 8])
    _v_skills  = st.session_state.vdb_type == 'skills'
    _v_content = st.session_state.vdb_type == 'content'
    with _vt1:
        if st.button("📚 Навыки", type='primary' if _v_skills else 'secondary',
                     use_container_width=True, key='vdb_tab_skills'):
            st.session_state.vdb_type = 'skills'
            st.session_state.vdb_df   = None
            st.session_state.vdb_reassign = False
            st.rerun()
    with _vt2:
        if st.button("📄 Содержание", type='primary' if _v_content else 'secondary',
                     use_container_width=True, key='vdb_tab_content'):
            st.session_state.vdb_type = 'content'
            st.session_state.vdb_df   = None
            st.session_state.vdb_reassign = False
            st.rerun()

    if not st.session_state.vdb_type:
        st.info("Выберите тип данных: Навыки или Содержание.")
        st.stop()

    # ── Загрузка данных ───────────────────────────────────────────────────────
    _v_table = 'skill_defs' if st.session_state.vdb_type == 'skills' else 'content_element_defs'
    _v_typename = 'навыки' if st.session_state.vdb_type == 'skills' else 'элементы содержания'

    if st.session_state.vdb_df is None:
        with st.spinner("Загружаю данные..."):
            st.session_state.vdb_df = load_view_data_cached(_v_table)

    _vdf: pd.DataFrame = st.session_state.vdb_df
    if _vdf.empty:
        st.warning("В базе нет данных.")
        if st.button("🔄 Обновить", key='vdb_reload'):
            load_view_data_cached.clear()
            st.session_state.vdb_df = None
            st.rerun()
        st.stop()

    # ── Фильтры (каскадные) ───────────────────────────────────────────────────
    _fc1, _fc2, _fc3, _fc4, _fc5 = st.columns([2, 1, 3, 3, 1])

    _vf_subj = _fc1.selectbox(
        "Предмет", [''] + sorted(_vdf['subject'].unique()), key='vdb_f_subj'
    )
    _fdf = _vdf if not _vf_subj else _vdf[_vdf['subject'] == _vf_subj]

    _vf_class = _fc2.selectbox(
        "Класс",
        [''] + sorted(_fdf['grade_class'].unique(), key=lambda x: int(x) if x.isdigit() else 99),
        key='vdb_f_class'
    )
    if _vf_class:
        _fdf = _fdf[_fdf['grade_class'] == _vf_class]

    _vf_sect = _fc3.selectbox(
        "Раздел", [''] + sorted(_fdf['section'].unique()), key='vdb_f_sect'
    )
    if _vf_sect:
        _fdf = _fdf[_fdf['section'] == _vf_sect]

    _vf_topic = _fc4.selectbox(
        "Тема", [''] + sorted(_fdf['topic'].unique()), key='vdb_f_topic'
    )
    if _vf_topic:
        _fdf = _fdf[_fdf['topic'] == _vf_topic]

    with _fc5:
        st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
        if st.button("🔄", key='vdb_reload2', help="Обновить данные"):
            load_view_data_cached.clear()
            st.session_state.vdb_df = None
            st.rerun()

    st.caption(f"Показано: {len(_fdf)} записей ({_v_typename})")

    # ── Кнопка "Сменить раздел/тему" ─────────────────────────────────────────
    _v_ra = st.session_state.vdb_reassign
    _ra_c1, _ra_c2 = st.columns([3, 9])
    with _ra_c1:
        if st.button(
            "✏️ Сменить раздел/тему" if not _v_ra else "✖ Отмена",
            key='vdb_ra_toggle', type='secondary'
        ):
            st.session_state.vdb_reassign = not _v_ra
            st.rerun()
    if _v_ra:
        with _ra_c2:
            st.info("Отметьте записи галочками, затем выберите новую тему ниже.")

    st.markdown("---")

    # ── Список записей ────────────────────────────────────────────────────────
    if _fdf.empty:
        st.warning("По выбранным фильтрам ничего не найдено.")
    else:
        _prev_subj  = None
        _prev_grade = None
        _prev_sect  = None
        _grp_map    = {}  # grp_key → [item_ids] — для режима переназначения

        for _, _row_group in _fdf.groupby(
            ['subject', 'grade_class', 'section', 'topic'],
            sort=False
        ):
            # _row_group is a DataFrame for one (subject, grade, section, topic)
            _g_subj  = _row_group['subject'].iloc[0]
            _g_grade = _row_group['grade_class'].iloc[0]
            _g_sect  = _row_group['section'].iloc[0]
            _g_topic = _row_group['topic'].iloc[0]

            # Section / subject-grade header
            if (_g_subj, _g_grade) != (_prev_subj, _prev_grade):
                st.markdown(f"## {_g_subj} — класс {_g_grade}")
                _prev_subj  = _g_subj
                _prev_grade = _g_grade
                _prev_sect  = None

            if _g_sect != _prev_sect:
                st.markdown(f"### {_g_sect}")
                _prev_sect = _g_sect

            st.markdown(f"**Тема: {_g_topic}**")

            # Group by frp_label within topic
            for _flabel, _ldf in _row_group.groupby('frp_label', sort=False):
                _items     = _ldf.to_dict('records')
                _item_ids  = [it['id'] for it in _items]
                _grp_key   = f'vdb_chk_g{min(_item_ids)}'
                _is_atomized = not (
                    len(_items) == 1 and _items[0]['label_normalized'] == _flabel
                )
                # Register group → ids mapping for use in reassign panel
                _grp_map[_grp_key] = _item_ids

                if _is_atomized:
                    if _v_ra:
                        st.checkbox(f"📌 {_flabel}", key=_grp_key)
                    else:
                        st.markdown(f"&nbsp;&nbsp;📌 *{_flabel}*", unsafe_allow_html=True)
                    for _it in _items:
                        st.markdown(
                            f"&nbsp;&nbsp;&nbsp;&nbsp;— {_it['label_normalized']}",
                            unsafe_allow_html=True
                        )
                else:
                    if _v_ra:
                        st.checkbox(_flabel, key=_grp_key)
                    else:
                        st.markdown(
                            f"&nbsp;&nbsp;— {_flabel}",
                            unsafe_allow_html=True
                        )

        # Сохраняем актуальную карту групп в session_state
        st.session_state.vdb_grp_map = _grp_map

    # ── Панель переназначения ─────────────────────────────────────────────────
    if _v_ra and not _fdf.empty:
        _grp_map_saved = st.session_state.get('vdb_grp_map', {})

        # Собираем все id из отмеченных групп
        _checked_ids = [
            rid
            for gk, ids in _grp_map_saved.items()
            if st.session_state.get(gk, False)
            for rid in ids
        ]

        st.markdown("---")
        with st.container(border=True):
            _sel_grps = sum(1 for gk in _grp_map_saved if st.session_state.get(gk, False))
            _sel_cnt  = len(_checked_ids)
            st.markdown(f"**Выбрано: {_sel_grps} групп ({_sel_cnt} записей)**")

            # "Выбрать все" / "Снять все"
            _chk_all_c, _unchk_all_c, _ = st.columns([2, 2, 8])
            with _chk_all_c:
                if st.button("☑ Выбрать все", key='vdb_check_all'):
                    for gk in _grp_map_saved:
                        st.session_state[gk] = True
                    st.rerun()
            with _unchk_all_c:
                if st.button("☐ Снять все", key='vdb_uncheck_all'):
                    for gk in _grp_map_saved:
                        st.session_state[gk] = False
                    st.rerun()

            if _sel_cnt > 0:
                st.markdown("**Новая тема:**")
                if st.session_state.get('db_frp_df') is None:
                    st.session_state.db_frp_df = load_frp_topics_cached()
                _frp_r = st.session_state.db_frp_df

                _nc1, _nc2, _nc3, _nc4 = st.columns(4)
                _nr_subj = _nc1.selectbox(
                    "Предмет", [''] + sorted(_frp_r['subject'].unique()), key='vdb_r_subj'
                )
                _nrdf = _frp_r if not _nr_subj else _frp_r[_frp_r['subject'] == _nr_subj]

                _nr_class = _nc2.selectbox(
                    "Класс",
                    [''] + sorted(_nrdf['grade_class'].unique(),
                                  key=lambda x: int(x) if x.isdigit() else 99),
                    key='vdb_r_class'
                )
                if _nr_class:
                    _nrdf = _nrdf[_nrdf['grade_class'] == _nr_class]

                _nr_sect = _nc3.selectbox(
                    "Раздел", [''] + sorted(_nrdf['section'].unique()), key='vdb_r_sect'
                )
                if _nr_sect:
                    _nrdf = _nrdf[_nrdf['section'] == _nr_sect]

                _nr_topic = _nc4.selectbox(
                    "Тема", [''] + sorted(_nrdf['topic'].unique()), key='vdb_r_topic'
                )

                if _nr_subj and _nr_class and _nr_sect and _nr_topic:
                    _new_row = _frp_r[
                        (_frp_r['subject']     == _nr_subj) &
                        (_frp_r['grade_class'] == _nr_class) &
                        (_frp_r['section']     == _nr_sect) &
                        (_frp_r['topic']       == _nr_topic)
                    ]
                    if not _new_row.empty:
                        _new_frp_id = int(_new_row.iloc[0]['id'])
                        st.caption(
                            f"→ {_nr_subj}, класс {_nr_class}, раздел «{_nr_sect}», тема «{_nr_topic}»"
                        )
                        if st.button(
                            f"✅ Переназначить {_sel_grps} групп ({_sel_cnt} записей)",
                            type='primary', key='vdb_apply_ra'
                        ):
                            _ra_conn = get_db_conn()
                            if _ra_conn:
                                try:
                                    with _ra_conn.cursor() as _ra_cur:
                                        _ra_cur.execute(
                                            f"UPDATE {_v_table} "
                                            f"SET frp_topic_id = %s WHERE id = ANY(%s)",
                                            (_new_frp_id, _checked_ids)
                                        )
                                    _ra_conn.commit()
                                    _ra_conn.close()
                                    st.success(
                                        f"✅ Обновлено {_sel_cnt} записей → "
                                        f"{_nr_sect} / {_nr_topic}"
                                    )
                                    load_view_data_cached.clear()
                                    st.session_state.vdb_df = None
                                    st.session_state.vdb_grp_map = {}
                                    st.session_state.vdb_reassign = False
                                    st.rerun()
                                except Exception as _ra_e:
                                    st.error(f"Ошибка обновления: {_ra_e}")
                            else:
                                st.error("Нет подключения к БД.")
            else:
                st.caption("Отметьте хотя бы одну группу в списке выше.")
